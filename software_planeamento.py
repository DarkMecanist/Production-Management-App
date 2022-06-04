import kivy
import sqlite3
import math
import string
import datetime
import time as tm
import re, fitz, os
import tkinter as tk
from tkinter import filedialog
import docx
from docx import Document
from docx.shared import Cm, Mm
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from PIL import Image, ImageDraw, ImageFont
from fpdf import FPDF
from itertools import combinations, cycle
import xlsxwriter as xls
import operator
import pandas as pd

kivy.require('1.11.0')
from kivy.config import Config
Config.set('graphics', 'window_state', 'maximized')
Config.set('graphics', 'fullscreen', 0)
Config.set('input', 'mouse', 'mouse,disable_multitouch')
from kivy.uix.label import Label
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.button import Button
from kivy.uix.scrollview import ScrollView
from kivy.uix.checkbox import CheckBox
from kivy.uix.textinput import TextInput
from kivy.uix.popup import Popup
from kivy.uix.dropdown import DropDown
from kivy.uix.screenmanager import ScreenManager, Screen, NoTransition
from kivy.core.window import Window
from kivy.app import App


def return_formatted_datetime(datetime_obj):
    return f'{datetime_obj.hour}:{datetime_obj.minute} - {datetime_obj.day}/{datetime_obj.month}/{datetime_obj.year}'


def round_time(time_obj):
    """Rounds time to 15 minute intervals. Ex 9:11 → 9:15"""

    if time_obj.minute > 0 and time_obj.minute < 15:
        new_minutes = 15
        time_obj = datetime.datetime(year=time_obj.year, month=time_obj.month, day=time_obj.day, hour=time_obj.hour, minute=new_minutes)
    elif time_obj.minute > 15 and time_obj.minute < 30:
        new_minutes = 30
        time_obj = datetime.datetime(year=time_obj.year, month=time_obj.month, day=time_obj.day, hour=time_obj.hour, minute=new_minutes)
    elif time_obj.minute > 30 and time_obj.minute < 45:
        new_minutes = 45
        time_obj = datetime.datetime(year=time_obj.year, month=time_obj.month, day=time_obj.day, hour=time_obj.hour, minute=new_minutes)
    elif time_obj.minute > 45:
        new_minutes = 0
        if time_obj.hour == 23:
            new_hour = 0
        else:
            new_hour = time_obj.hour + 1

            time_obj = datetime.datetime(year=time_obj.year, month=time_obj.month, day=time_obj.day, hour=new_hour, minute=new_minutes)

    return time_obj


def return_times_between_times(time_start, time_finish, interval_step):
    """Returns all the times, ignoring weekend days, equally spaced, includes start time but excludes finish time"""
    times_list = []
    # print(f'LATEST DATE IS {latest_date}')

    # if time_start.day != latest_date.day or time_start.month != latest_date.month or time_start.year != latest_date.year:
    #     time_start = datetime.datetime(year=latest_date.year, month=latest_date.month, day=latest_date.day, hour=time_start.hour, minute=time_start.minute)
    #
    # if time_finish.day != latest_date.day or time_finish.month != latest_date.month or time_finish.year != latest_date.year:
    #     time_finish = datetime.datetime(year=latest_date.year, month=latest_date.month, day=latest_date.day, hour=time_finish.hour, minute=time_finish.minute)

    duration = return_duration_between_dates(time_start, time_finish)
    num_times = round((duration / interval_step))

    # print(f'RETURNING TIMES BETWEEN {str(time_start)} and {str(time_finish)}')
    for _ in range(num_times):
        times_list.append(time_start)
        time_start = return_next_weekday(time_start + interval_step)

    # print(times_list)
    # print('___________________________________________________')
    return times_list


def return_duration_between_dates(date_start, date_finish):
    duration = date_finish - date_start

    if duration.days == 2:
        return duration - datetime.timedelta(days=2)

    return duration


def return_next_weekday(datetime_obj):
    '''checks if an day object is a week day. If it isn't, it outputs the next week day'''

    while datetime_obj.weekday() > 4:
        datetime_obj += datetime.timedelta(days=1)

    return datetime_obj


def check_date_is_valid(date_string):
    valid = False

    try:
        date_info = [int(elem) for elem in date_string.split('/')]
        if len(date_info) == 3:
            datetime.datetime(day=date_info[0], month=date_info[1], year=date_info[2])

            valid = True
        else:
            return valid
    except ValueError:
        pass

    return valid


def return_formatted_time_string(time_mins):
    minutes = time_mins

    if time_mins < 60:
        return f'{minutes} min'
    else:
        hours = int(str(minutes / 60).split('.')[0])
        minutes = round(float('0.' + str(minutes / 60).split('.')[1]) * 60)

        if minutes < 10:
            return f'{hours}:0{minutes} h'
        else:
            return f'{hours}:{minutes} h'


def return_dictionary_of_two_letter_combinations():
    two_letter_combo_dict = {}

    combo_1 = combinations(list(string.ascii_uppercase), 1)
    combo_2 = combinations(list(string.ascii_uppercase), 2)

    list_1 = [','.join(comb) for comb in combo_1]
    list_2 = [','.join(comb) for comb in combo_2]

    combo_list = list_1 + list_2

    for combo in combo_list:
        two_letter_combo_dict.update({combo: ''})

    return two_letter_combo_dict


def round_number_if_integer(number):
    if str(number)[-2:] == '.0':
        number = int(number)

    return number


def remove_invalid_file_name_characters(filename_string):
    invalid_characters = ['<', '>', ':', '"', '/', '\\', '|', '?', '*']

    for invalid_character in invalid_characters:
        filename_string = filename_string.replace(invalid_character, '')

    return filename_string


def wrapper_function(list, type):
    new_list = []

    if type == 'Part Name':
        for item in list:
            new_list.append(item.replace('\n', ''))
    elif type == 'Part Qtty':
        for item in list:
            # print(f'Previous quantity: {item}')
            new_item = item.replace(' ', '').replace('.', '').replace(',', '').replace('\n', '')
            # print(f'Previous quantity: {int(new_item)/10}')

            new_list.append(int(new_item)/10)

            # if new_item[-3:] == '000':
            #     new_item = new_item[:-3]
            #     new_list.append(str(float(new_item)))
            # else:
            #     new_item = new_item[:-3] + '.' + new_item[-3:]
            #     new_list.append(new_item)
    return new_list


def get_search_pattern(type):
    if type == 'Client':
        return r'Exmo\.\(s\) Sr.\(s\)\n(.+)\n', 1
    elif type == 'Order':
        return r'\b(.+)\nNº Requisição', 1
    elif type == 'Part Ref':
        return r'(\b[0-9]{8}|\b[A-Z0-9]{12}|\b1)\n(.+|.+\n.+|.+\n.+\n.+)\nUN', 1 #Part Ref #r'(\b[A-Z0-9]{12}\n|\b1\n)'
    elif type == 'Part Name':
        return r'(\b[0-9]{8}|\b[A-Z0-9]{12}|\b1)\n(.+|.+\n.+|.+\n.+\n.+)\nUN', 2 #Part Name r'\b(.+|.+\n.+|.+\n.+\n.+)\nUN'
    elif type == 'Part Qtty':
        return r'\bUN\n ([0-9]+,[0-9]+)', 1 #Part qtty
    elif type == 'Part Unit':
        return r'\b(UN)\n [0-9]+,[0-9]+', 1 #Part unit
    elif type == 'Part Material':
        return r'\b(UN)\n [0-9]+,[0-9]+\n(.+)', 2 #Part material


def extract_data(type, text_to_search):
    search_pattern = get_search_pattern(type)

    search = re.compile(search_pattern[0])
    results = search.finditer(text_to_search)

    results_list = []

    for result in results:
        results_list.append(result.group(search_pattern[1]))


    return results_list


def identify_PDF_type(filepath):
    '''Returns what type of PDF, a given entry file is (Part LC5, Sheet LC5 or Sheet LF-3015)'''
    doc = fitz.Document(filepath)
    page = doc.loadPage()
    text = page.getText('text')

    if text[:15] == 'Exmo.(s) Sr.(s)':
        type = 'Order Prilux'
        return type
    else:
        print('File Matches no known type')


def return_order_parts_from_PDF():
    root = tk.Tk()
    root.withdraw()
    path_list = filedialog.askopenfilenames()
    path_dict = {}
    task_list = []
    for path in path_list:
        doc = fitz.Document(path)
        PDF_type = identify_PDF_type(path)

        if PDF_type == 'Order Prilux':
            print('INSIDE GENERATE ORDER FROM PDF FUNC')

            num_pages = doc.pageCount

            full_text = ''
            for num in range(num_pages):
                page = doc.loadPage(num)

                text = page.getText('text')

                if 'Original' in text:
                    full_text += text

            print('__________________________________________________________________________________________________')
            print(full_text)
            print('__________________________________________________________________________________________________')
            client = list(set(extract_data('Client', full_text)))
            order = list(set(extract_data('Order', full_text)))
            part_refs = extract_data('Part Ref', full_text)
            part_names = extract_data('Part Name', full_text)
            part_qttys = extract_data('Part Qtty', full_text)
            part_units = extract_data('Part Unit', full_text)
            part_materials = extract_data('Part Material', full_text)

            print(f'LENGTH OF PART REFS: {len(part_refs)}')
            print(f'LENGTH OF PART NAMES: {len(part_names)}')
            print(f'LENGTH OF PART QTTYS: {len(part_qttys)}')
            print(f'LENGTH OF PART UNITS: {len(part_units)}')
            print(f'LENGTH OF PART MATERIALS: {len(part_materials)}')

            part_names = wrapper_function(part_names, 'Part Name')
            part_qttys = wrapper_function(part_qttys, 'Part Qtty')

            # print(path)
            # print(client)
            # print(order)
            # print(f'{part_refs} Num elems: {len(part_refs)}')
            # print(f'{part_names} Num elems: {len(part_names)}')
            # print(f'{part_qttys} Num elems: {len(part_qttys)}')
            # print(f'{part_units} Num elems: {len(part_units)}')

            #Creating the temporary dict
            path_dict[path] = {'client': client, 'order_num': order, 'part_refs': part_refs,
                                'part_names': part_names, 'part_qttys': part_qttys, 'part_units': part_units,
                               'part_materials': part_materials}


            for result in path_dict:
                # print(path_dict[result])
                # print((path_dict[result]['client'], len(path_dict[result]['client'])))
                # print((path_dict[result]['order_num'], len(path_dict[result]['order_num'])))
                # print((path_dict[result]['part_refs'], len(path_dict[result]['part_refs'])))
                # print((path_dict[result]['part_names'], len(path_dict[result]['part_names'])))
                # print((path_dict[result]['part_qttys'], len(path_dict[result]['part_qttys'])))

                for index in range(len(path_dict[result]['part_refs'])):

                    task = (path_dict[result]['client'][0], path_dict[result]['order_num'][0], path_dict[result]['part_refs'][index],
                            path_dict[result]['part_names'][index], path_dict[result]['part_qttys'][index], path_dict[result]['part_materials'][index])
                    task_list.append(task)

    return task_list


def return_formatted_time(time):
    return tm.strftime('%H:%M:%S', tm.gmtime(time))


def load_tasks_from_database(rowid=None):
    if rowid == None:
        cursor.execute('select *, rowid from tasks order by priority')
    else:
        cursor.execute('''select * from tasks where rowid = :rowid''', {'rowid': rowid})
    return cursor.fetchall()


def insert_new_task_database(current_path, original_path, machine, material_ref, notes, estimated_sheets_required, estimated_time, start_date, end_date, priority, order_parts, aggregated_tasks, aggregated_index='', auto_fill_rowid=True):
    cursor.execute('insert into tasks values(:current_path, :original_path, :machine, :material_ref, :notes, :estimated_sheets_required, :estimated_time, :start_date, :end_date, :priority, :order_parts, :aggregated_tasks, :aggregated_index)',
                   {'current_path': current_path,
                    'original_path': original_path,
                    'machine': machine,
                    'material_ref': material_ref,
                    'notes': notes,
                    'estimated_sheets_required': estimated_sheets_required,
                    'estimated_time': estimated_time,
                    'start_date': start_date,
                    'end_date': end_date,
                    'priority': priority,
                    'order_parts': order_parts,
                    'aggregated_tasks': aggregated_tasks,
                    'aggregated_index': aggregated_index
                    })

    if auto_fill_rowid:
        tasks_list = load_tasks_from_database()
        for task in tasks_list:
            if current_path == task[0] and original_path == task[1] and machine == task[2] and material_ref == task[3] and order_parts == task[10]:
                change_value_task_database(str(task[-1]), 'aggregated_tasks', task[-1])
                print(f'Sucessfully changed rowid upon task creation: {task[-1]}')


def remove_task_database(rowid):
    with conn:
        cursor.execute('delete from tasks where rowid = :rowid', {'rowid': rowid})


def change_value_task_database(value, field, rowid):
    with conn:
        if field == 'current_path':
            cursor.execute('update tasks set current_path = :current_path where rowid = :rowid', {'current_path': value, 'rowid': rowid})
        elif field == 'original_path':
            cursor.execute('update tasks set original_path = :original_path where rowid = :rowid', {'original_path': value, 'rowid': rowid})
        elif field == 'machine':
            cursor.execute('update tasks set machine = :machine where rowid = :rowid', {'machine': value, 'rowid': rowid})
        elif field == 'material_ref':
            cursor.execute('update tasks set material_ref = :material_ref where rowid = :rowid', {'material_ref': value, 'rowid': rowid})
        elif field == 'notes':
            cursor.execute('update tasks set notes = :notes where rowid = :rowid', {'notes': value, 'rowid': rowid})
        elif field == 'estimated_sheets_required':
            cursor.execute('update tasks set estimated_sheets_required = :estimated_sheets_required where rowid = :rowid', {'estimated_sheets_required': value, 'rowid': rowid})
        elif field == 'estimated_time':
            cursor.execute('update tasks set estimated_time = :estimated_time where rowid = :rowid', {'estimated_time': value, 'rowid': rowid})
        elif field == 'start_date':
            cursor.execute('update tasks set start_date = :start_date where rowid = :rowid', {'start_date': value, 'rowid': rowid})
        elif field == 'end_date':
            cursor.execute('update tasks set end_date = :end_date where rowid = :rowid', {'end_date': value, 'rowid': rowid})
        elif field == 'priority':
            cursor.execute('update tasks set priority = :priority where rowid = :rowid', {'priority': value, 'rowid': rowid})
        elif field == 'order_parts':
            cursor.execute('update tasks set order_parts = :order_parts where rowid = :rowid', {'order_parts': value, 'rowid': rowid})
        elif field == 'aggregated_tasks':
            cursor.execute('update tasks set aggregated_tasks = :aggregated_tasks where rowid = :rowid', {'aggregated_tasks': value, 'rowid': rowid})
        elif field == 'aggregated_index':
            cursor.execute('update tasks set aggregated_index = :aggregated_index where rowid = :rowid', {'aggregated_index': value, 'rowid': rowid})


def load_order_parts_from_database(mode='all', material_ref=None, rowid=None, *args):
    if mode == 'all':
        cursor.execute('''select *, rowid from order_parts order by client, order_num_client, name''')
        return cursor.fetchall()
    elif mode == 'new_task':
        cursor.execute('''select name, rowid from order_parts where material_ref = :material_ref order by client''', {'material_ref': material_ref})
        return cursor.fetchall()
    elif mode == 'rowid':
        cursor.execute('''select * from order_parts where rowid = :rowid''', {'rowid': rowid})
        return cursor.fetchall()
    elif mode == 'produced_quantity':
        cursor.execute('select produced_quantity from order_parts where rowid = :rowid', {'rowid': rowid})
        return cursor.fetchone()


def insert_new_order_part_database(ref, name, material_ref, quantity, produced_quantity, order_num, order_num_client, client, date_modified, due_date, additional_operations, *args):
    cursor.execute('insert into order_parts values(:ref, :name, :material_ref, :quantity, :produced_quantity, :order_num, :order_num_client, :client, :date_modified, :due_date, :additional_operations)',
                     {'ref': ref,
                      'name': name,
                      'material_ref': material_ref,
                      'quantity': quantity,
                      'produced_quantity': produced_quantity,
                      'order_num': order_num,
                      'order_num_client': order_num_client,
                      'client': client,
                      'date_modified': date_modified,
                      'due_date': due_date,
                      'additional_operations': additional_operations
                      })


def remove_order_part_database(rowid):
    with conn:
        cursor.execute('''delete from order_parts where rowid = :rowid''', {'rowid': rowid})


def change_value_order_part_database(value, field, rowid):
    with conn:
        if field == 'ref':
            cursor.execute('''update order_parts set ref = :ref where rowid = :rowid''', {'ref': value, 'rowid': rowid})
        elif field == 'name':
            cursor.execute('''update order_parts set name = :name where rowid = :rowid''', {'name': value, 'rowid': rowid})
        elif field == 'material_ref':
            cursor.execute('''update order_parts set material_ref = :material_ref where rowid = :rowid''', {'material_ref': value, 'rowid': rowid})
            print(f'Changing {field} to {rowid}')
        elif field == 'quantity':
            cursor.execute('''update order_parts set quantity = :quantity where rowid = :rowid''', {'quantity': value, 'rowid': rowid})
        elif field == 'produced_quantity':
            cursor.execute('''update order_parts set produced_quantity = :produced_quantity where rowid = :rowid''', {'produced_quantity': value, 'rowid': rowid})
        elif field == 'order_num':
            cursor.execute('''update order_parts set order_num = :order_num where rowid = :rowid''', {'order_num': value, 'rowid': rowid})
        elif field == 'order_num_client':
            cursor.execute('''update order_parts set order_num_client = :order_num_client where rowid = :rowid''', {'order_num_client': value, 'rowid': rowid})
        elif field == 'client':
            cursor.execute('''update order_parts set client = :client where rowid = :rowid''', {'client': value, 'rowid': rowid})
        elif field == 'date_modified':
            cursor.execute('''update order_parts set date_modified = :date_modified where rowid = :rowid''', {'date_modified': value, 'rowid': rowid})
        elif field == 'due_date':
            cursor.execute('''update order_parts set due_date = :due_date where rowid = :rowid''', {'due_date': value, 'rowid': rowid})
        elif field == 'additional_operations':
            cursor.execute('''update order_parts set additional_operations = :additional_operations where rowid = :rowid''', {'additional_operations': value, 'rowid': rowid})


def get_part_info_by_ref(part_ref):
    try:
        cursor.execute('''select weight, material_ref, time, name from parts where ref = :ref''', {'ref': part_ref})
        part_info = cursor.fetchall()[0]
        part_weight = part_info[0]
        part_material_ref = part_info[1]
        part_time = part_info[2]
        name = part_info[3]

        return part_weight, part_material_ref, part_time, name
    except IndexError:
        return 0.0, '', 0, ''


def get_part_info_by_name(part_name):
    try:
        cursor.execute('''select weight, material_ref, time, ref from parts where name = :name''', {'name': part_name})
        part_info = cursor.fetchall()[0]
        part_weight = part_info[0]
        part_material_ref = part_info[1]
        part_time = part_info[2]
        name = part_info[3]

        return part_weight, part_material_ref, part_time, name
    except IndexError:
        return 0.0, '', 0, ''


def load_parts_from_database(*args):
    cursor.execute('''select *, rowid from parts order by name asc''')
    return cursor.fetchall()


def insert_new_part_database(ref, name, weight, material_ref, time, client, date_modified, *args):
    cursor.execute('insert into parts values(:ref, :name, :weight, :material_ref, :time, :client, :date_modified)',
                     {'ref': ref,
                      'name': name,
                      'weight': weight,
                      'material_ref': material_ref,
                      'time': time,
                      'client': client,
                      'date_modified': date_modified
                      })


def remove_part_database(rowid):
    with conn:
        cursor.execute('''delete from parts where rowid = :rowid''', {'rowid': rowid})


def change_value_part_database(value, field, rowid):
    with conn:
        if field == 'ref':
            cursor.execute('''update parts set ref = :ref where rowid = :rowid''', {'ref': value, 'rowid': rowid})
        elif field == 'name':
            cursor.execute('''update parts set name = :name where rowid = :rowid''', {'name': value, 'rowid': rowid})
        elif field == 'weight':
            cursor.execute('''update parts set weight = :weight where rowid = :rowid''', {'weight': value, 'rowid': rowid})
        elif field == 'material_ref':
            cursor.execute('''update parts set material_ref = :material_ref where rowid = :rowid''', {'material_ref': value, 'rowid': rowid})
        elif field == 'time':
            cursor.execute('''update parts set time = :time where rowid = :rowid''', {'time': value, 'rowid': rowid})
        elif field == 'client':
            cursor.execute('''update parts set client = :client where rowid = :rowid''', {'client': value, 'rowid': rowid})
        elif field == 'date_modified':
            cursor.execute('''update parts set date_modified = :date_modified where rowid = :rowid''', {'date_modified': value, 'rowid': rowid})


def get_material_name(ref):
    try:
        cursor.execute('''select type, spec, thickness, client from materials where ref = :ref''', {'ref': ref})
        material = cursor.fetchall()[0]

        # return f'{material[0]} {material[1]}, {material[2]}' Caso tenha problemas voltar a ativar isto
        return f'{material[0]} {material[1]}, {material[2]}mm {material[3]}'
    except IndexError:
        return ref


def get_material_ref_by_name(material_name):
    material_type = material_name.split(' ')[0]
    material_spec = material_name.split(' ')[1].split(',')[0]
    material_thickness = float(material_name.split(' ')[2].split('mm')[0])
    material_client = material_name.split(' ')[3]
    material_ref = get_material_ref(material_type, material_spec, material_thickness, material_client)

    return material_ref


def get_material_ref(type, spec, thickness, client):
    cursor.execute('''select ref from materials where type = :type and spec = :spec and thickness = :thickness and client = :client''', {'type': type, 'spec': spec, 'thickness': thickness, 'client': client})
    return cursor.fetchone()[0]


def load_materials_from_database(material_type=None):
    cursor.execute('select *, rowid from materials order by thickness asc')
    materials_list = cursor.fetchall()

    if material_type == 'ferro':
        return [material for material in materials_list if 'Ferro' in material[1] and material[11] == 'Prilux']
    elif material_type == 'zincado':
        return [material for material in materials_list if ('Zincado' in material[1] or 'Zincor' in material[1]) and material[11] == 'Prilux']
    elif material_type == 'inox':
        return [material for material in materials_list if 'Inox' in material[1] and material[11] == 'Prilux']
    elif material_type == 'alumínio':
        return [material for material in materials_list if 'Alumínio' in material[1] or 'Aluminio' in material[1] and material[11] == 'Prilux']
    elif material_type == 'clients':
        return [material for material in materials_list if material[11] != 'Prilux']
    elif material_type == None:
        return materials_list


def insert_new_material_database(ref, type, spec, thickness, length, width, density, stock_weight, stock_num_sheets, min_stock, date_modified, client, *args):
    cursor.execute('insert into materials values(:ref, :type, :spec, :thickness, :length, :width, :density, :stock_weight, :stock_num_sheets, :min_stock, :date_modified, :client)',
                     {'ref': ref,
                      'type': type,
                      'spec': spec,
                      'thickness': thickness,
                      'length': length,
                      'width': width,
                      'density': density,
                      'stock_weight': stock_weight,
                      'stock_num_sheets': stock_num_sheets,
                      'min_stock': min_stock,
                      'date_modified': date_modified,
                      'client': client
                      })


def remove_material_database(rowid):
    with conn:
        cursor.execute('''delete from materials where rowid = :rowid''', {'rowid': rowid})


def change_value_material_database(value, field, rowid):
    with conn:
        if field == 'ref':
            cursor.execute('''update materials set ref = :ref where rowid = :rowid''', {'ref':value, 'rowid':rowid})
        elif field == 'spec':
            cursor.execute('''update materials set spec = :spec where rowid = :rowid''', {'spec': value, 'rowid': rowid})
        elif field == 'thickness':
            cursor.execute('''update materials set thickness = :thickness where rowid = :rowid''', {'thickness': value, 'rowid': rowid})
        elif field == 'length':
            cursor.execute('''update materials set length = :length where rowid = :rowid''', {'length': value, 'rowid': rowid})
        elif field == 'width':
            cursor.execute('''update materials set width = :width where rowid = :rowid''', {'width': value, 'rowid': rowid})
        elif field == 'stock_weight':
            cursor.execute('''update materials set stock_weight = :stock_weight where rowid = :rowid''', {'stock_weight': value, 'rowid': rowid})
        elif field == 'stock_num_sheets':
            cursor.execute('''update materials set stock_num_sheets = :stock_num_sheets where rowid = :rowid''', {'stock_num_sheets': value, 'rowid': rowid})
        elif field == 'min_stock':
            cursor.execute('''update materials set min_stock = :min_stock where rowid = :rowid''', {'min_stock': value, 'rowid': rowid})
        elif field == 'date_modified':
            cursor.execute('''update materials set date_modified = :date_modified where rowid = :rowid''', {'date_modified': value, 'rowid': rowid})


def load_shifts(machine=None):
    if machine == None:
        cursor.execute('''select * from shifts''')
    else:
        cursor.execute('''select * from shifts where machine = :machine''', {'machine': machine})

    return cursor.fetchall()


def insert_new_shift(machine, time_start, time_finish, time_break, break_duration):
    cursor.execute('''insert into shifts values(:machine, :time_start, :time_finish, :time_break, :break_duration)''',
                   {'machine': machine,
                    'time_start': time_start,
                    'time_finish': time_finish,
                    'time_break': time_break,
                    'break_duration': break_duration})


def remove_shifts(rowid=''):
    if rowid == '':
        with conn:
            cursor.execute('delete from shifts')
    else:
        with conn:
            cursor.execute('delete from shifts where rowid = :rowid', {'rowid': rowid})


def change_value_shifts_database(value, field, rowid):
    with conn:
        if field == 'time_start':
            cursor.execute('update shifts set time_start = :time_start where rowid = :rowid', {'time_start': value, 'rowid': rowid})
        elif field == 'time_finish':
            cursor.execute('update shifts set time_finish = :time_finish where rowid = :rowid', {'time_finish': value, 'rowid': rowid})
        elif field == 'time_break':
            cursor.execute('update shifts set time_break = :time_break where rowid = :rowid', {'time_break': value, 'rowid': rowid})
        elif field == 'break_duration':
            cursor.execute('update shifts set break_duration = :break_duration where rowid = :rowid', {'break_duration': value, 'rowid': rowid})


def convert_rgb_to_kivy_float(rgb_tuple):
    converted_tuple = []

    for color in rgb_tuple:
        converted_tuple.append(color/255)

    converted_tuple.append(1)

    return tuple(converted_tuple)


class Planning(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.shifts_LF = load_shifts('LF3015')
        self.shifts_LC5 = load_shifts('LC5')
        self.added_shifts_LF = []
        self.added_shifts_LC5 = []

        self.orientation = 'vertical'

        current_date = datetime.datetime.now()

        self.label_machine_LF = Label(text='LF-3015', font_size='20')
        self.row_main_info_LF = BoxLayout(orientation='horizontal', size_hint_y=0.3)
        self.label_num_shifts_LF = Label(text=f'Nº Turnos:   {str(len(self.shifts_LF))}    ', size_hint_x=0.3)
        self.label_time_start_plan_LF = Label(text='Hora de Início (HH:MM): ')
        self.ti_time_start_plan_LF = TextInput(text=f'{current_date.hour}:{current_date.minute}', size_hint_x=0.2)
        self.label_date_start_plan_LF = Label(text='Data de Início (DD/MM/AAAA): ')
        self.ti_date_start_plan_LF = TextInput(text=f'{current_date.day}/{current_date.month}/{current_date.year}', size_hint_x=0.3)
        self.button_remove_shift_LF = Button(text='-', background_color=color_dark_red, size_hint_x=0.1)
        self.button_add_shift_LF = Button(text='+', background_color=color_light_green, size_hint_x=0.1)

        self.layout_shifts_LF = BoxLayout(orientation='vertical')

        self.label_machine_LC5 = Label(text='LC5', font_size='20')
        self.row_main_info_LC5 = BoxLayout(orientation='horizontal', size_hint_y=0.3)
        self.label_num_shifts_LC5 = Label(text=f'Nº Turnos:   {str(len(self.shifts_LC5))}    ', size_hint_x=0.3)
        self.label_time_start_plan_LC5 = Label(text='Hora de Início (HH:MM): ')
        self.ti_time_start_plan_LC5 = TextInput(text=f'{current_date.hour}:{current_date.minute}', size_hint_x=0.2)
        self.label_date_start_plan_LC5 = Label(text='Data de Início (DD/MM/AAAA): ')
        self.ti_date_start_plan_LC5 = TextInput(text=f'{current_date.day}/{current_date.month}/{current_date.year}', size_hint_x=0.3)
        self.button_remove_shift_LC5 = Button(text='-', background_color=color_dark_red, size_hint_x=0.1)
        self.button_add_shift_LC5 = Button(text='+', background_color=color_light_green, size_hint_x=0.1)

        self.layout_shifts_LF = BoxLayout(orientation='vertical')
        self.header_layout_shifts_LF = BoxLayout(orientation='horizontal')
        self.header_layout_id_LF = Label(text='ID')
        self.header_layout_time_start_LF = Label(text='Hora Início (HH:MM)')
        self.header_layout_time_finish_LF = Label(text='Hora Fim (HH:MM)')
        self.header_layout_time_break_LF = Label(text='Hora Pausa (HH:MM)')
        self.header_layout_dur_break_LF = Label(text='Dur. Pausa (min)')

        self.layout_shifts_LC5 = BoxLayout(orientation='vertical')
        self.header_layout_shifts_LC5 = BoxLayout(orientation='horizontal')
        self.header_layout_id_LC5 = Label(text='ID')
        self.header_layout_time_start_LC5 = Label(text='Hora Início (HH:MM)')
        self.header_layout_time_finish_LC5 = Label(text='Hora Fim (HH:MM)')
        self.header_layout_time_break_LC5 = Label(text='Hora Pausa (HH:MM)')
        self.header_layout_dur_break_LC5 = Label(text='Dur. Pausa (min)')

        self.row_buttons = BoxLayout(orientation='horizontal', size_hint_y=0.5)
        self.button_return = Button(text='<<<', background_color=color_light_blue)
        self.button_submit = Button(text='Ok', background_color=color_light_blue)

        self.button_add_shift_LF.bind(on_press=lambda x: self.add_shift('LF3015'))
        self.button_add_shift_LC5.bind(on_press=lambda x: self.add_shift('LC5'))
        self.button_remove_shift_LF.bind(on_press=lambda x: self.remove_shift('LF3015'))
        self.button_remove_shift_LC5.bind(on_press=lambda x: self.remove_shift('LC5'))
        self.button_return.bind(on_press=self.close_popup_window)
        self.button_submit.bind(on_press=self.validate_data)

        self.row_main_info_LF.add_widget(self.label_num_shifts_LF)
        self.row_main_info_LF.add_widget(self.button_remove_shift_LF)
        self.row_main_info_LF.add_widget(self.button_add_shift_LF)
        self.row_main_info_LF.add_widget(self.label_time_start_plan_LF)
        self.row_main_info_LF.add_widget(self.ti_time_start_plan_LF)
        self.row_main_info_LF.add_widget(self.label_date_start_plan_LF)
        self.row_main_info_LF.add_widget(self.ti_date_start_plan_LF)
        self.row_main_info_LF.add_widget(Label(size_hint_x=1))

        self.row_main_info_LC5.add_widget(self.label_num_shifts_LC5)
        self.row_main_info_LC5.add_widget(self.button_remove_shift_LC5)
        self.row_main_info_LC5.add_widget(self.button_add_shift_LC5)
        self.row_main_info_LC5.add_widget(self.label_time_start_plan_LC5)
        self.row_main_info_LC5.add_widget(self.ti_time_start_plan_LC5)
        self.row_main_info_LC5.add_widget(self.label_date_start_plan_LC5)
        self.row_main_info_LC5.add_widget(self.ti_date_start_plan_LC5)
        self.row_main_info_LC5.add_widget(Label(size_hint_x=1))

        self.header_layout_shifts_LF.add_widget(self.header_layout_id_LF)
        self.header_layout_shifts_LF.add_widget(self.header_layout_time_start_LF)
        self.header_layout_shifts_LF.add_widget(self.header_layout_time_finish_LF)
        self.header_layout_shifts_LF.add_widget(self.header_layout_time_break_LF)
        self.header_layout_shifts_LF.add_widget(self.header_layout_dur_break_LF)
        self.layout_shifts_LF.add_widget(self.header_layout_shifts_LF)

        self.header_layout_shifts_LC5.add_widget(self.header_layout_id_LC5)
        self.header_layout_shifts_LC5.add_widget(self.header_layout_time_start_LC5)
        self.header_layout_shifts_LC5.add_widget(self.header_layout_time_finish_LC5)
        self.header_layout_shifts_LC5.add_widget(self.header_layout_time_break_LC5)
        self.header_layout_shifts_LC5.add_widget(self.header_layout_dur_break_LC5)
        self.layout_shifts_LC5.add_widget(self.header_layout_shifts_LC5)

        self.row_buttons.add_widget(self.button_return)
        self.row_buttons.add_widget(self.button_submit)

        self.add_widget(self.label_machine_LF)
        self.add_widget(self.row_main_info_LF)
        self.add_widget(self.layout_shifts_LF)
        self.add_widget(Label(size_hint_y=0.1))
        self.add_widget(self.label_machine_LC5)
        self.add_widget(self.row_main_info_LC5)
        self.add_widget(self.layout_shifts_LC5)
        self.add_widget(Label(size_hint_y=0.1))
        self.add_widget(self.row_buttons)

        self.update_shifts_layout('LF3015')
        self.update_shifts_layout('LC5')

    def close_popup_window(self, *args):
        production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sidebar.pop_plan_tasks.dismiss()

    def update_shifts_layout(self, machine):
        if machine == 'LF3015':
            for shift in self.added_shifts_LF:
                self.layout_shifts_LF.remove_widget(shift)

            counter = 1
            for shift in self.shifts_LF:
                new_shift = Shift(counter, shift[0], shift[1], shift[2], shift[3], shift[4])
                self.layout_shifts_LF.add_widget(new_shift)
                self.added_shifts_LF.append(new_shift)
                counter += 1
        else:
            for shift in self.added_shifts_LC5:
                self.layout_shifts_LC5.remove_widget(shift)

            counter = 1
            for shift in self.shifts_LC5:
                new_shift = Shift(counter, shift[0], shift[1], shift[2], shift[3], shift[4])
                self.layout_shifts_LC5.add_widget(new_shift)
                self.added_shifts_LC5.append(new_shift)
                counter += 1

    def add_shift(self, machine):
        if machine == 'LF3015':
            if len(self.added_shifts_LF) == 3:
                return None
            else:
                new_shift = Shift()
                self.added_shifts_LF.append(new_shift)
                new_shift.shift_id = str(self.added_shifts_LF.index(new_shift) + 1)
                new_shift.label_id.text = f'T{new_shift.shift_id}'
                self.layout_shifts_LF.add_widget(new_shift)

        else:
            if len(self.added_shifts_LC5) == 3:
                return None
            else:
                new_shift = Shift()
                self.added_shifts_LC5.append(new_shift)
                new_shift.shift_id = str(self.added_shifts_LC5.index(new_shift) + 1)
                new_shift.label_id.text = f'T{new_shift.shift_id}'
                self.layout_shifts_LC5.add_widget(new_shift)

        self.update_num_shifts_label(machine)

    def remove_shift(self, machine):

        if machine == 'LF3015':
            if len(self.added_shifts_LF) == 1:
                return None
            else:
                self.layout_shifts_LF.remove_widget(self.added_shifts_LF[-1])
                self.added_shifts_LF.remove(self.added_shifts_LF[-1])

        else:
            if len(self.added_shifts_LC5) == 1:
                return None
            else:
                self.layout_shifts_LC5.remove_widget(self.added_shifts_LC5[-1])
                self.added_shifts_LC5.remove(self.added_shifts_LC5[-1])

        self.update_num_shifts_label(machine)

    def update_num_shifts_label(self, machine):

        if machine == 'LF3015':
            self.label_num_shifts_LF.text = f'Nº Turnos:   {str(len(self.added_shifts_LF))}'
        else:
            self.label_num_shifts_LC5.text = f'Nº Turnos:   {str(len(self.added_shifts_LC5))}'

    def validate_data(self, *args):
        total_added_shifts = self.added_shifts_LF + self.added_shifts_LC5
        valid_time = True
        valid_date_LF = check_date_is_valid(self.ti_date_start_plan_LF.text)
        valid_date_LC5 = check_date_is_valid(self.ti_date_start_plan_LC5.text)

        for added_shift in total_added_shifts:
            is_valid_1 = self.check_valid_time(added_shift.ti_time_start.text)
            is_valid_2 = self.check_valid_time(added_shift.ti_time_finish.text)
            is_valid_3 = self.check_valid_time(added_shift.ti_time_break.text)
            is_valid_4 = self.check_valid_time(self.ti_time_start_plan_LF.text)
            is_valid_5 = self.check_valid_time(self.ti_time_start_plan_LC5.text)

            if is_valid_1 == False or is_valid_2 == False or is_valid_3 == False or is_valid_4 == False or is_valid_5 == False:
                valid_time = False
                break

            valid_minutes = self.check_valid_minutes(added_shift.ti_break_duration.text)

            if valid_minutes == False:
                break

        if valid_time and valid_minutes and valid_date_LF and valid_date_LC5:
            try:
                production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sidebar.plan_selected_tasks(self.ti_time_start_plan_LF.text, self.ti_time_start_plan_LC5.text, self.ti_date_start_plan_LF.text, self.ti_date_start_plan_LC5.text)
                production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sidebar.pop_plan_tasks.dismiss()
            except PermissionError:
                popup_warning_file_open = PopupWarningMessage(message='Alguém tem o Excel aberto, pede para fecharem', type='warning')
                popup_warning_file_open.open()


        else:
            self.popup_warning_delete = PopupWarningMessage(message='Alguma coisa está mal preenchida', type='warning')
            self.popup_warning_delete.open()

    def check_valid_time(self, time_string):
        num_list = time_string.split(':')

        if len(num_list) != 2:
            return False

        for num in num_list:
            try:
                int(num)
            except ValueError:
                return False

        if int(num_list[0]) < 0 or int(num_list[0]) > 24 or int(num_list[1]) < 0 or int(num_list[1]) > 60:
            return False

        return True

    def check_valid_minutes(self, minutes):
        try:
            int(minutes)
        except ValueError:
            return False

        if int(minutes) < 0 or int(minutes) > 60:
            return False

        return True


class Shift(BoxLayout):
    def __init__(self, shift_id=0, machine='', time_start='', time_finish='', time_break='', break_duration='', **kwargs):
        super().__init__(**kwargs)

        self.orientation = 'horizontal'

        self.shift_id = shift_id
        self.machine = machine
        self.time_start = time_start
        self.time_finish = time_finish
        self.time_break = time_break
        self.break_duration = break_duration

        self.label_id = Label(text=f'T{self.shift_id}')
        self.ti_time_start = TextInput(text=self.time_start)
        self.ti_time_finish = TextInput(text=self.time_finish)
        self.ti_time_break = TextInput(text=self.time_break)
        self.ti_break_duration = TextInput(text=str(self.break_duration))

        self.add_widget(self.label_id)
        self.add_widget(self.ti_time_start)
        self.add_widget(self.ti_time_finish)
        self.add_widget(self.ti_time_break)
        self.add_widget(self.ti_break_duration)


class Shiftspage(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.shifts_LF = load_shifts('LF3015')
        self.shifts_LC5 = load_shifts('LC5')
        self.added_shifts_LF = []
        self.added_shifts_LC5 = []

        self.orientation = 'vertical'

        self.label_machine_LF = Label(text='LF-3015', font_size='20')
        self.row_main_info_LF = BoxLayout(orientation='horizontal', size_hint_y=0.3)
        self.label_num_shifts_LF = Label(text=f'Nº Turnos:   {str(len(self.shifts_LF))}', size_hint_x=0.2)
        self.button_remove_shift_LF = Button(text='-', background_color=color_dark_red, size_hint_x=0.1)
        self.button_add_shift_LF = Button(text='+', background_color=color_light_green, size_hint_x=0.1)

        self.layout_shifts_LF = BoxLayout(orientation='vertical')

        self.label_machine_LC5 = Label(text='LC5', font_size='20')
        self.row_main_info_LC5 = BoxLayout(orientation='horizontal', size_hint_y=0.3)
        self.label_num_shifts_LC5 = Label(text=f'Nº Turnos:   {str(len(self.shifts_LC5))}', size_hint_x=0.2)
        self.button_remove_shift_LC5 = Button(text='-', background_color=color_dark_red, size_hint_x=0.1)
        self.button_add_shift_LC5 = Button(text='+', background_color=color_light_green, size_hint_x=0.1)

        self.layout_shifts_LF = BoxLayout(orientation='vertical')
        self.header_layout_shifts_LF = BoxLayout(orientation='horizontal')
        self.header_layout_id_LF = Label(text='ID')
        self.header_layout_time_start_LF = Label(text='Hora Início (HH:MM)')
        self.header_layout_time_finish_LF = Label(text='Hora Fim (HH:MM)')
        self.header_layout_time_break_LF = Label(text='Hora Pausa (HH:MM)')
        self.header_layout_dur_break_LF = Label(text='Dur. Pausa (min)')

        self.layout_shifts_LC5 = BoxLayout(orientation='vertical')
        self.header_layout_shifts_LC5 = BoxLayout(orientation='horizontal')
        self.header_layout_id_LC5 = Label(text='ID')
        self.header_layout_time_start_LC5 = Label(text='Hora Início (HH:MM)')
        self.header_layout_time_finish_LC5 = Label(text='Hora Fim (HH:MM)')
        self.header_layout_time_break_LC5 = Label(text='Hora Pausa (HH:MM)')
        self.header_layout_dur_break_LC5 = Label(text='Dur. Pausa (min)')

        self.row_buttons = BoxLayout(orientation='horizontal', size_hint_y=0.5)
        self.button_return = Button(text='<<<', background_color=color_light_blue)
        self.button_submit = Button(text='Ok', background_color=color_light_blue)

        self.button_add_shift_LF.bind(on_press=lambda x: self.add_shift('LF3015'))
        self.button_add_shift_LC5.bind(on_press=lambda x: self.add_shift('LC5'))
        self.button_remove_shift_LF.bind(on_press=lambda x: self.remove_shift('LF3015'))
        self.button_remove_shift_LC5.bind(on_press=lambda x: self.remove_shift('LC5'))
        self.button_return.bind(on_press=self.close_popup_window)
        self.button_submit.bind(on_press=self.validate_data)

        self.row_main_info_LF.add_widget(self.label_num_shifts_LF)
        self.row_main_info_LF.add_widget(self.button_remove_shift_LF)
        self.row_main_info_LF.add_widget(self.button_add_shift_LF)
        self.row_main_info_LF.add_widget(Label(size_hint_x=1))

        self.row_main_info_LC5.add_widget(self.label_num_shifts_LC5)
        self.row_main_info_LC5.add_widget(self.button_remove_shift_LC5)
        self.row_main_info_LC5.add_widget(self.button_add_shift_LC5)
        self.row_main_info_LC5.add_widget(Label(size_hint_x=1))

        self.header_layout_shifts_LF.add_widget(self.header_layout_id_LF)
        self.header_layout_shifts_LF.add_widget(self.header_layout_time_start_LF)
        self.header_layout_shifts_LF.add_widget(self.header_layout_time_finish_LF)
        self.header_layout_shifts_LF.add_widget(self.header_layout_time_break_LF)
        self.header_layout_shifts_LF.add_widget(self.header_layout_dur_break_LF)
        self.layout_shifts_LF.add_widget(self.header_layout_shifts_LF)

        self.header_layout_shifts_LC5.add_widget(self.header_layout_id_LC5)
        self.header_layout_shifts_LC5.add_widget(self.header_layout_time_start_LC5)
        self.header_layout_shifts_LC5.add_widget(self.header_layout_time_finish_LC5)
        self.header_layout_shifts_LC5.add_widget(self.header_layout_time_break_LC5)
        self.header_layout_shifts_LC5.add_widget(self.header_layout_dur_break_LC5)
        self.layout_shifts_LC5.add_widget(self.header_layout_shifts_LC5)

        self.row_buttons.add_widget(self.button_return)
        self.row_buttons.add_widget(self.button_submit)

        self.add_widget(self.label_machine_LF)
        self.add_widget(self.row_main_info_LF)
        self.add_widget(self.layout_shifts_LF)
        self.add_widget(Label(size_hint_y=0.1))
        self.add_widget(self.label_machine_LC5)
        self.add_widget(self.row_main_info_LC5)
        self.add_widget(self.layout_shifts_LC5)
        self.add_widget(Label(size_hint_y=0.1))
        self.add_widget(self.row_buttons)

        self.update_shifts_layout('LF3015')
        self.update_shifts_layout('LC5')

    def close_popup_window(self, *args):
        production_planning.homepage.layout_sidebar.pop_shifts.dismiss()


    def update_shifts_layout(self, machine):
        if machine == 'LF3015':
            for shift in self.added_shifts_LF:
                self.layout_shifts_LF.remove_widget(shift)

            counter = 1
            for shift in self.shifts_LF:
                new_shift = Shift(counter, shift[0], shift[1], shift[2], shift[3], shift[4])
                self.layout_shifts_LF.add_widget(new_shift)
                self.added_shifts_LF.append(new_shift)
                counter += 1
        else:
            for shift in self.added_shifts_LC5:
                self.layout_shifts_LC5.remove_widget(shift)

            counter = 1
            for shift in self.shifts_LC5:
                new_shift = Shift(counter, shift[0], shift[1], shift[2], shift[3], shift[4])
                self.layout_shifts_LC5.add_widget(new_shift)
                self.added_shifts_LC5.append(new_shift)
                counter += 1

    def add_shift(self, machine):
        if machine == 'LF3015':
            if len(self.added_shifts_LF) == 3:
                return None
            else:
                new_shift = Shift()
                self.added_shifts_LF.append(new_shift)
                new_shift.shift_id = str(self.added_shifts_LF.index(new_shift) + 1)
                new_shift.label_id.text = f'T{new_shift.shift_id}'
                self.layout_shifts_LF.add_widget(new_shift)

        else:
            if len(self.added_shifts_LC5) == 3:
                return None
            else:
                new_shift = Shift()
                self.added_shifts_LC5.append(new_shift)
                new_shift.shift_id = str(self.added_shifts_LC5.index(new_shift) + 1)
                new_shift.label_id.text = f'T{new_shift.shift_id}'
                self.layout_shifts_LC5.add_widget(new_shift)

        self.update_num_shifts_label(machine)

    def remove_shift(self, machine):

        if machine == 'LF3015':
            if len(self.added_shifts_LF) == 1:
                return None
            else:
                self.layout_shifts_LF.remove_widget(self.added_shifts_LF[-1])
                self.added_shifts_LF.remove(self.added_shifts_LF[-1])

        else:
            if len(self.added_shifts_LC5) == 1:
                return None
            else:
                self.layout_shifts_LC5.remove_widget(self.added_shifts_LC5[-1])
                self.added_shifts_LC5.remove(self.added_shifts_LC5[-1])

        self.update_num_shifts_label(machine)

    def update_num_shifts_label(self, machine):


        if machine == 'LF3015':
            self.label_num_shifts_LF.text = f'Nº Turnos:   {str(len(self.added_shifts_LF))}'
        else:
            self.label_num_shifts_LC5.text = f'Nº Turnos:   {str(len(self.added_shifts_LC5))}'

    def validate_data(self, *args):
        total_added_shifts = self.added_shifts_LF + self.added_shifts_LC5
        valid_time = True

        for added_shift in total_added_shifts:
            is_valid_1 = self.check_valid_time(added_shift.ti_time_start.text)
            is_valid_2 = self.check_valid_time(added_shift.ti_time_finish.text)
            is_valid_3 = self.check_valid_time(added_shift.ti_time_break.text)

            if is_valid_1 == False or is_valid_2 == False or is_valid_3 == False:
                valid_time = False
                break

            valid_minutes = self.check_valid_minutes(added_shift.ti_break_duration.text)

            if valid_minutes == False:
                break

        if valid_time and valid_minutes:
            remove_shifts()

            for shift in self.added_shifts_LF:
                insert_new_shift('LF3015', shift.ti_time_start.text, shift.ti_time_finish.text, shift.ti_time_break.text, shift.ti_break_duration.text)

            for shift in self.added_shifts_LC5:
                insert_new_shift('LC5', shift.ti_time_start.text, shift.ti_time_finish.text, shift.ti_time_break.text, shift.ti_break_duration.text)

            production_planning.homepage.layout_sidebar.pop_shifts.dismiss()



        else:
            self.popup_warning_delete = PopupWarningMessage(message='Alguma coisa está mal preenchida', type='warning')
            self.popup_warning_delete.open()

    def check_valid_time(self, time_string):
        num_list = time_string.split(':')

        if len(num_list) != 2:
            return False

        for num in num_list:
            try:
                int(num)
            except ValueError:
                return False

        if int(num_list[0]) < 0 or int(num_list[0]) > 24 or int(num_list[1]) < 0 or int(num_list[1]) > 60:
            return False

        return True

    def check_valid_minutes(self, minutes):
        try:
            int(minutes)
        except ValueError:
            return False

        if int(minutes) < 0 or int(minutes) > 60:
            return False

        return True


class Production():
    pass


class Clientspage_Sideframe(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)


class Clientspage_Sidebar(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.orientation = 'vertical'
        self.size_hint_x = 0.08

        self.button_return = Button(text='<<<', font_size='16', halign='center', background_color=color_light_blue)
        self.button_new_client = Button(text='Novo Cliente', font_size='16', halign='center', background_color=color_light_blue)
        self.button_delete_client = Button(text='Apagar Clientes\nSelecionados', font_size='16', halign='center', background_color=color_light_blue)
        self.button_consume_history = Button(text='Histórico de\nConsumos', font_size='16', halign='center', background_color=color_light_blue)

        self.button_return.bind(on_press=self.close_popup_window)

        self.add_widget(self.button_return)
        self.add_widget(self.button_new_client)
        self.add_widget(self.button_delete_client)
        self.add_widget(self.button_consume_history)

    def close_popup_window(self, *args):
        production_planning.homepage.layout_sidebar.pop_clients.dismiss()


class Clientspage(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.orientation = 'horizontal'

        self.layout_sidebar = Clientspage_Sidebar()
        self.layout_sideframe = Clientspage_Sideframe()

        self.add_widget(self.layout_sidebar)
        self.add_widget(self.layout_sideframe)


class Order(BoxLayout):
    def __init__(self, order_num, order_num_client, client, is_selected=False, **kwargs):
        super().__init__(**kwargs)
        #TODO Change the button to red when any order part from the order is one day away from end date

        self.orientation = 'vertical'
        self.size_hint_y = None
        self.height = 45

        self.order_num = order_num
        self.order_num_client = order_num_client
        self.client = client
        self.is_selected = is_selected
        self.num_parts_additional_operations = 0
        self.num_total_order_parts = 0
        self.num_completed_order_parts = 0
        self.percent_completed = 0
        self.earliest_due_date = ''
        self.button_idle_color = color_light_blue

        self.calculate_num_order_parts()
        self.set_button_idle_color()

        self.text_display_button = f'{self.client}   -   Enc. Int: {self.order_num}  /  Enc. Cliente:  {self.order_num_client}\n[{self.num_parts_additional_operations} Peças c/ operações adicionais] - [{self.num_completed_order_parts}/{self.num_total_order_parts} Peças Prod.] - [{self.percent_completed} Completa]'

        self.button_order = RightClickableButton(text=self.text_display_button, font_size='17', halign='center', background_color=self.button_idle_color)
        self.button_order.bind(on_press=self.select_order)
        self.add_widget(self.button_order)

        self.set_earliest_due_date()
        self.check_due_date()

    def select_order(self, *args):
        if self.button_order.mouse_button == 'left':
            if not self.is_selected:
                self.is_selected = True
            self.change_button_color()

            production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sideframe.update_order_parts_display(self.order_num)
        elif self.button_order.mouse_button == 'right':
            # print(production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sideframe.order_parts)
            self.layout_popup_change_general_order_info = BoxLayout(orientation='vertical')
            self.layout_popup_line_1 = BoxLayout(orientation='horizontal')
            self.layout_popup_line_2 = BoxLayout(orientation='horizontal')
            self.layout_popup_line_3 = BoxLayout(orientation='horizontal')
            self.label_popup_order_num = Label(text='Enc. Interna: ')
            self.ti_popup_order_num = TextInput(text=self.order_num)
            self.label_popup_order_num_client = Label(text='Enc. Cliente: ')
            self.ti_popup_order_num_client = TextInput(text=self.order_num_client)
            self.label_popup_client = Label(text='Cliente: ')
            self.ti_popup_client = TextInput(text=self.client)
            self.popup_button_ok = Button(text='OK', font_size='14', halign='right', background_color=color_light_blue)

            self.popup_button_ok.bind(on_press=self.change_general_order_info)

            self.layout_popup_line_1.add_widget(self.label_popup_order_num)
            self.layout_popup_line_1.add_widget(self.ti_popup_order_num)
            self.layout_popup_line_2.add_widget(self.label_popup_order_num_client)
            self.layout_popup_line_2.add_widget(self.ti_popup_order_num_client)
            self.layout_popup_line_3.add_widget(self.label_popup_client)
            self.layout_popup_line_3.add_widget(self.ti_popup_client)
            self.layout_popup_change_general_order_info.add_widget(self.layout_popup_line_1)
            self.layout_popup_change_general_order_info.add_widget(self.layout_popup_line_2)
            self.layout_popup_change_general_order_info.add_widget(self.layout_popup_line_3)
            self.layout_popup_change_general_order_info.add_widget(self.popup_button_ok)

            self.popup_change_general_order_info = Popup(title='Alterar dados gerais da encomenda', content=self.layout_popup_change_general_order_info, size_hint=(0.4, 0.25))

            self.popup_change_general_order_info.open()

    def change_general_order_info(self, *args):
        previous_order_num = self.order_num
        previous_order_num_client = self.order_num_client
        previous_client = self.client
        new_order_num = self.ti_popup_order_num.text
        new_order_num_client = self.ti_popup_order_num_client.text
        new_client = self.ti_popup_client.text

        for order_part in production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sideframe.order_parts:
            if order_part[5] == previous_order_num and order_part[6] == previous_order_num_client and order_part[7] == previous_client:
                order_part[5] = new_order_num
                order_part[6] = new_order_num_client
                order_part[7] = new_client
                change_value_order_part_database(new_order_num, 'order_num', order_part[-1])
                change_value_order_part_database(new_order_num_client, 'order_num_client', order_part[-1])
                change_value_order_part_database(new_client, 'client', order_part[-1])

        self.order_num = new_order_num
        self.order_num_client = new_order_num_client
        self.client = new_client
        self.text_display_button = f'{self.client}   -   Enc. Int: {self.order_num}  /  Enc. Cliente:  {self.order_num_client}\n[{self.num_parts_additional_operations} Peças c/ operações adicionais] - [{self.num_completed_order_parts}/{self.num_total_order_parts} Peças Prod.] - [{self.percent_completed} Completa]'
        self.button_order.text = self.text_display_button

        self.popup_change_general_order_info.dismiss()

    def change_button_color(self, *args):
        if self.is_selected:
            self.button_order.background_color = color_dark_green
        else:
            self.button_order.background_color = self.button_idle_color

        for order in production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sideframe.added_orders:
            if order == self:
                pass
            else:
                order.button_order.background_color = order.button_idle_color
                order.is_selected = False

    def calculate_num_order_parts(self, *args):
        total_parts = 0
        completed_parts = 0
        parts_additional_operations = 0
        for order_part in load_order_parts_from_database():
            if order_part[5] == self.order_num and order_part[6] == self.order_num_client and order_part[7] == self.client:
                total_parts += 1
                if order_part[4] >= order_part[3]:
                    completed_parts += 1
                if order_part[10]:
                    parts_additional_operations += 1

        percent_completed = int(round((completed_parts/total_parts)*100, 0))

        self.num_total_order_parts = str(total_parts)
        self.num_completed_order_parts = str(completed_parts)
        self.num_parts_additional_operations = str(parts_additional_operations)
        self.percent_completed = f'{str(percent_completed)}%'
        self.text_display_button = f'{self.client}   -   Enc. Int: {self.order_num}  /  Enc. Cliente:  {self.order_num_client}\n[{self.num_parts_additional_operations} Peças c/ operações adicionais] - [{self.num_completed_order_parts}/{self.num_total_order_parts} Peças Prod.] - [{self.percent_completed} Completa]'

    def update_button_info_text(self, *args):
        self.button_order.text = self.text_display_button
        self.check_due_date()

    def set_button_idle_color(self, *args):
        if self.percent_completed == '100%':
            self.button_idle_color = color_light_green
        else:
            self.button_idle_color = color_light_blue

    def set_earliest_due_date(self, *args):
        order_parts_due_dates = []

        for order_part in load_order_parts_from_database():
            if order_part[5] == self.order_num and order_part[6] == self.order_num_client and order_part[7] == self.client:
                try:
                    due_date_info = [int(date_elem) for date_elem in order_part[9].split('/')]
                    order_parts_due_dates.append(datetime.datetime(day=due_date_info[0], month=due_date_info[1], year=due_date_info[2], hour=17, minute=45))
                except ValueError:
                    pass

        try:
            self.earliest_due_date = min(order_parts_due_dates)
        except ValueError:
            pass

        # print(self.earliest_due_date)

    def check_due_date(self, *args):
        if self.earliest_due_date != '':
            curr = datetime.datetime.now()
            current_date = datetime.datetime(year=curr.year, month=curr.month, day=curr.day, hour=17, minute=45)

            time_interval = self.earliest_due_date - current_date

            if time_interval.days < 0:
                self.button_order.text += f' - Encomenda atrasada {abs(time_interval.days)} dias'
            elif time_interval.days > 0:
                self.button_order.text += f' - Data de entrega em {time_interval.days} dias'
            else:
                self.button_order.text += f' - Data de Entrega Hoje'


class OrderPart(BoxLayout):
    def __init__(self, ref, name, material_ref, quantity, produced_quantity, order_num, order_num_client, client,
                 date_modified, due_date, additional_operations, rowid, is_selected=False, **kwargs):
        super().__init__(**kwargs)

        self.orientation = 'horizontal'
        self.size_hint_y = None
        self.height = 35

        self.ref = ref
        self.name = name
        self.material_ref = material_ref
        self.material_name = get_material_name(self.material_ref)
        self.quantity = str(quantity)
        self.produced_quantity = str(produced_quantity)
        self.order_num = order_num
        self.order_num_client = order_num_client
        self.client = client
        self.date_modified = date_modified
        self.due_date = due_date
        self.additional_operations = bool(additional_operations)
        self.rowid = str(rowid)
        self.is_selected = is_selected

        self.checkbox_selected = CheckBox(size_hint_x=0.03)
        self.button_ref = RightClickableButton(text=self.ref, font_size='16', background_color=color_dark_blue, size_hint_x=0.1)
        self.button_name = RightClickableButton(text=self.name[0:78], font_size='14', background_color=color_dark_blue, size_hint_x=0.32)
        self.button_material_name = RightClickableButton(text=self.material_name, font_size='16', background_color=color_dark_blue, size_hint_x=0.17)
        self.button_quantity = RightClickableButton(text=self.quantity, font_size='16', background_color=color_dark_blue, size_hint_x=0.05)
        self.button_produced_quantity = RightClickableButton(text=self.produced_quantity, font_size='16', background_color=color_dark_blue, size_hint_x=0.05)
        self.button_date_modified = RightClickableButton(text=self.date_modified, font_size='16', background_color=color_dark_blue, size_hint_x=0.1)
        self.button_due_date = RightClickableButton(text=self.due_date, font_size='16', background_color=color_dark_blue, size_hint_x=0.1)
        self.checkbox_additional_operations = CheckBox(size_hint_x=0.03)

        self.checkbox_selected.bind(on_press=self.set_selected)
        self.button_ref.bind(on_press=lambda button: self.choose_on_press_function('ref', self.rowid, button))
        self.button_name.bind(on_press=lambda button: self.choose_on_press_function('name', self.rowid, button))
        self.button_material_name.bind(on_press=lambda button: self.choose_on_press_function('material_name', self.rowid, button))
        self.button_quantity.bind(on_press=lambda button: self.choose_on_press_function('quantity', self.rowid, button))
        self.button_produced_quantity.bind(on_press=lambda button: self.choose_on_press_function('produced_quantity', self.rowid, button))
        self.button_due_date.bind(on_press=lambda button: self.choose_on_press_function('due_date', self.rowid, button))
        self.checkbox_additional_operations.bind(on_press=self.set_additional_operations)

        self.add_widget(self.checkbox_selected)
        self.add_widget(self.button_ref)
        self.add_widget(self.button_name)
        self.add_widget(self.button_material_name)
        self.add_widget(self.button_quantity)
        self.add_widget(self.button_produced_quantity)
        self.add_widget(self.button_date_modified)
        self.add_widget(self.button_due_date)
        self.add_widget(self.checkbox_additional_operations)

        self.set_initial_state_additional_operations_checkbox()
        self.set_color_buttons()
        self.set_ref_button_text()

    def set_selected(self, *args):
        self.is_selected = not self.is_selected

    def set_initial_state_additional_operations_checkbox(self):
        if self.additional_operations:
            self.checkbox_additional_operations.state = 'down'

    def set_ref_button_text(self):
        if self.ref == '1':
            self.button_ref.text = 'S/ Ref'

    def set_additional_operations(self, *args):
        self.additional_operations = not self.additional_operations
        change_value_order_part_database(self.additional_operations, 'additional_operations', self.rowid)
        self.update_button_order_info()

    def choose_on_press_function(self, field, rowid, button, *args):
        if button.mouse_button == 'left':
            for order_part in production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sideframe.added_order_parts:
                if button == order_part.button_ref or button == order_part.button_name or button == order_part.button_material_name or button == order_part.button_quantity or button == order_part.button_produced_quantity or button == order_part.button_date_modified or button == order_part.button_due_date:
                    order_part.button_ref.background_color = color_light_black
                    order_part.button_name.background_color = color_light_black
                    order_part.button_material_name.background_color = color_light_black
                    order_part.button_quantity.background_color = color_light_black
                    order_part.button_produced_quantity.background_color = color_light_black
                    order_part.button_date_modified.background_color = color_light_black
                    order_part.button_due_date.background_color = color_light_black
                else:
                    order_part.set_color_buttons()
                    # order_part.button_ref.background_color = color_dark_blue
                    # order_part.button_name.background_color = color_dark_blue
                    # order_part.button_material_name.background_color = color_dark_blue
                    # order_part.button_quantity.background_color = color_dark_blue
                    # order_part.button_produced_quantity.background_color = color_dark_blue
                    # order_part.button_date_modified.background_color = color_dark_blue
                    # order_part.button_due_date.background_color = color_dark_blue

        elif button.mouse_button == 'right':
            self.display_popup_change_value(field, rowid)

    def display_popup_change_value(self, field, rowid):

        self.layout_new_value = BoxLayout(orientation='vertical')
        self.layout_horizontal_line = BoxLayout(orientation='horizontal', padding=(10, 10))
        self.label_new_value = Label(text='Novo valor: ', font_size='16', halign='right', size_hint_x=0.3)
        if field == 'material_name':
            self.drp_material = DropDown()
            self.button_material = Button(text=self.material_ref, font_size='16', halign='right', size_hint_y=0.75)
            self.button_material.bind(on_release=self.drp_material.open)
            self.drp_material.bind(on_select=lambda instance, new_val: setattr(self.button_material, 'text', new_val))
            self.button_submit = Button(text='Submeter', font_size='18', background_color=color_light_green)

            for material in load_materials_from_database():
                material_type = material[1]
                material_spec = str(material[2])
                material_thickness = str(material[3])
                material_client = material[11]

                material_name = material_type + ' ' + material_spec + ', ' + material_thickness + 'mm ' + material_client

                button_material_name = Button(text=material_name, font_size='16', size_hint_y=None, height=30, background_color=color_dark_blue)
                button_material_name.bind(on_release=lambda material_object: self.drp_material.select(material_object.text))

                self.drp_material.add_widget(button_material_name)

            self.layout_horizontal_line.add_widget(self.label_new_value)
            self.layout_horizontal_line.add_widget(self.button_material)
            self.layout_new_value.add_widget(self.layout_horizontal_line)
            self.layout_new_value.add_widget(self.button_submit)

            self.button_submit.bind(on_press=lambda x: self.change_value(self.button_material.text, field, rowid))

            self.pop_change_value = Popup(title='Novo valor: ', content=self.layout_new_value, size_hint=(0.50, 0.20))
            self.pop_change_value.open()

        else:
            if field == 'ref':
                self.ti_new_value = TextInput(text=self.button_ref.text)
            elif field == 'name':
                self.ti_new_value = TextInput(text=self.name)
            else:
                self.ti_new_value = TextInput()

            self.button_submit = Button(text='Submeter', font_size='18', background_color=color_light_green)

            self.layout_horizontal_line.add_widget(self.label_new_value)
            self.layout_horizontal_line.add_widget(self.ti_new_value)
            self.layout_new_value.add_widget(self.layout_horizontal_line)
            self.layout_new_value.add_widget(self.button_submit)

            self.button_submit.bind(on_press=lambda x: self.check_validity(self.ti_new_value.text, field, rowid))

            self.pop_change_value = Popup(title='Novo valor: ', content=self.layout_new_value, size_hint=(0.50, 0.20))
            self.pop_change_value.bind(on_open=self.set_focus_text_input)
            self.pop_change_value.open()

    def check_validity(self, value, field, rowid):
        valid = False

        if field == 'due_date':
            valid = check_date_is_valid(value)
        else:
            valid = True

        if valid:
            self.change_value(value, field, rowid)
        else:
            if field == 'due_date':
                warning_popup = PopupWarningMessage('Inserir data no formato: dia/mês/ano')
                warning_popup.open()

    def change_value(self, value, field, rowid, *args):
        if field == 'ref':
            change_value_order_part_database(value, field, rowid)
            self.button_ref.text = str(value)
        elif field == 'name':
            change_value_order_part_database(value, field, rowid)
            self.button_name.text = str(value)
        elif field == 'material_name':
            material_type = value.split(' ')[0]
            material_spec = value.split(' ')[1].split(',')[0]
            material_thickness = float(value.split(' ')[2].split('mm')[0])
            material_client = value.split(' ')[3]
            material_ref = get_material_ref(material_type, material_spec, material_thickness, material_client)
            change_value_order_part_database(material_ref, 'material_ref', rowid)
            material_name = get_material_name(material_ref)
            self.button_material_name.text = material_name
        elif field == 'quantity':
            value = int(value)
            change_value_order_part_database(value, field, rowid)
            self.button_quantity.text = str(value)
            self.set_color_buttons()
        elif field == 'produced_quantity':
            value = int(value)
            change_value_order_part_database(value, field, rowid)
            self.button_produced_quantity.text = str(value)
            self.set_color_buttons()
        elif field == 'due_date':
            change_value_order_part_database(value, field, rowid)
            self.button_due_date.text = str(value)

        date_modified = f'{datetime.datetime.now().day}/{datetime.datetime.now().month}/{datetime.datetime.now().year}'
        self.button_date_modified.text = date_modified
        change_value_order_part_database(date_modified, 'date_modified', rowid)
        self.pop_change_value.dismiss()
        self.update_button_order_info()

    def update_button_order_info(self, *args):
        for order in production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sideframe.added_orders:
            if order.is_selected:
                order.button_order.background_color = color_dark_green
                order.calculate_num_order_parts()
                order.update_button_info_text()
                order.set_button_idle_color()

    def set_color_buttons(self):
        if int(self.button_produced_quantity.text) >= int(self.button_quantity.text):
            self.button_ref.background_color = color_light_green
            self.button_name.background_color = color_light_green
            self.button_material_name.background_color = color_light_green
            self.button_quantity.background_color = color_light_green
            self.button_produced_quantity.background_color = color_light_green
            self.button_date_modified.background_color = color_light_green
            self.button_due_date.background_color = color_light_green

        elif self.check_part_in_tasks():
            self.button_ref.background_color = color_orange
            self.button_name.background_color = color_orange
            self.button_material_name.background_color = color_orange
            self.button_quantity.background_color = color_orange
            self.button_produced_quantity.background_color = color_orange
            self.button_date_modified.background_color = color_orange
            self.button_due_date.background_color = color_orange
        else:
            self.button_ref.background_color = color_dark_blue
            self.button_name.background_color = color_dark_blue
            self.button_material_name.background_color = color_dark_blue
            self.button_quantity.background_color = color_dark_blue
            self.button_produced_quantity.background_color = color_dark_blue
            self.button_date_modified.background_color = color_dark_blue
            self.button_due_date.background_color = color_dark_blue

    def set_focus_text_input(self, *args): # Return to this later, autoselect and focus all textinput widgets
        self.ti_new_value.focus = True
        self.ti_new_value.select_all()

    def check_part_in_tasks(self):
        is_part_in_tasks = False

        for order_part_rowid in production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sideframe.order_parts_in_tasks:
            if self.rowid == order_part_rowid:
                is_part_in_tasks = True

        return is_part_in_tasks


class OrderPartspage_Sideframe(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        #TODO Develop way to automatically create tasks from multiple orders simultaneously

        self.orientation = 'vertical'
        self.orders = set([f'{order_part[6]}, {order_part[7]}' for order_part in load_order_parts_from_database()])
        self.added_orders = []
        self.order_parts = [list(elem) for elem in load_order_parts_from_database()]
        self.added_order_parts = []
        self.select_all = False
        self.order_parts_in_tasks = self.get_rowids_order_parts_in_tasks()

        self.layout_orders = BoxLayout(orientation='vertical', size_hint_y=0.35)
        self.layout_header = BoxLayout(orientation='horizontal', size_hint_y=0.08, padding=5)
        self.layout_order_parts = BoxLayout(orientation='vertical', padding=5)

        self.layout_orders_list = ScrollView()
        self.layout_scroll_orders = BoxLayout(orientation='vertical', size_hint_y=None, padding=5)
        self.layout_scroll_orders.bind(minimum_height=self.layout_scroll_orders.setter('height'))
        self.checkbox_select_all = CheckBox(size_hint_x=0.035)
        self.label_ref = Label(text='Ref.', halign='center', size_hint_x=0.1)
        self.label_name = Label(text='Nome', halign='center', size_hint_x=0.32)
        self.label_material = Label(text='Material', halign='center', size_hint_x=0.17)
        self.label_quantity = Label(text='Qtd.', halign='center', size_hint_x=0.05)
        self.label_produced_quantity = Label(text='Qtd\nProd.', halign='center', size_hint_x=0.05)
        self.label_date_modified = Label(text='Data\nModif.', halign='center', size_hint_x=0.1)
        self.label_due_date = Label(text='Data\nEntrega', halign='center', size_hint_x=0.1)
        self.label_additional_operations = Label(text='Oper.\nAdic.', halign='center', size_hint_x=0.03)
        self.layout_order_parts_list = ScrollView()
        self.layout_scroll_order_parts = BoxLayout(orientation='vertical', size_hint_y=None, padding=5)
        self.layout_scroll_order_parts.bind(minimum_height=self.layout_scroll_order_parts.setter('height'))

        self.checkbox_select_all.bind(on_press=self.select_all_order_parts)

        self.layout_orders_list.add_widget(self.layout_scroll_orders)
        self.layout_order_parts_list.add_widget(self.layout_scroll_order_parts)
        self.layout_orders.add_widget(self.layout_orders_list)
        self.layout_order_parts.add_widget(self.layout_order_parts_list)
        self.layout_header.add_widget(self.checkbox_select_all)
        self.layout_header.add_widget(self.label_ref)
        self.layout_header.add_widget(self.label_name)
        self.layout_header.add_widget(self.label_material)
        self.layout_header.add_widget(self.label_quantity)
        self.layout_header.add_widget(self.label_produced_quantity)
        self.layout_header.add_widget(self.label_date_modified)
        self.layout_header.add_widget(self.label_due_date)
        self.layout_header.add_widget(self.label_additional_operations)
        self.add_widget(self.layout_orders)
        self.add_widget(self.layout_header)
        self.add_widget(self.layout_order_parts)

        self.update_orders_display()
        # print('ORDER PARTS IN TASKS')
        # print(self.order_parts_in_tasks)

    def update_orders_display(self, *args):
        if len(self.added_orders) != 0:
            for added_order in self.added_orders:
                self.layout_scroll_orders.remove_widget(added_order)

            self.added_orders = []

        self.orders = set([(order_part[5], order_part[6], order_part[7]) for order_part in load_order_parts_from_database()])

        for order in self.orders:
            self.o = Order(order[0], order[1], order[2])
            # self.layout_scroll_orders.add_widget(self.o)
            self.added_orders.append(self.o)

        self.added_orders = sorted(self.added_orders, key=operator.attrgetter('client'))

        for order in self.added_orders:
            self.layout_scroll_orders.add_widget(order)

    def update_order_parts_display(self, order_num, *args):
        if len(self.added_order_parts) != 0:
            for added_order_part in self.added_order_parts:
                self.layout_scroll_order_parts.remove_widget(added_order_part)

            self.added_order_parts = []

        self.order_parts = [list(elem) for elem in load_order_parts_from_database()]

        for order_part in self.order_parts:
            if order_num == order_part[5]:
                op = OrderPart(order_part[0], order_part[1], order_part[2], order_part[3], order_part[4], order_part[5],
                              order_part[6], order_part[7], order_part[8], order_part[9], order_part[10], order_part[-1])
                # print((order_part[1], order_part[11]))
                self.layout_scroll_order_parts.add_widget(op)
                self.added_order_parts.append(op)

    def select_all_order_parts(self, *args):
        #TODO Reset checkbox all state when a new order is selected

        self.select_all = not self.select_all

        if self.select_all:
            state = 'down'
        else:
            state = 'normal'

        for part in self.added_order_parts:
            part.is_selected = self.select_all
            part.checkbox_selected.state = state

    def get_rowids_order_parts_in_tasks(self):
        rowid_list = []

        for db_task in load_tasks_from_database():
            for order_rowid in db_task[10].split(','):
                rowid_list.append(order_rowid)

        return rowid_list


class NewOrderPart(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.orientation = 'horizontal'
        self.size_hint_y = None
        self.height = 35

        self.button_x = Button(text='X', font_size='16', background_color=color_dark_red, size_hint_x=0.05, padding=(5, 5))
        self.layout_ref_name = BoxLayout(orientation='horizontal', size_hint_x=0.3)
        self.drp_ref = DropDown()
        self.button_ref = Button(text='', font_size='16', halign='right', size_hint_x=0.35, background_color=color_light_blue)
        self.drp_name = DropDown()
        self.button_name = Button(text='', font_size='16', halign='right', size_hint_x=0.65, background_color=color_light_blue)
        self.ti_name = TextInput(size_hint_x=0.65)
        self.drp_material = DropDown()
        self.button_material = Button(text='', font_size='16', halign='right', size_hint_x=0.28, background_color=color_light_blue)
        self.ti_quantity = TextInput(size_hint_x=0.1)
        self.checkbox_additional_operations = CheckBox(size_hint_x=0.03)

        self.button_x.bind(on_press=self.delete_new_order_part)
        self.button_ref.bind(on_release=self.drp_ref.open)
        self.drp_ref.bind(on_select=lambda instance, new_val: setattr(self.button_ref, 'text', new_val))
        self.button_name.bind(on_press=self.drp_name.open)
        self.drp_name.bind(on_select=lambda instance, new_val: setattr(self.button_name, 'text', new_val))
        self.button_material.bind(on_release=self.drp_material.open)
        self.drp_material.bind(on_select=lambda instance, new_val: setattr(self.button_material, 'text', new_val))

        button_no_ref = Button(text='S/ Ref', font_size='16', size_hint_y=None, height=30, background_color=color_dark_blue)
        button_no_ref.bind(on_release=lambda x: self.update_widgets(button_no_ref))
        button_with_ref = Button(text='C/ Ref', font_size='16', size_hint_y=None, height=30, background_color=color_dark_blue)
        button_with_ref.bind(on_release=lambda x: self.update_widgets(button_with_ref))

        self.drp_ref.add_widget(button_no_ref)
        self.drp_ref.add_widget(button_with_ref)

        for part in load_parts_from_database():

            part_name = part[1]

            button_part_name = Button(text=part_name, font_size='16', size_hint_y=None, height=30, background_color=color_dark_blue)
            button_part_name.bind(on_release=lambda part_object: self.select_part(part_object.text))

            self.drp_name.add_widget(button_part_name)

        for material in load_materials_from_database():
            material_type = material[1]
            material_spec = str(material[2])
            material_thickness = str(material[3])
            material_client = material[11]

            material_name = material_type + ' ' + material_spec + ', ' + material_thickness + 'mm ' + material_client

            button_material_name = Button(text=material_name, font_size='16', size_hint_y=None, height=30, background_color=color_dark_blue)
            button_material_name.bind(on_release=lambda material_object: self.drp_material.select(material_object.text))

            self.drp_material.add_widget(button_material_name)

        self.add_widget(self.button_x)
        self.add_widget(self.layout_ref_name)
        self.layout_ref_name.add_widget(self.button_ref)
        self.layout_ref_name.add_widget(self.button_name)
        self.add_widget(self.button_material)
        self.add_widget(self.ti_quantity)
        self.add_widget(self.checkbox_additional_operations)


    def delete_new_order_part(self, *args):
        production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sidebar.layout_popup_new_order.layout_scroll.remove_widget(self)
        del production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sidebar.layout_popup_new_order.added_new_order_parts[
            production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sidebar.layout_popup_new_order.added_new_order_parts.index(self)]
        # print(production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sidebar.layout_popup_new_order.added_new_order_parts)


    def update_widgets(self, button_no_ref, *args):
        previous_value = self.button_ref.text
        self.drp_ref.select(button_no_ref.text)

        # print((previous_value, button_no_ref.text))

        if button_no_ref.text == 'S/ Ref':
            if previous_value == 'S/ Ref':
                pass
            else:
                self.layout_ref_name.remove_widget(self.button_name)
                self.layout_ref_name.add_widget(self.ti_name)
        else:
            if previous_value == 'S/ Ref':
                self.layout_ref_name.remove_widget(self.ti_name)
                self.layout_ref_name.add_widget(self.button_name)


    def select_part(self, part_name, *args):
        self.drp_name.select(part_name)

        part_info = get_part_info_by_name(part_name) #weight, material_ref, time, ref
        part_ref = part_info[3]
        part_material_ref = part_info[1]
        part_material_name = get_material_name(part_material_ref)

        self.button_ref.text = part_ref
        self.button_material.text = part_material_name


class NewOrderpage(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.added_new_order_parts = []

        self.orientation = 'vertical'

        self.layout_order_info = BoxLayout(orientation='horizontal', padding=5, size_hint_y=0.15)
        self.layout_order_part_header = BoxLayout(orientation='horizontal', padding=5, size_hint_y=0.2)
        self.layout_order_parts_list = ScrollView()
        self.layout_scroll = BoxLayout(orientation='vertical', size_hint_y=None)
        self.layout_scroll.bind(minimum_height=self.layout_scroll.setter('height'))
        # Futuramente alinhar melhor as labels com a posição dos widgets da classe NewOrderPart
        self.label_order_num = Label(text='Num Enc. Interna: ')
        self.ti_order_num = TextInput()
        self.label_order_num_client = Label(text='Num Enc. Cliente: ')
        self.ti_order_num_client = TextInput()
        self.label_client = Label(text='Cliente: ')
        self.ti_client = TextInput()
        self.label_due_date = Label(text='Data Entrega: ')
        self.ti_due_date = TextInput()
        self.label_empty = Label(text='', size_hint_x=0.05)
        self.label_ref = Label(text='Ref', size_hint_x=0.11)
        self.label_name = Label(text='Nome', size_hint_x=0.2)
        self.label_material = Label(text='Material', size_hint_x=0.28)
        self.label_quantity = Label(text='Qtd.', size_hint_x=0.1)
        self.label_additional_operations = Label(text='Oper.\nAdic.', size_hint_x=0.03)
        self.button_add_new_order_part = Button(text='+', font_size='16', background_color=color_light_blue, size_hint_x=0.05, size_hint_y=None, height=35)
        self.button_submit_order = Button(text='Submeter Encomenda', font_size='16', background_color=color_dark_blue, size_hint_x=0.3, size_hint_y=None, height=35,
                                          halign='center')

        self.button_add_new_order_part.bind(on_press=self.add_order_part)
        self.button_submit_order.bind(on_press=self.create_new_order)

        self.layout_order_info.add_widget(self.label_order_num)
        self.layout_order_info.add_widget(self.ti_order_num)
        self.layout_order_info.add_widget(self.label_order_num_client)
        self.layout_order_info.add_widget(self.ti_order_num_client)
        self.layout_order_info.add_widget(self.label_client)
        self.layout_order_info.add_widget(self.ti_client)
        self.layout_order_info.add_widget(self.label_due_date)
        self.layout_order_info.add_widget(self.ti_due_date)
        self.layout_order_part_header.add_widget(self.label_empty)
        self.layout_order_part_header.add_widget(self.label_ref)
        self.layout_order_part_header.add_widget(self.label_name)
        self.layout_order_part_header.add_widget(self.label_material)
        self.layout_order_part_header.add_widget(self.label_quantity)
        self.layout_order_part_header.add_widget(self.label_additional_operations)
        self.add_widget(self.layout_order_info)
        self.add_widget(self.layout_order_part_header)
        self.layout_order_parts_list.add_widget(self.layout_scroll)
        self.layout_scroll.add_widget(self.button_add_new_order_part)
        self.add_widget(self.layout_order_parts_list)
        self.add_widget(self.button_submit_order)

        self.add_order_part()


    def add_order_part(self, *args):
        new_order_part = NewOrderPart()
        self.layout_scroll.remove_widget(self.button_add_new_order_part)
        self.layout_scroll.add_widget(new_order_part)
        self.added_new_order_parts.append(new_order_part)
        self.layout_scroll.add_widget(self.button_add_new_order_part)

    def create_new_order(self, *args):

        if self.ti_order_num.text == '' or self.ti_client == '': # Falta colocar impedimentos de criar encomenda aqui
            pass
        else:
            order_num = self.ti_order_num.text
            order_num_client = self.ti_order_num_client.text
            client = self.ti_client.text
            date_modified = f'{datetime.datetime.now().day}/{datetime.datetime.now().month}/{datetime.datetime.now().year}'
            due_date = self.ti_due_date.text
            for added_new_order_part in self.added_new_order_parts:
                if added_new_order_part.button_ref.text == '' or (added_new_order_part.button_ref.text == 'S/ Ref' and added_new_order_part.ti_name.text == '') or (added_new_order_part.button_ref.text == 'C/ Ref' and added_new_order_part.button_name.text == '') or added_new_order_part.ti_quantity.text == '' or added_new_order_part.button_material.text == '':
                    print('Invalid')
                else:
                    ref = added_new_order_part.button_ref.text
                    name = added_new_order_part.button_name.text
                    if added_new_order_part.button_ref.text == 'S/ Ref':
                        name = added_new_order_part.ti_name.text
                    material_name = added_new_order_part.button_material.text
                    material_type = material_name.split(' ')[0]
                    material_spec = material_name.split(' ')[1].split(',')[0]
                    material_thickness = float(material_name.split(' ')[2].split('mm')[0])
                    material_client = material_name.split(' ')[3]
                    material_ref = get_material_ref(material_type, material_spec, material_thickness, material_client)
                    quantity = int(added_new_order_part.ti_quantity.text)
                    produced_quantity = 0
                    additional_operations = False
                    if added_new_order_part.checkbox_additional_operations.state == 'down':
                        additional_operations = True


                    insert_new_order_part_database(ref, name, material_ref, quantity, produced_quantity, order_num, order_num_client, client, date_modified, due_date, additional_operations)

                    production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sideframe.update_orders_display()
                    production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sideframe.update_order_parts_display(order_num)
                    production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sidebar.pop_new_order.dismiss()

        production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sidebar.set_num_current_orders()


class OrderPartspage_Sidebar(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.orientation = 'vertical'
        self.size_hint_x = 0.08

        self.images = []
        self.num_current_orders = 0

        self.button_return = Button(text='<<<', font_size='16', halign='center', background_color=color_light_blue)
        self.button_new_order = Button(text='Nova\nEncomenda', font_size='14', halign='center', background_color=color_light_blue)
        self.button_import_orders = Button(text='Inserir\nEncomenda(s) PDF', font_size='14', halign='center', background_color=color_light_blue)
        self.button_delete_item = Button(text='Apagar\nItens Selecionados', font_size='14', halign='center', background_color=color_light_blue)
        self.button_generate_tasks = Button(text='Gerar Tarefas\nItens Selecionados', font_size='14', halign='center', background_color=color_light_blue)
        self.button_generate_tag = Button(text='Gerar Etiquetas\nItens Selecionados', font_size='14', halign='center', background_color=color_light_blue)
        self.button_export_list_order_parts = Button(text='Exportar Listagem\nItens Selecionados', font_size='14', halign='center', background_color=color_light_blue)
        # self.button_email_selected_items = Button(text='Enviar Email\nItens Selecionados', font_size='14', halign='center', background_color=color_light_blue)
        self.label_empty = Label(text='', size_hint_y=1)
        self.label_num_current_orders = Label(text=f'{self.num_current_orders}\nEncomendas\nem\nCurso', halign='center')

        self.button_return.bind(on_press=self.close_popup_window)
        self.button_new_order.bind(on_press=self.display_popup_new_order)
        self.button_import_orders.bind(on_press=self.insert_orders_from_PDF)
        self.button_delete_item.bind(on_press=self.display_popup_warning)
        self.button_generate_tasks.bind(on_press=self.generate_tasks_selected_parts)
        self.button_generate_tag.bind(on_press=self.generate_order_tags)
        self.button_export_list_order_parts.bind(on_press=self.export_list_order_parts)
        # self.button_email_selected_items.bind(on_press=self.display_popup_send_email)

        self.add_widget(self.button_return)
        self.add_widget(self.button_new_order)
        self.add_widget(self.button_import_orders)
        self.add_widget(self.button_delete_item)
        self.add_widget(self.button_generate_tasks)
        self.add_widget(self.button_generate_tag)
        self.add_widget(self.button_export_list_order_parts)
        # self.add_widget(self.button_email_selected_items)
        self.add_widget(self.label_empty)
        self.add_widget(self.label_num_current_orders)

        self.set_num_current_orders()

    def close_popup_window(self, *args):
        production_planning.homepage.layout_sidebar.pop_order_parts.dismiss()

    def display_popup_new_order(self, *args):
        self.layout_popup_new_order = NewOrderpage()
        self.pop_new_order = Popup(title='Nova Encomenda', content=self.layout_popup_new_order, size_hint=(0.9, 0.6))

        self.pop_new_order.open()

    def insert_orders_from_PDF(self, *args):
        task_list = return_order_parts_from_PDF()
        print(f'NUM TOTAL PARTS FROM PDF: {len(task_list)}')
        for task in task_list:
            client = task[0]
            order_num = task[1]
            ref = task[2]
            name = task[3]
            quantity = task[4]
            part_info = get_part_info_by_ref(ref)
            if part_info[1] == '':
                material_ref = task[5]
            else:
                material_ref = part_info[1]
            produced_quantity = 0
            order_num_client = 'N/A'
            date_modified = f'{datetime.datetime.now().day}/{datetime.datetime.now().month}/{datetime.datetime.now().year}'
            due_date = f'{datetime.datetime.now().day}/{datetime.datetime.now().month}/{datetime.datetime.now().year}'

            # print((ref, name, material_ref, quantity, produced_quantity, order_num, client, date_modified, due_date))
            insert_new_order_part_database(ref, name, material_ref, quantity, produced_quantity, order_num, order_num_client, client, date_modified, due_date, additional_operations=False)
            production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sideframe.update_orders_display()
            production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sideframe.update_order_parts_display(order_num)
            production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sidebar.set_num_current_orders()

    def display_popup_new_part(self, *args):
        self.layout_popup_new_part = NewPartpage()
        self.pop_new_part = Popup(title='Novo Artigo', content=self.layout_popup_new_part, size_hint=(0.60, 0.35))

        self.pop_new_part.open()

    def display_popup_warning(self, *args):
        self.popup_warning_delete = PopupWarningMessage(message='Isto irá apagar permanentemente as peças selecionadas!', type='choice', continue_func=self.delete_selected_part)
        self.popup_warning_delete.open()

    def delete_selected_part(self, *args):
        for order_part in production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sideframe.added_order_parts:
            if order_part.is_selected:
                remove_order_part_database(order_part.rowid)
                production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sideframe.update_orders_display()
                production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sideframe.update_order_parts_display(order_part.order_num)

                production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sideframe.checkbox_select_all.state = 'normal'
                production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sideframe.select_all = False

        self.update_button_order_info()
        self.set_num_current_orders()

    def generate_tasks_selected_parts(self, *args):
        parts_by_material_ref = {}

        for order_part in production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sideframe.added_order_parts:
            if order_part.is_selected and int(order_part.produced_quantity) < int(order_part.quantity) and not self.check_order_part_in_any_task(order_part.rowid):
                if order_part.material_ref not in parts_by_material_ref.keys():
                    parts_by_material_ref.update({order_part.material_ref: []})

                parts_by_material_ref[order_part.material_ref].append(order_part.rowid)

        for material_ref in parts_by_material_ref:
            current_path = ''
            original_path = current_path
            machine = ''
            notes = ''
            estimated_sheets_required = 0.1
            estimated_time = 15
            start_date = ''
            end_date = ''
            priority = 0
            order_parts = ','.join(parts_by_material_ref[material_ref])
            aggregated_tasks = ''

            insert_new_task_database(current_path, original_path, machine, material_ref, notes, estimated_sheets_required, estimated_time, start_date, end_date, priority,
                                     order_parts, aggregated_tasks)

    def check_order_part_in_any_task(self, rowid):
        '''Checks if the rowid passed in is already added to one of the database tasks'''
        for task in load_tasks_from_database():
            for task_order_part_rowid in task[10].split(','):
                if task_order_part_rowid == rowid:
                    return True
        return False

    def update_button_order_info(self, *args):
        try:
            any_order_part = production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sideframe.added_order_parts[0]

            for order in production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sideframe.added_orders:
                if any_order_part.order_num == order.order_num and any_order_part.order_num_client == order.order_num_client and any_order_part.client == order.client:
                    order.is_selected = True
                if order.is_selected:
                    order.button_order.background_color = color_dark_green
                    order.calculate_num_order_parts()
                    order.update_button_info_text()
                    order.set_button_idle_color()
        except IndexError:
            pass

    def delete_images_from_folder(self):
        for image in self.images:
            if image != 'Etiqueta.jpg':
                os.remove(image)

    def add_images_to_word(self, doc_word, num_empty_order_images):

        p = doc_word.add_paragraph()
        r = p.add_run()

        for _ in range(num_empty_order_images):
            self.images.append('Etiqueta.jpg')

        image_gen = self.image_generator(self.images)
        num_pages = int(len(self.images) / 12)
        for _ in range(num_pages):
            for __ in range(12):
                image = next(image_gen)
                r.add_picture(image, width=Cm(6.75))
            p = doc_word.add_paragraph()
            r = p.add_run()

    def image_generator(self, images):
        for image in images:
            yield image

    def create_order_image(self, image_path, client, order_num_client, part_description, part_qtty, order_counter):
        current_date = f'{datetime.datetime.now().day}/{datetime.datetime.now().month}/{datetime.datetime.now().year}'
        multiline_part_description = part_description[:34] + '\n' + part_description[34:]

        blank_order_image = image_path
        image = Image.open(blank_order_image)
        text = ImageDraw.ImageDraw(image, mode="RGB")

        client_text_position = (135, 132)
        order_num_text_position = (438, 322)
        part_description_text_position = (110, 228)
        part_qtty_text_position = (90, 322)
        date_position = (410, 417)

        fontsize = 25
        fontsize_description = 24
        font = ImageFont.truetype("arial.ttf", fontsize)
        font_description = ImageFont.truetype("arial.ttf", fontsize_description)

        text.text(client_text_position, client, fill=(0, 0, 0), font=font)
        text.text(order_num_text_position, order_num_client, fill=(0, 0, 0), font=font)
        text.text(part_description_text_position, multiline_part_description, fill=(0, 0, 0), font=font_description)
        text.text(part_qtty_text_position, part_qtty + ' Un.', fill=(0, 0, 0), font=font)
        text.text(date_position, current_date, fill=(0, 0, 0), font=font)

        # new_width = 300
        # new_height = new_width - 90
        # image = image.resize((new_width, new_height), Image.ANTIALIAS)

        image.save(f'{order_counter}.jpg')

        return f'{order_counter}.jpg'

    def generate_order_tags(self, *args):
        order_counter = 1
        for order_part in production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sideframe.added_order_parts:
            if order_part.is_selected:
                order_image = self.create_order_image('Etiqueta.jpg', order_part.client, order_part.order_num_client, order_part.name, order_part.quantity, str(order_counter))
                self.images.append(order_image)
                order_counter += 1

        doc_word = Document()

        sections = doc_word.sections
        # top_margin = 2
        # bottom_margin = 2
        # left_margin = 0.5
        # right_margin = 0.5

        for section in sections:

            section.top_margin = Cm(0.45)
            section.bottom_margin = Cm(0.45)
            section.left_margin = Cm(0.45)
            section.right_margin = Cm(0.45)

            section.page_height = Mm(297)
            section.page_width = Mm(210)

            new_height, new_width = section.page_height, section.page_width
            section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
            section.page_width = new_height
            section.page_height = new_width


        num_empty_order_images = 15 - (len(self.images) % 15)

        self.add_images_to_word(doc_word, num_empty_order_images)

        self.delete_images_from_folder()

        doc_word.save('Etiquetas_Produção.docx')

        layout_popup_inform_success = BoxLayout(orientation='vertical')
        layout_popup_inform_success.add_widget(Label(text='Etiquetas geradas com sucesso'))
        button_open_doc = Button(text='Abrir', font_size='16', halign='center', background_color=color_light_blue)
        button_open_doc.bind(on_press=lambda x: os.startfile("Etiquetas_Produção.docx"))
        layout_popup_inform_success.add_widget(button_open_doc)

        popup_inform_generated_successfully = Popup(title='Sucesso', content=layout_popup_inform_success, size_hint=(0.2, 0.15))
        popup_inform_generated_successfully.open()

    def export_list_order_parts(self, *args):

        pdf = FPDF()
        pdf.add_page()

        pdf.set_font('Times', size=10)
        effective_page_width = pdf.w - 2*pdf.l_margin
        col_width = effective_page_width/4

        pdf.set_font('Times', size=5)
        pdf.cell(effective_page_width, pdf.font_size, '(Ficheiro Gerado Automaticamente)', align='C')

        pdf.ln(10)

        order_part_info = []
        client = ''
        order_num_client = ''

        for order_part in production_planning.homepage.layout_sidebar.layout_popup_order_parts.layout_sideframe.added_order_parts:
            if order_part.is_selected:
                order_part_info.append((order_part.name, order_part.quantity, order_part.produced_quantity, order_part.additional_operations))
                client = order_part.client
                order_num_client = order_part.order_num_client

        pdf.set_font('Times', 'B', size=12)
        pdf.cell(effective_page_width, 0.0, f'Listagem de Estados - Enc n°: {order_num_client}, {client}', align='C')

        pdf.ln(10)
        text_height = pdf.font_size

        pdf.set_font('Times', 'B', size=6)
        pdf.set_fill_color(169, 169, 169)
        pdf.cell(col_width * 2.5, text_height, 'Ref/Nome', border=1, fill=1, align='C')
        pdf.cell(col_width * 0.5, text_height, 'Quantidade', border=1, fill=1, align='C')
        pdf.cell(col_width * 0.5, text_height, 'Quantidade\nProduzida', border=1, fill=1, align='C')
        pdf.cell(col_width * 0.5, text_height, 'Operações\nAdicionais', border=1, fill=1, align='C')

        pdf.ln(text_height)
        pdf.set_font('Times', size=6)

        for info in order_part_info:
            if int(info[2]) >= int(info[1]) and not info[3]:
                pdf.set_fill_color(0, 179, 0)
                fill = 1
            elif int(info[2]) >= int(info[1]) and info[3]:
                pdf.set_fill_color(179, 143, 0)
                fill = 1
            else:
                fill = 0

            pdf.cell(col_width * 2.5, text_height, str(info[0]).strip('\n'), border=1, fill=fill, align='C')
            pdf.cell(col_width * 0.5, text_height, str(info[1]), border=1, fill=fill, align='C')
            pdf.cell(col_width * 0.5, text_height, str(info[2]), border=1, fill=fill, align='C')
            has_additional_operations = 'Não'
            if info[3]:
                has_additional_operations = 'Sim'
            pdf.cell(col_width * 0.5, text_height, has_additional_operations, border=1, fill=fill, align='C')

            pdf.ln(text_height)

        filename = remove_invalid_file_name_characters(f'Listagem_{order_num_client}_{client}')

        pdf.output(f'{filename}.pdf', 'F')

    # def display_popup_send_email(self, *args):
    #     self.layout_popup_send_email = BoxLayout(orientation='vertical')
    #     self.pop_send_email = Popup(title='Enviar email', content=self.layout_popup_send_email, size_hint=(0.60, 0.60))
    #
    #     layout_line_receivers = BoxLayout(orientation='horizontal', size_hint_y=0.2, padding=(5, 5))
    #     layout_line_topic = BoxLayout(orientation='horizontal', size_hint_y=0.2, padding=(5, 5))
    #     layout_line_message = BoxLayout(orientation='vertical', padding=(5, 5))
    #     layout_line_buttons = BoxLayout(orientation='horizontal', size_hint_y=0.3, padding=(5, 5))
    #
    #     label_receivers = Label(text='Destinatários: ', size_hint_x=0.2)
    #     ti_receivers = TextInput()
    #     label_topic = Label(text='Assunto: ', size_hint_x=0.2)
    #     ti_topic = TextInput()
    #     ti_message = TextInput()
    #     button_return = Button(text='<<<', font_size='16', halign='center', background_color=color_light_blue, padding=(10, 10))
    #     button_send = Button(text='Enviar', font_size='16', halign='center', background_color=color_light_blue, padding=(10, 10))
    #
    #     button_return.bind(on_press=self.pop_send_email.dismiss)
    #     button_send.bind(on_press=lambda x: self.send_email(ti_receivers.text, ti_topic.text, ti_message.text))
    #
    #     layout_line_receivers.add_widget(label_receivers)
    #     layout_line_receivers.add_widget(ti_receivers)
    #     layout_line_topic.add_widget(label_topic)
    #     layout_line_topic.add_widget(ti_topic)
    #     layout_line_message.add_widget(ti_message)
    #     layout_line_buttons.add_widget(button_return)
    #     layout_line_buttons.add_widget(button_send)
    #     self.layout_popup_send_email.add_widget(layout_line_receivers)
    #     self.layout_popup_send_email.add_widget(layout_line_topic)
    #     self.layout_popup_send_email.add_widget(layout_line_message)
    #     self.layout_popup_send_email.add_widget(layout_line_buttons)
    #
    #     self.pop_send_email.open()

    def send_email(self, receivers, topic, message, *args): # Falta adicionar funcionalidade ao email e averiguar como enviar para mais do que um destinatário
        print((receivers, topic, message))

    def set_num_current_orders(self, *args):
        current_orders = set([f'{order_part[5]}, {order_part[6]}, {order_part[7]}' for order_part in load_order_parts_from_database()])

        self.num_current_orders = len(current_orders)
        self.label_num_current_orders.text = f'{self.num_current_orders}\nEncomendas\nem\nCurso'


class OrderPartspage(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.orientation = 'horizontal'

        self.layout_sidebar = OrderPartspage_Sidebar()
        self.layout_sideframe = OrderPartspage_Sideframe()

        self.add_widget(self.layout_sidebar)
        self.add_widget(self.layout_sideframe)


class Part(BoxLayout):
    def __init__(self, ref, name, weight, material_ref, time, client, date_modified, id, is_selected=False, **kwargs):
        super().__init__(**kwargs)

        self.orientation = 'horizontal'
        self.size_hint_y = None
        self.height = 35

        self.ref = ref
        self.name = name
        self.weight = str(weight)
        self.material_name = get_material_name(material_ref)
        self.time = return_formatted_time(time)
        self.client = client
        self.date_modified = date_modified
        self.id = str(id)
        self.is_selected = is_selected

        self.checkbox_selected = CheckBox(size_hint_x=0.03)
        self.button_ref = Button(text=self.ref, font_size='18', background_color=color_dark_blue, size_hint_x=0.1)
        self.button_name = Button(text=self.name, font_size='18', background_color=color_dark_blue, size_hint_x=0.25)
        self.button_weight = Button(text=self.weight, font_size='18', background_color=color_dark_blue, size_hint_x=0.07)
        self.button_material_name = Button(text=self.material_name, font_size='18', background_color=color_dark_blue, size_hint_x=0.15)
        self.button_time = Button(text=self.time, font_size='18', background_color=color_dark_blue, size_hint_x=0.07)
        self.button_client = Button(text=self.client, font_size='18', background_color=color_dark_blue, size_hint_x=0.1)
        self.button_date_modified = Button(text=self.date_modified, font_size='18', background_color=color_dark_blue, size_hint_x=0.1)

        self.checkbox_selected.bind(active=self.on_selector_active)
        self.button_ref.bind(on_press=lambda x: self.display_popup_change_value('ref', self.id))
        self.button_name.bind(on_press=lambda x: self.display_popup_change_value('name', self.id))
        self.button_weight.bind(on_press=lambda x: self.display_popup_change_value('weight', self.id))
        self.button_material_name.bind(on_press=lambda x: self.display_popup_change_value('material_name', self.id))
        self.button_time.bind(on_press=lambda x: self.display_popup_change_value('time', self.id))
        self.button_client.bind(on_press=lambda x: self.display_popup_change_value('client', self.id))

        self.add_widget(self.checkbox_selected)
        self.add_widget(self.button_ref)
        self.add_widget(self.button_name)
        self.add_widget(self.button_weight)
        self.add_widget(self.button_material_name)
        self.add_widget(self.button_time)
        self.add_widget(self.button_client)
        self.add_widget(self.button_date_modified)

    def on_selector_active(self, *args):
        if self.is_selected:
            self.is_selected = False
        else:
            self.is_selected = True

    def display_popup_change_value(self, field, rowid, *args):

        self.layout_new_value = BoxLayout(orientation='vertical')
        self.layout_horizontal_line = BoxLayout(orientation='horizontal', padding=(10, 10))
        self.label_new_value = Label(text='Novo valor: ', font_size='16', halign='right', size_hint_x=0.3)
        if field == 'material_name':
            self.drp_material = DropDown()
            self.button_material = Button(text='', font_size='16', halign='right', size_hint_y=0.75)
            self.button_material.bind(on_release=self.drp_material.open)
            self.drp_material.bind(on_select=lambda instance, new_val: setattr(self.button_material, 'text', new_val))
            self.button_submit = Button(text='Submeter', font_size='18', background_color=color_light_green)

            for material in load_materials_from_database():
                material_type = material[1]
                material_spec = str(material[2])
                material_thickness = str(material[3])
                material_client = material[11]

                material_name = material_type + ' ' + material_spec + ', ' + material_thickness + 'mm ' + material_client

                button_material_name = Button(text=material_name, font_size='16', size_hint_y=None, height=30, background_color=color_dark_blue)
                button_material_name.bind(on_release=lambda material_object: self.drp_material.select(material_object.text))

                self.drp_material.add_widget(button_material_name)

            self.layout_horizontal_line.add_widget(self.label_new_value)
            self.layout_horizontal_line.add_widget(self.button_material)
            self.layout_new_value.add_widget(self.layout_horizontal_line)
            self.layout_new_value.add_widget(self.button_submit)

            self.button_submit.bind(on_press=lambda x: self.change_value(self.button_material.text, field, rowid))

            self.pop_change_value = Popup(title='Novo valor: ', content=self.layout_new_value, size_hint=(0.30, 0.20))
            self.pop_change_value.open()

        else:
            if field == 'ref':
                self.ti_new_value = TextInput(text=self.button_ref.text)
            elif field == 'name':
                self.ti_new_value = TextInput(text=self.button_name.text)
            else:
                self.ti_new_value = TextInput()

            self.button_submit = Button(text='Submeter', font_size='18', background_color=color_light_green)

            self.layout_horizontal_line.add_widget(self.label_new_value)
            self.layout_horizontal_line.add_widget(self.ti_new_value)
            self.layout_new_value.add_widget(self.layout_horizontal_line)
            self.layout_new_value.add_widget(self.button_submit)

            self.button_submit.bind(on_press=lambda x: self.change_value(self.ti_new_value.text, field, rowid))

            self.pop_change_value = Popup(title='Novo valor: ', content=self.layout_new_value, size_hint=(0.30, 0.20))
            self.pop_change_value.open()

    def change_value(self, value, field, rowid, *args):
        '''Falta atualizar os outros campos quando um valor de cálculos é alterado'''
        if field == 'ref':
            change_value_part_database(value, field, rowid)
            self.button_ref.text = str(value)
        elif field == 'name':
            change_value_part_database(value, field, rowid)
            self.button_name.text = str(value)
        elif field == 'weight':
            value = float(value)
            change_value_part_database(value, field, rowid)
            self.button_weight.text = str(value)
        elif field == 'material_name':
            material_type = value.split(' ')[0]
            material_spec = value.split(' ')[1].split(',')[0]
            material_thickness = float(value.split(' ')[2].split('mm')[0])
            material_client = value.split(' ')[3]
            material_ref = get_material_ref(material_type, material_spec, material_thickness, material_client)
            change_value_part_database(material_ref, 'material_ref', rowid)
            material_name = get_material_name(material_ref)
            self.button_material_name.text = material_name
        elif field == 'time':
            value = int(value)
            change_value_part_database(value, field, rowid)
            self.button_time.text = return_formatted_time(value)
        elif field == 'client':
            change_value_part_database(value, field, rowid)
            self.button_client.text = str(value)

        date_modified = f'{datetime.datetime.now().day}/{datetime.datetime.now().month}/{datetime.datetime.now().year}'
        self.button_date_modified.text = date_modified
        change_value_part_database(date_modified, 'date_modified', rowid)
        self.pop_change_value.dismiss()


class Partspage_Sideframe(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.parts_list = []

        self.orientation = 'vertical'
        self.padding = 5
        self.layout_labels = BoxLayout(orientation='horizontal', size_hint_y=0.1)
        self.label_empty = Label(text='', size_hint_x=0.03)
        self.label_ref = Label(text='Ref.', font_size='16', size_hint_x=0.1)
        self.label_name = Label(text='Nome', font_size='16', size_hint_x=0.25)
        self.label_weight = Label(text='Peso', font_size='16', size_hint_x=0.07)
        self.label_material_ref = Label(text='Material', font_size='16', size_hint_x=0.15)
        self.label_time = Label(text='Tempo', font_size='16', size_hint_x=0.07)
        self.label_client = Label(text='Cliente', font_size='16', size_hint_x=0.1)
        self.label_date_modified = Label(text='Últ. Modificação', font_size='16', size_hint_x=0.1)
        self.layout_list = ScrollView()
        self.layout_scroll = BoxLayout(orientation='vertical', size_hint_y=None, padding=5)
        self.layout_scroll.bind(minimum_height=self.layout_scroll.setter('height'))

        self.add_widget(self.layout_labels)
        self.layout_labels.add_widget(self.label_empty)
        self.layout_labels.add_widget(self.label_ref)
        self.layout_labels.add_widget(self.label_name)
        self.layout_labels.add_widget(self.label_weight)
        self.layout_labels.add_widget(self.label_material_ref)
        self.layout_labels.add_widget(self.label_time)
        self.layout_labels.add_widget(self.label_client)
        self.layout_labels.add_widget(self.label_date_modified)
        self.layout_list.add_widget(self.layout_scroll)
        self.add_widget(self.layout_list)

        self.update_display()

    def update_display(self, *args):
        if len(self.parts_list) != 0:
            for part in self.parts_list:
                self.layout_scroll.remove_widget(part)

            self.parts_list = []

        for part in load_parts_from_database():
            p = Part(part[0], part[1], part[2], part[3], part[4], part[5], part[6], part[-1])
            self.layout_scroll.add_widget(p)
            self.parts_list.append(p)


class NewPartpage(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.is_client = False

        self.orientation = 'vertical'

        self.layout_line_1 = BoxLayout(orientation='horizontal', padding=(5, 5))
        self.layout_line_2 = BoxLayout(orientation='horizontal', padding=(5, 5))
        self.layout_line_3 = BoxLayout(orientation='horizontal', padding=(5, 5))

        self.label_ref = Label(text='Referência: ', font_size='16', halign='right', size_hint_y=0.75, size_hint_x=0.4)
        self.ti_ref = TextInput(size_hint_y=0.75)
        self.label_name = Label(text='Nome: ', font_size='16', halign='right', size_hint_y=0.75, size_hint_x=0.4)
        self.ti_name = TextInput(size_hint_y=0.75)
        self.label_weight = Label(text='Peso: ', font_size='16', halign='right', size_hint_y=0.75, size_hint_x=0.4)
        self.ti_weight = TextInput(size_hint_y=0.75, size_hint_x = 0.3)
        self.label_material = Label(text='Material: ', font_size='16', halign='right', size_hint_y=0.75, size_hint_x=0.5)
        self.drp_material = DropDown()
        self.button_material = Button(text='', font_size='16', halign='right', size_hint_y=0.75)
        self.button_material.bind(on_release=self.drp_material.open)
        self.drp_material.bind(on_select=lambda instance, new_val: setattr(self.button_material, 'text', new_val))

        for material in load_materials_from_database():
            material_type = material[1]
            material_spec = str(material[2])
            material_thickness = str(material[3])
            material_client = material[11]

            material_name = material_type + ' ' + material_spec + ', ' + material_thickness + 'mm ' + material_client

            button_material_name = Button(text=material_name, font_size='16', size_hint_y=None, height=30, background_color=color_dark_blue)
            button_material_name.bind(on_release=lambda material_object: self.drp_material.select(material_object.text))

            self.drp_material.add_widget(button_material_name)

        self.label_time = Label(text='Tempo(seg): ', font_size='16', halign='right', size_hint_y=0.75, size_hint_x=0.5)
        self.ti_time = TextInput(size_hint_y=0.75, size_hint_x = 0.3)
        self.label_is_client = Label(text='M.P. Cliente?: ', font_size='16', halign='right', size_hint_y=0.75)
        self.check_is_client = CheckBox(size_hint_x=0.2)
        self.label_empty = Label(text='', size_hint_y=0.25)
        self.button_submit = Button(text='Submeter Novo Artigo', font_size='16', halign='center', background_color=color_dark_blue, padding=(10, 10))

        self.check_is_client.bind(active=self.select_is_client)
        self.button_submit.bind(on_press=self.insert_new_part)

        self.add_widget(self.layout_line_1)
        self.add_widget(self.layout_line_2)
        self.add_widget(self.layout_line_3)

        self.layout_line_1.add_widget(self.label_ref)
        self.layout_line_1.add_widget(self.ti_ref)
        self.layout_line_1.add_widget(self.label_name)
        self.layout_line_1.add_widget(self.ti_name)
        self.layout_line_2.add_widget(self.label_weight)
        self.layout_line_2.add_widget(self.ti_weight)
        self.layout_line_2.add_widget(self.label_material)
        self.layout_line_2.add_widget(self.button_material)
        self.layout_line_3.add_widget(self.label_time)
        self.layout_line_3.add_widget(self.ti_time)
        self.layout_line_3.add_widget(self.check_is_client)
        self.add_widget(self.label_empty)
        self.add_widget(self.button_submit)


    def insert_new_part(self, *args):
        ref = self.ti_ref.text
        name = self.ti_name.text
        weight = float(self.ti_weight.text)
        material = self.button_material.text
        material_type = material.split(' ')[0]
        material_spec = material.split(' ')[1].split(',')[0]
        material_thickness = material.split(' ')[2].split('mm')[0]
        material_client = material.split(' ')[3]
        time = int(self.ti_time.text)

        client = 'Prilux'

        if self.is_client:
            client = self.ti_client.text
        material_ref = get_material_ref(material_type, material_spec, material_thickness, material_client)
        date_modified = f'{datetime.datetime.now().day}/{datetime.datetime.now().month}/{datetime.datetime.now().year}'

        insert_new_part_database(ref, name, weight, material_ref, time, client, date_modified)
        production_planning.homepage.layout_sidebar.layout_popup_parts.layout_sideframe.update_display()
        production_planning.homepage.layout_sidebar.layout_popup_parts.layout_sidebar.pop_new_part.dismiss()



    def select_is_client(self, *args):
        self.is_client = not self.is_client

        if self.is_client:
            self.label_client = Label(text='Cliente: ', font_size='16', halign='right', size_hint_y=0.75)
            self.ti_client = TextInput(size_hint_y=0.75)
            self.layout_line_3.add_widget(self.label_client)
            self.layout_line_3.add_widget(self.ti_client)
        else:
            self.layout_line_3.remove_widget(self.label_client)
            self.layout_line_3.remove_widget(self.ti_client)


class Partspage_Sidebar(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.orientation = 'vertical'
        self.size_hint_x = 0.08

        self.button_return = Button(text='<<<', font_size='16', halign='center', background_color=color_light_blue)
        self.button_new_part = Button(text='Novo Artigo', font_size='16', halign='center', background_color=color_light_blue)
        self.button_delete_part = Button(text='Apagar Artigos\nSelecionados', font_size='16', halign='center', background_color=color_light_blue)
        self.label_empty = Label(text='', size_hint_y=4)

        self.button_return.bind(on_press=self.close_popup_window)
        self.button_new_part.bind(on_press=self.display_popup_new_part)
        self.button_delete_part.bind(on_press=self.delete_selected_part)

        self.add_widget(self.button_return)
        self.add_widget(self.button_new_part)
        self.add_widget(self.button_delete_part)
        self.add_widget(self.label_empty)

    def close_popup_window(self, *args):
        production_planning.homepage.layout_sidebar.pop_parts.dismiss()

    def display_popup_new_part(self, *args):
        self.layout_popup_new_part = NewPartpage()
        self.pop_new_part = Popup(title='Novo Artigo', content=self.layout_popup_new_part, size_hint=(0.60, 0.35))

        self.pop_new_part.open()

    def delete_selected_part(self, *args):
        for part in production_planning.homepage.layout_sidebar.layout_popup_parts.layout_sideframe.parts_list:
            if part.is_selected:
                remove_part_database(part.id)
                production_planning.homepage.layout_sidebar.layout_popup_parts.layout_sideframe.update_display()


class Partspage(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.orientation = 'horizontal'

        self.layout_sidebar = Partspage_Sidebar()
        self.layout_sideframe = Partspage_Sideframe()

        self.add_widget(self.layout_sidebar)
        self.add_widget(self.layout_sideframe)


class Material(BoxLayout):
    def __init__(self, ref, material_type, spec, thickness, length, width, density, stock_weight, min_stock, date_modified, client, id, is_selected=False, **kwargs):
        super().__init__(**kwargs)

        # print((ref, material_type, spec, thickness, length, width, density, stock_weight, min_stock, date_modified, client))
        # print('__________________________________________________________________')

        self.ref = str(ref)
        self.type = str(material_type)
        self.spec = str(spec)
        self.stock_weight = str(stock_weight)
        self.stock_num_sheets = str(round(float(stock_weight)/(float(thickness)*float(density)*(float(length)/1000)*(float(width)/1000)), 1))
        self.min_stock = str(min_stock)
        self.thickness = str(thickness)
        self.sheet_length = str(length)
        self.sheet_width = str(width)
        self.density = str(density)
        self.date_modified = str(date_modified)
        self.client = str(client)
        self.id = str(id)
        self.is_selected = is_selected

        # print((self.ref, self.type, self.spec, self.thickness, self.sheet_length, self.sheet_width, self.density, self.stock_weight, self.min_stock, self.date_modified, self.client))

        self.orientation = 'horizontal'
        self.size_hint_y = None
        self.height = 35

        if self.client == 'Prilux':
            self.checkbox_selected = CheckBox(size_hint_x=0.05)
            self.button_ref = Button(text=self.ref, font_size='18', background_color=color_dark_blue, size_hint_x=0.1)
            self.button_spec = Button(text=self.spec, font_size='18', background_color=color_dark_blue, size_hint_x=0.1)
            self.button_thickness = Button(text=self.thickness, font_size='18', background_color=color_dark_blue, size_hint_x=0.05)
            self.button_length = Button(text=self.sheet_length, font_size='18', background_color=color_dark_blue, size_hint_x=0.1)
            self.button_width = Button(text=self.sheet_width, font_size='18', background_color=color_dark_blue, size_hint_x=0.1)
            self.button_stock_weight = Button(text=self.stock_weight, font_size='18', background_color=color_dark_blue, size_hint_x=0.1)
            self.button_stock_num_sheets = Button(text=self.stock_num_sheets, font_size='18', background_color=color_dark_blue, size_hint_x=0.05)
            self.button_min_stock = Button(text=self.min_stock, font_size='18', background_color=color_dark_blue, size_hint_x=0.05)
            self.button_date_modified = Button(text=self.date_modified, font_size='18', background_color=color_dark_blue, size_hint_x=0.1)

            self.add_widget(self.checkbox_selected)
            self.add_widget(self.button_ref)
            self.add_widget(self.button_spec)
            self.add_widget(self.button_thickness)
            self.add_widget(self.button_length)
            self.add_widget(self.button_width)
            self.add_widget(self.button_stock_weight)
            self.add_widget(self.button_stock_num_sheets)
            self.add_widget(self.button_min_stock)
            self.add_widget(self.button_date_modified)

        else:
            self.checkbox_selected = CheckBox(size_hint_x=0.05)
            self.button_client = Button(text=self.client, font_size='17', background_color=color_dark_blue, size_hint_x=0.175)
            self.button_ref = Button(text=self.ref, font_size='17', background_color=color_dark_blue, size_hint_x=0.15)
            self.button_type = Button(text=self.type, font_size='17', background_color=color_dark_blue, size_hint_x=0.075)
            self.button_spec = Button(text=self.spec, font_size='17', background_color=color_dark_blue, size_hint_x=0.1)
            self.button_thickness = Button(text=self.thickness, font_size='17', background_color=color_dark_blue, size_hint_x=0.05)
            self.button_length = Button(text=self.sheet_length, font_size='17', background_color=color_dark_blue, size_hint_x=0.1)
            self.button_width = Button(text=self.sheet_width, font_size='17', background_color=color_dark_blue, size_hint_x=0.1)
            self.button_stock_weight = Button(text=self.stock_weight, font_size='17', background_color=color_dark_blue, size_hint_x=0.075)
            self.button_stock_num_sheets = Button(text=self.stock_num_sheets, font_size='17', background_color=color_dark_blue, size_hint_x=0.05)
            self.button_min_stock = Button(text=self.min_stock, font_size='17', background_color=color_dark_blue, size_hint_x=0.05)
            self.button_date_modified = Button(text=self.date_modified, font_size='17', background_color=color_dark_blue, size_hint_x=0.15)

            self.add_widget(self.checkbox_selected)
            self.add_widget(self.button_client)
            self.add_widget(self.button_ref)
            self.add_widget(self.button_type)
            self.add_widget(self.button_spec)
            self.add_widget(self.button_thickness)
            self.add_widget(self.button_length)
            self.add_widget(self.button_width)
            self.add_widget(self.button_stock_weight)
            self.add_widget(self.button_stock_num_sheets)
            self.add_widget(self.button_min_stock)
            self.add_widget(self.button_date_modified)

        self.checkbox_selected.bind(active=self.on_selector_active)
        self.button_ref.bind(on_press=lambda x: self.display_popup_change_value('ref', self.id))
        self.button_spec.bind(on_press=lambda x: self.display_popup_change_value('spec', self.id))
        self.button_thickness.bind(on_press=lambda x: self.display_popup_change_value('thickness', self.id))
        self.button_width.bind(on_press=lambda x: self.display_popup_change_value('width', self.id))
        self.button_length.bind(on_press=lambda x: self.display_popup_change_value('length', self.id))
        self.button_stock_weight.bind(on_press=lambda x: self.display_popup_change_value('stock_weight', self.id))
        self.button_stock_num_sheets.bind(on_press=lambda x: self.display_popup_change_value('stock_num_sheets', self.id))
        self.button_min_stock.bind(on_press=lambda x: self.display_popup_change_value('min_stock', self.id))

    def on_selector_active(self, *args):
        if self.is_selected:
            self.is_selected = False
        else:
            self.is_selected = True

    def display_popup_change_value(self, field, rowid, *args):

        self.layout_new_value = BoxLayout(orientation='vertical')
        self.layout_horizontal_line = BoxLayout(orientation='horizontal', padding=(10, 10))
        self.label_new_value = Label(text='Novo valor: ', font_size='16', halign='right')
        self.ti_new_value = TextInput()
        self.button_submit = Button(text='Submeter', font_size='18', background_color=color_light_green)

        self.layout_horizontal_line.add_widget(self.label_new_value)
        self.layout_horizontal_line.add_widget(self.ti_new_value)
        self.layout_new_value.add_widget(self.layout_horizontal_line)
        self.layout_new_value.add_widget(self.button_submit)

        self.button_submit.bind(on_press=lambda x: self.change_value(self.ti_new_value.text, field, rowid))

        self.pop_change_value = Popup(title='Novo valor: ', content=self.layout_new_value, size_hint=(0.30, 0.20))
        self.pop_change_value.open()

    def change_value(self, value, field, rowid, *args):
        '''Falta atualizar os outros campos quando um valor de cálculos é alterado'''
        # print((self.stock_weight, self.sheet_length, self.sheet_width))
        if field == 'ref':
            change_value_material_database(value, field, rowid)
            self.ref = str(value)
            self.button_ref.text = self.ref
        elif field == 'type':
            change_value_material_database(value, field, rowid)
            self.type = str(value)
            self.button_type.text = self.type
        elif field == 'spec':
            change_value_material_database(value, field, rowid)
            self.spec = str(value)
            self.button_spec.text = self.spec
        elif field == 'thickness':
            value = float(value)
            change_value_material_database(value, field, rowid)
            self.thickness = str(value)
            self.button_thickness.text = self.thickness
        elif field == 'length':
            value = float(value)
            change_value_material_database(value, field, rowid)
            self.sheet_length = str(value)
            self.button_length.text = self.length
        elif field == 'width':
            value = float(value)
            change_value_material_database(value, field, rowid)
            self.sheet_width = str(value)
            self.button_width.text = self.width
        elif field == 'stock_weight':
            value = float(value)
            change_value_material_database(value, field, rowid)
            self.stock_weight = str(value)
            self.button_stock_weight.text = self.stock_weight
            total_weight = float(self.stock_weight)
            thickness = float(self.thickness)
            density = float(self.density)
            length = float(self.sheet_length)/1000
            width = float(self.sheet_width)/1000
            new_stock_num_sheets = total_weight/(thickness * density * length * width)
            # print(new_stock_num_sheets)
            change_value_material_database(new_stock_num_sheets, 'stock_num_sheets', rowid)
            self.stock_num_sheets = str(new_stock_num_sheets)
            self.button_stock_num_sheets.text = self.stock_num_sheets
        elif field == 'stock_num_sheets':
            value = float(value)
            change_value_material_database(value, field, rowid)
            self.button_stock_num_sheets.text = str(value)
        elif field == 'min_stock':
            value = float(value)
            change_value_material_database(value, field, rowid)
            self.button_min_stock.text = str(value)

        date_modified = f'{datetime.datetime.now().day}/{datetime.datetime.now().month}/{datetime.datetime.now().year}'
        self.button_date_modified.text = date_modified
        change_value_material_database(date_modified, 'date_modified', rowid)
        self.pop_change_value.dismiss()


class Materialspage_Sideframe(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.active_list = ''
        self.materials_list = []

        self.orientation = 'vertical'
        self.padding = 5

        self.button_steel_menu = Button(text='Ferro', font_size='16', background_color=color_orange, size_hint_y=0.15)
        self.button_zinc_menu = Button(text='Zincado', font_size='16', background_color=color_light_blue, size_hint_y=0.15)
        self.button_stainless_steel_menu = Button(text='Inox', font_size='16', background_color=color_light_gray, size_hint_y=0.15)
        self.button_aluminium_menu = Button(text='Alumínio', font_size='16', background_color=color_light_green, size_hint_y=0.15)
        self.button_clients_menu = Button(text='Material Clientes', font_size='16', background_color=color_light_yellow, size_hint_y=0.15)

        self.layout_header = BoxLayout(orientation='vertical', size_hint_y=0.12)
        self.layout_header_non_client_materials = BoxLayout(orientation='horizontal')
        self.layout_header_client_materials = BoxLayout(orientation='horizontal')
        self.label_empty_non_client_materials = Label(text='', size_hint_x=0.054)
        self.label_empty_client_materials = Label(text='', size_hint_x=0.055)
        self.label_ref_non_client_materials = Label(text='Ref.', size_hint_x=0.1, font_size='12', halign='center')
        self.label_ref_client_materials = Label(text='Ref.', size_hint_x=0.15, font_size='12', halign='center')
        self.label_client = Label(text='Cliente', size_hint_x=0.175, font_size='12', halign='center')
        self.label_type_non_client_materials = Label(text='Tipo', size_hint_x=0.075, font_size='12', halign='center')
        self.label_type_client_materials = Label(text='Tipo', size_hint_x=0.075, font_size='12', halign='center')
        self.label_spec_non_client_materials = Label(text='Liga', size_hint_x=0.1, font_size='12', halign='center')
        self.label_spec_client_materials = Label(text='Liga', size_hint_x=0.1, font_size='12', halign='center')
        self.label_thickness_non_client_materials = Label(text='Esp.', size_hint_x=0.05, font_size='12', halign='center')
        self.label_thickness_client_materials = Label(text='Esp.', size_hint_x=0.05, font_size='12', halign='center')
        self.label_length_non_client_materials = Label(text='Comp.', size_hint_x=0.1, font_size='12', halign='center')
        self.label_length_client_materials = Label(text='Comp.', size_hint_x=0.1, font_size='12', halign='center')
        self.label_width_non_client_materials = Label(text='Larg.', size_hint_x=0.1, font_size='12', halign='center')
        self.label_width_client_materials = Label(text='Larg.', size_hint_x=0.1, font_size='12', halign='center')
        self.label_stock_weight_non_client_materials = Label(text='Stock\nPeso', size_hint_x=0.1, font_size='12', halign='center')
        self.label_stock_weight_client_materials = Label(text='Stock\nPeso', size_hint_x=0.075, font_size='12', halign='center')
        self.label_stock_num_sheets_non_client_materials = Label(text='Stock\nNum\nChapas', size_hint_x=0.05, font_size='12', halign='center')
        self.label_stock_num_sheets_client_materials = Label(text='Stock\nNum\nChapas', size_hint_x=0.05, font_size='12', halign='center')
        self.label_min_stock_non_client_materials = Label(text='Stock Min.\nNum.\nChapas', size_hint_x=0.05, font_size='12', halign='center')
        self.label_min_stock_client_materials = Label(text='Stock Min.\nNum.\nChapas', size_hint_x=0.05, font_size='12', halign='center')
        self.label_date_modified_non_client_materials = Label(text='Última\nModificação', size_hint_x=0.1, font_size='12', halign='center')
        self.label_date_modified_client_materials = Label(text='Última\nModificação', size_hint_x=0.15, font_size='12', halign='center')

        self.layout_list = ScrollView(size_hint_y=1)
        self.layout_scroll_steel = BoxLayout(orientation='vertical', size_hint_y=None, padding=5)
        self.layout_scroll_steel.bind(minimum_height=self.layout_scroll_steel.setter('height'))
        self.layout_scroll_zinc = BoxLayout(orientation='vertical', size_hint_y=None, padding=5)
        self.layout_scroll_zinc.bind(minimum_height=self.layout_scroll_zinc.setter('height'))
        self.layout_scroll_stainless_steel = BoxLayout(orientation='vertical', size_hint_y=None, padding=5)
        self.layout_scroll_stainless_steel.bind(minimum_height=self.layout_scroll_stainless_steel.setter('height'))
        self.layout_scroll_aluminium = BoxLayout(orientation='vertical', size_hint_y=None, padding=5)
        self.layout_scroll_aluminium.bind(minimum_height=self.layout_scroll_aluminium.setter('height'))
        self.layout_scroll_clients = BoxLayout(orientation='vertical', size_hint_y=None, padding=5)
        self.layout_scroll_clients.bind(minimum_height=self.layout_scroll_clients.setter('height'))

        self.steel_list_is_open = False
        self.zinc_list_is_open = False
        self.stainless_steel_list_is_open = False
        self.aluminium_list_is_open = False
        self.clients_list_is_open = False

        self.button_steel_menu.bind(on_press=lambda x: self.open_menu('steel'))
        self.button_zinc_menu.bind(on_press=lambda x: self.open_menu('zinc'))
        self.button_stainless_steel_menu.bind(on_press=lambda x: self.open_menu('stainless_steel'))
        self.button_aluminium_menu.bind(on_press=lambda x: self.open_menu('aluminium'))
        self.button_clients_menu.bind(on_press=lambda x: self.open_menu('clients'))

        self.add_widget(self.button_steel_menu)
        self.add_widget(self.button_zinc_menu)
        self.add_widget(self.button_stainless_steel_menu)
        self.add_widget(self.button_aluminium_menu)
        self.add_widget(self.button_clients_menu)
        self.layout_header_non_client_materials.add_widget(self.label_empty_non_client_materials)
        self.layout_header_non_client_materials.add_widget(self.label_ref_non_client_materials)
        self.layout_header_non_client_materials.add_widget(self.label_spec_non_client_materials)
        self.layout_header_non_client_materials.add_widget(self.label_thickness_non_client_materials)
        self.layout_header_non_client_materials.add_widget(self.label_length_non_client_materials)
        self.layout_header_non_client_materials.add_widget(self.label_width_non_client_materials)
        self.layout_header_non_client_materials.add_widget(self.label_stock_weight_non_client_materials)
        self.layout_header_non_client_materials.add_widget(self.label_stock_num_sheets_non_client_materials)
        self.layout_header_non_client_materials.add_widget(self.label_min_stock_non_client_materials)
        self.layout_header_non_client_materials.add_widget(self.label_date_modified_non_client_materials)
        self.layout_header_client_materials.add_widget(self.label_empty_client_materials)
        self.layout_header_client_materials.add_widget(self.label_client)
        self.layout_header_client_materials.add_widget(self.label_ref_client_materials)
        self.layout_header_client_materials.add_widget(self.label_type_client_materials)
        self.layout_header_client_materials.add_widget(self.label_spec_client_materials)
        self.layout_header_client_materials.add_widget(self.label_thickness_client_materials)
        self.layout_header_client_materials.add_widget(self.label_length_client_materials)
        self.layout_header_client_materials.add_widget(self.label_width_client_materials)
        self.layout_header_client_materials.add_widget(self.label_stock_weight_client_materials)
        self.layout_header_client_materials.add_widget(self.label_stock_num_sheets_client_materials)
        self.layout_header_client_materials.add_widget(self.label_min_stock_client_materials)
        self.layout_header_client_materials.add_widget(self.label_date_modified_client_materials)
        self.add_widget(self.layout_header)
        self.add_widget(self.layout_list)

        self.update_display()
        # print(f'Materials list holds {len(self.materials_list)} materials')


    def update_display(self, *args):
        if len(self.materials_list) != 0:
            for material in self.materials_list:
                if material.type == 'Ferro' and material.client == 'Prilux':
                    self.layout_scroll_steel.remove_widget(material)
                elif (material.type == 'Zincado' or material.type == 'Zincor') and material.client == 'Prilux':
                    self.layout_scroll_zinc.remove_widget(material)
                elif material.type == 'Inox' and material.client == 'Prilux':
                    self.layout_scroll_stainless_steel.remove_widget(material)
                elif material.type == 'Alumínio' and material.client == 'Prilux':
                    self.layout_scroll_aluminium.remove_widget(material)
                elif material.client != 'Prilux':
                    self.layout_scroll_clients.remove_widget(material)

            self.materials_list = []

        for material in load_materials_from_database('ferro'):
            m = Material(material[0], material[1], material[2], material[3], material[4], material[5], material[6], material[7], material[9], material[10],
                         material[11], material[-1])
            self.layout_scroll_steel.add_widget(m)
            self.materials_list.append(m)

        for material in load_materials_from_database('zincado'):
            m = Material(material[0], material[1], material[2], material[3], material[4], material[5], material[6], material[7], material[9], material[10],
                         material[11], material[-1])
            self.layout_scroll_zinc.add_widget(m)
            self.materials_list.append(m)

        for material in load_materials_from_database('inox'):
            m = Material(material[0], material[1], material[2], material[3], material[4], material[5], material[6], material[7], material[9], material[10],
                         material[11], material[-1])
            self.layout_scroll_stainless_steel.add_widget(m)
            self.materials_list.append(m)

        for material in load_materials_from_database('alumínio'):
            m = Material(material[0], material[1], material[2], material[3], material[4], material[5], material[6], material[7], material[9], material[10],
                         material[11], material[-1])
            self.layout_scroll_aluminium.add_widget(m)
            self.materials_list.append(m)

        for material in load_materials_from_database('clients'):
            m = Material(material[0], material[1], material[2], material[3], material[4], material[5], material[6], material[7], material[9], material[10],
                         material[11], material[-1])
            self.layout_scroll_clients.add_widget(m)
            self.materials_list.append(m)

    def open_menu(self, name, *args):
        # print(f'Clicked button {name}')
        # print(f'active_list = {self.active_list}')
        if name != self.active_list:
            if self.active_list == 'steel':
                self.layout_list.remove_widget(self.layout_scroll_steel)
                self.layout_header.remove_widget(self.layout_header_non_client_materials)
                # print('Removed layout steel')
            elif self.active_list == 'zinc':
                self.layout_list.remove_widget(self.layout_scroll_zinc)
                self.layout_header.remove_widget(self.layout_header_non_client_materials)
                # print('Removed layout zinc')
            elif self.active_list == 'stainless_steel':
                self.layout_list.remove_widget(self.layout_scroll_stainless_steel)
                self.layout_header.remove_widget(self.layout_header_non_client_materials)
                # print('Removed layout stainles_steel')
            elif self.active_list == 'aluminium':
                self.layout_list.remove_widget(self.layout_scroll_aluminium)
                self.layout_header.remove_widget(self.layout_header_non_client_materials)
                # print('Remvoed layout aluminium')
            elif self.active_list == 'clients':
                self.layout_list.remove_widget(self.layout_scroll_clients)
                self.layout_header.remove_widget(self.layout_header_client_materials)
                # print('Removed layour clients')

            self.active_list = name

            if self.active_list == 'steel':
                self.layout_header.add_widget(self.layout_header_non_client_materials)
                self.layout_list.add_widget(self.layout_scroll_steel)
                # print('Added Layout steel')

            elif self.active_list == 'zinc':
                self.layout_header.add_widget(self.layout_header_non_client_materials)
                self.layout_list.add_widget(self.layout_scroll_zinc)
                # print('Added layout zinc')

            elif self.active_list == 'stainless_steel':
                self.layout_header.add_widget(self.layout_header_non_client_materials)
                self.layout_list.add_widget(self.layout_scroll_stainless_steel)
                # print('Added layout stainless_steel')

            elif self.active_list == 'aluminium':
                self.layout_header.add_widget(self.layout_header_non_client_materials)
                self.layout_list.add_widget(self.layout_scroll_aluminium)
                # print('Added layout aluminium')

            elif self.active_list == 'clients':
                self.layout_header.add_widget(self.layout_header_client_materials)
                self.layout_list.add_widget(self.layout_scroll_clients)
                # print('Added layout clients')

            self.change_button_color(name)

    def change_button_color(self, name, *args):
        if name == 'steel':
            self.button_steel_menu.background_color = color_dark_green
            self.button_zinc_menu.background_color = color_light_blue
            self.button_stainless_steel_menu.background_color = color_light_gray
            self.button_aluminium_menu.background_color = color_light_green
            self.button_clients_menu.background_color = color_light_yellow
        elif name == 'zinc':
            self.button_steel_menu.background_color = color_orange
            self.button_zinc_menu.background_color = color_dark_green
            self.button_stainless_steel_menu.background_color = color_light_gray
            self.button_aluminium_menu.background_color = color_light_green
            self.button_clients_menu.background_color = color_light_yellow
        elif name == 'stainless_steel':
            self.button_steel_menu.background_color = color_orange
            self.button_zinc_menu.background_color = color_light_blue
            self.button_stainless_steel_menu.background_color = color_dark_green
            self.button_aluminium_menu.background_color = color_light_green
            self.button_clients_menu.background_color = color_light_yellow
        elif name == 'aluminium':
            self.button_steel_menu.background_color = color_orange
            self.button_zinc_menu.background_color = color_light_blue
            self.button_stainless_steel_menu.background_color = color_light_gray
            self.button_aluminium_menu.background_color = color_dark_green
            self.button_clients_menu.background_color = color_light_yellow
        elif name == 'clients':
            self.button_steel_menu.background_color = color_orange
            self.button_zinc_menu.background_color = color_light_blue
            self.button_stainless_steel_menu.background_color = color_light_gray
            self.button_aluminium_menu.background_color = color_light_green
            self.button_clients_menu.background_color = color_dark_green


class NewMaterialpage(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.is_client = False

        self.orientation = 'vertical'

        self.layout_line_1 = BoxLayout(orientation='horizontal', padding=(5, 5))
        self.layout_line_2 = BoxLayout(orientation='horizontal', padding=(5, 5))
        self.layout_line_3 = BoxLayout(orientation='horizontal', padding=(5, 5))
        self.layout_line_4 = BoxLayout(orientation='horizontal', padding=(5, 5))

        self.label_ref = Label(text='Referência: ', font_size='16', halign='right', size_hint_y=0.75)
        self.ti_ref = TextInput(size_hint_y=0.75)
        self.label_type = Label(text='Tipo: ', font_size='16', halign='right', size_hint_y=0.75)
        self.drp_type = DropDown()
        self.button_type = Button(text='', size_hint_y=0.75, background_color=color_dark_blue)
        self.button_type.bind(on_release=self.drp_type.open)
        self.drp_type.bind(on_select=lambda instance, new_val: setattr(self.button_type, 'text', new_val))

        # Buttons for dropdown list
        self.button_steel = Button(text='Ferro', font_size='16', size_hint_y=None, height=30, background_color=color_dark_blue)
        self.button_zinc = Button(text='Zincado', font_size='16', size_hint_y=None, height=30, background_color=color_dark_blue)
        self.button_stainless_steel = Button(text='Inox', font_size='16', size_hint_y=None, height=30, background_color=color_dark_blue)
        self.button_aluminium = Button(text='Alumínio', font_size='16', size_hint_y=None, height=30, background_color=color_dark_blue)

        self.button_steel.bind(on_release=lambda type_object: self.drp_type.select(type_object.text))
        self.button_zinc.bind(on_release=lambda type_object: self.drp_type.select(type_object.text))
        self.button_stainless_steel.bind(on_release=lambda type_object: self.drp_type.select(type_object.text))
        self.button_aluminium.bind(on_release=lambda type_object: self.drp_type.select(type_object.text))

        self.drp_type.add_widget(self.button_steel)
        self.drp_type.add_widget(self.button_zinc)
        self.drp_type.add_widget(self.button_stainless_steel)
        self.drp_type.add_widget(self.button_aluminium)
        # End of dropdown section

        self.label_spec = Label(text='Especif.: ', font_size='16', halign='right', size_hint_y=0.75)
        self.ti_spec = TextInput(size_hint_y=0.75)
        self.label_stock_weight = Label(text='Stock Kg: ', font_size='16', halign='right', size_hint_y=0.75)
        self.ti_stock_weight = TextInput(size_hint_y=0.75)
        self.label_density = Label(text='Densidade: ', font_size='16', halign='right', size_hint_y=0.75)
        self.ti_density = TextInput(size_hint_y=0.75, size_hint_x=0.2)
        self.label_thickness = Label(text='Espessura: ', font_size='16', halign='right', size_hint_y=0.75)
        self.ti_thickness = TextInput(size_hint_y=0.75, size_hint_x=0.2)
        self.label_length = Label(text='Comprimento: ', font_size='16', halign='right', size_hint_y=0.75)
        self.ti_length = TextInput(size_hint_y=0.75, size_hint_x=0.7)
        self.label_width = Label(text='Largura: ', font_size='16', halign='right', size_hint_y=0.75, size_hint_x=0.7)
        self.ti_width = TextInput(size_hint_y=0.75, size_hint_x=0.7)
        self.label_min_stock = Label(text='Stock Mínimo (chapas): ', font_size='16', halign='right', size_hint_y=0.75)
        self.ti_min_stock = TextInput(size_hint_y=0.75, size_hint_x=0.2)
        self.label_is_client = Label(text='M.P. Cliente?: ', font_size='16', halign='right', size_hint_y=0.75)
        self.check_is_client = CheckBox()
        self.label_empty = Label(text='', size_hint_y=0.25)
        self.button_submit = Button(text='Submeter Novo Material', font_size='16', halign='center', background_color=color_dark_blue, padding=(10, 10))

        self.check_is_client.bind(active=self.select_is_client)
        self.button_submit.bind(on_press=self.insert_new_material)

        self.add_widget(self.layout_line_1)
        self.add_widget(self.layout_line_2)
        self.add_widget(self.layout_line_3)
        self.add_widget(self.layout_line_4)

        self.layout_line_1.add_widget(self.label_ref)
        self.layout_line_1.add_widget(self.ti_ref)
        self.layout_line_1.add_widget(self.label_type)
        self.layout_line_1.add_widget(self.button_type)
        self.layout_line_1.add_widget(self.label_spec)
        self.layout_line_1.add_widget(self.ti_spec)
        self.layout_line_2.add_widget(self.label_stock_weight)
        self.layout_line_2.add_widget(self.ti_stock_weight)
        self.layout_line_2.add_widget(self.label_density)
        self.layout_line_2.add_widget(self.ti_density)
        self.layout_line_2.add_widget(self.label_thickness)
        self.layout_line_2.add_widget(self.ti_thickness)
        self.layout_line_3.add_widget(self.label_length)
        self.layout_line_3.add_widget(self.ti_length)
        self.layout_line_3.add_widget(self.label_width)
        self.layout_line_3.add_widget(self.ti_width)
        self.layout_line_3.add_widget(self.label_min_stock)
        self.layout_line_3.add_widget(self.ti_min_stock)
        self.layout_line_4.add_widget(self.label_is_client)
        self.layout_line_4.add_widget(self.check_is_client)
        self.add_widget(self.label_empty)
        self.add_widget(self.button_submit)

    def insert_new_material(self, *args):
        ref = self.ti_ref.text
        type = self.button_type.text
        spec = self.ti_spec.text
        thickness = float(self.ti_thickness.text)
        length = float(self.ti_length.text)
        width = float(self.ti_width.text)
        density = float(self.ti_density.text)
        stock_weight = float(self.ti_stock_weight.text)
        stock_num_sheets = math.floor(stock_weight/(length/1000*width/1000*density*thickness))
        min_stock = int(self.ti_min_stock.text)
        date_modified = f'{datetime.datetime.now().day}/{datetime.datetime.now().month}/{datetime.datetime.now().year}'
        client = 'Prilux'
        if self.is_client:
            client = self.ti_client.text

        # print((ref, type, spec, thickness, length, width, density, stock_weight, stock_num_sheets, min_stock, date_modified, client))
        insert_new_material_database(ref, type, spec, thickness, length, width, density, stock_weight, stock_num_sheets, min_stock, date_modified, client)
        production_planning.homepage.layout_sidebar.layout_popup_materials.layout_sideframe.update_display()
        production_planning.homepage.layout_sidebar.layout_popup_materials.layout_sidebar.pop_new_material.dismiss()

    def select_is_client(self, *args):
        self.is_client = not self.is_client

        if self.is_client:
            self.label_client = Label(text='Cliente: ', font_size='16', halign='right', size_hint_y=0.75)
            self.ti_client = TextInput(size_hint_y=0.75)
            self.layout_line_4.add_widget(self.label_client)
            self.layout_line_4.add_widget(self.ti_client)
        else:
            self.layout_line_4.remove_widget(self.label_client)
            self.layout_line_4.remove_widget(self.ti_client)


class Materialspage_Sidebar(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.orientation = 'vertical'
        self.size_hint_x = 0.08

        self.button_return = Button(text='<<<', font_size='16', halign='center', background_color=color_light_blue)
        self.button_new_material = Button(text='Novo Material', font_size='14', halign='center', background_color=color_light_blue)
        self.button_delete_material = Button(text='Apagar Materiais\nSelecionados', font_size='14', halign='center', background_color=color_light_blue)
        # self.button_material_machine_dependencies = Button(text='Dependências\nMaterial/Máquina', font_size='14', halign='center', background_color=color_light_blue)
        # self.button_consume_history = Button(text='Histórico de\nConsumos', font_size='14', halign='center', background_color=color_light_blue)
        # self.button_order_recommendations = Button(text='Recomendações de\nEncomendas', font_size='14', halign='center', background_color=color_light_blue)
        self.label_empty1 = Label(size_hint_y=2)
        self.label_empty2 = Label(size_hint_y=2)
        self.label_empty3 = Label(size_hint_y=2)

        self.button_return.bind(on_press=self.close_popup_window)
        self.button_new_material.bind(on_press=self.display_popup_new_material)
        self.button_delete_material.bind(on_press=self.delete_selected_materials)

        self.add_widget(self.button_return)
        self.add_widget(self.button_new_material)
        self.add_widget(self.button_delete_material)
        # self.add_widget(self.button_material_machine_dependencies)
        # self.add_widget(self.button_consume_history)
        # self.add_widget(self.button_order_recommendations)
        self.add_widget(self.label_empty1)
        self.add_widget(self.label_empty2)
        self.add_widget(self.label_empty3)

    def close_popup_window(self, *args):
        production_planning.homepage.layout_sidebar.pop_materials.dismiss()

    def display_popup_new_material(self, *args):
        self.layout_popup_new_material = NewMaterialpage()
        self.pop_new_material = Popup(title='Novo Material', content=self.layout_popup_new_material, size_hint=(0.60, 0.40))

        self.pop_new_material.open()

    def delete_selected_materials(self, *args):
        for material in production_planning.homepage.layout_sidebar.layout_popup_materials.layout_sideframe.materials_list:
            if material.is_selected:
                remove_material_database(material.id)
                production_planning.homepage.layout_sidebar.layout_popup_materials.layout_sideframe.update_display()


class Materialspage(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.orientation = 'horizontal'

        self.layout_sidebar = Materialspage_Sidebar()
        self.layout_sideframe = Materialspage_Sideframe()

        self.add_widget(self.layout_sidebar)
        self.add_widget(self.layout_sideframe)


# class Popup_ManualTask(BoxLayout):
#     def __init__(self, **kwargs):
#         super().__init__(**kwargs)
#
#         self.orientation = 'horizontal'
#         self.size_hint_y = None
#         self.height = 30
#
#         self.button_delete = Button(text='X', background_color=color_dark_red, size_hint_x=0.05)
#         self.label_part_ref = Label(text='Ref. Peça: ', font_size='18', font_name='Calibri', size_hint_x=0.15)
#         self.ti_part_ref = TextInput(size_hint_x=0.17)
#         self.label_part_name = Label(text='Designação Peça: ', font_size='18', font_name='Calibri', size_hint_x=0.22)
#         self.ti_part_name = TextInput(size_hint_x=0.4)
#         self.label_part_qtty = Label(text='Qtd. Peça: ', font_size='18', font_name='Calibri', size_hint_x=0.15)
#         self.ti_part_qtty = TextInput(size_hint_x=0.15)
#         self.drp_unit = DropDown()
#         self.button_unit = Button(text='Unid.', font_size='18', font_name='Calibri', size_hint_x=0.1)
#         self.button_qtty = Button(text='UN', font_size='18', font_name='Calibri', size_hint_y=None, height=30, background_color=color_light_green)
#         self.button_meters = Button(text='MT', font_size='18', font_name='Calibri', size_hint_y=None, height=30, background_color=color_light_green)
#         self.button_kgs = Button(text='KG', font_size='18', font_name='Calibri', size_hint_y=None, height=30, background_color=color_light_green)
#
#
#         self.button_delete.bind(on_press=self.delete_self)
#         self.button_unit.bind(on_release=self.drp_unit.open)
#         self.button_qtty.bind(on_release=lambda x: self.drp_unit.select(self.button_qtty.text))
#         self.button_meters.bind(on_release=lambda x: self.drp_unit.select(self.button_meters.text))
#         self.button_kgs.bind(on_release=lambda x: self.drp_unit.select(self.button_kgs.text))
#         self.drp_unit.bind(on_select=lambda instance, x: setattr(self.button_unit, 'text', x))
#
#         self.add_widget(self.button_delete)
#         self.add_widget(self.label_part_ref)
#         self.add_widget(self.ti_part_ref)
#         self.add_widget(self.label_part_name)
#         self.add_widget(self.ti_part_name)
#         self.add_widget(self.label_part_qtty)
#         self.add_widget(self.ti_part_qtty)
#         self.add_widget(self.button_unit)
#         self.drp_unit.add_widget(self.button_qtty)
#         self.drp_unit.add_widget(self.button_meters)
#         self.drp_unit.add_widget(self.button_kgs)
#
#     def delete_self(self, *args):
#         pass
#         # for manual_task in order_viewer.home_page.layout_sidebar.layout_manual_parts_list.manual_task_list:
#         #     if self == manual_task:
#         #         order_viewer.home_page.layout_sidebar.layout_manual_parts_list.delete_manual_task(self)


# class Popup_ManualTaskList(ScrollView):
#     def __init__(self, **kwargs):
#         super().__init__(**kwargs)
#
#         self.manual_task_list = []
#
#         self.layout_scroll = BoxLayout(orientation='vertical', padding=10, spacing=(5, 5), size=(10, 10), size_hint_y=None)
#         self.layout_scroll.bind(minimum_height=self.layout_scroll.setter('height'))
#
#         self.add_widget(self.layout_scroll)
#
#
#     def generate_manual_tasks_list_display(self):
#
#         self.manual_task_list.append(Popup_ManualTask())
#
#         for manual_task in self.manual_task_list:
#             self.layout_scroll.add_widget(manual_task)
#
#
#     def create_new_manual_task(self, *args):
#         new_manual_task = Popup_ManualTask()
#
#         self.manual_task_list.append(new_manual_task)
#         self.layout_scroll.add_widget(new_manual_task)
#
#
#     def delete_manual_task(self, manual_task_to_delete, *args):
#         self.layout_scroll.remove_widget(manual_task_to_delete)
#         del self.manual_task_list[self.manual_task_list.index(manual_task_to_delete)]
#
#
#     def submit_manual_tasks(self, manual_task_client, manual_task_order):
#         print(f'Manual Task Client: {manual_task_client}')
#         print(f'Manual Task Order: {manual_task_order}')
#         manual_task_id = 1
#         manual_order_path = ''
#         manual_part_produced_qtty = '0'
#
#         # for manual_task in self.manual_task_list:
#         #     new_manual_task = HomePage_ManualTask(manual_order_path, manual_task_client, manual_task_order, manual_task.ti_part_ref.text, manual_task.ti_part_name.text,
#         #                                     manual_task.ti_part_qtty.text, manual_task.button_unit.text, manual_part_produced_qtty, manual_task_id)
#         #
#         #     order_viewer.task_list.append(new_manual_task)
#         #     manual_task_id += 1
#         #
#         # order_viewer.home_page.layout_sideframe.layout_task_list.update_task_list_display()


class Taskpage_Filter(CheckBox):
    def __init__(self, name, state='normal', is_selected=False, **kwargs):
        super().__init__(**kwargs)

        self.name = name
        self.state = state
        self.is_selected = is_selected

        self.bind(active=self.on_filter_pressed)

    def on_filter_pressed(self, *args):

        if self.name == 'Machine':
            production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.update_display(sort_by='machine')
        elif self.name == 'Material':
            production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.update_display(sort_by='material')
        elif self.name == 'Start Date':
            production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.update_display(sort_by='start_date')
        elif self.name == 'End Date':
            production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.update_display(sort_by='end_date')
        elif self.name == 'Priority':
            production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.update_display(sort_by='priority')

        production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.layout_filters_bar.button_expand_collapse_all.text = '+\n+'
        production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.layout_filters_bar.expand_all = True


class Taskpage_CheckboxSelector(CheckBox):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.state = 'normal'
        self.is_selected = False

        self.bind(on_press=self.on_selector_active)

    def on_selector_active(self, *args):
        if self.is_selected:
            self.is_selected = False
        else:
            self.is_selected = True

        # order_viewer.create_task_list(is_selected=self.is_selected)
        # order_viewer.home_page.layout_sideframe.layout_task_list.update_task_list_display()


class Taskpage_Filters_Bar(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.orientation = 'horizontal'
        self.size_hint_y = 0.1
        self.padding = (12, 12)

        self.expand_all = True

        self.button_expand_collapse_all = Button(text='+\n+', font_size='20', background_color=color_light_blue, size_hint_x=0.05)

        self.button_expand_collapse_all.bind(on_press=self.expand_collapse_all)

        # Select All Checkbox Block
        self.layout_checkbox_select_all = BoxLayout(orientation='vertical', size_hint_x=0.03)

        self.label_checkbox_select_all = Label(text='', font_size='15', halign='center', valign='center')

        self.checkbox_select_all = Taskpage_CheckboxSelector()

        self.checkbox_select_all.bind(on_press=self.select_all_tasks)

        self.layout_checkbox_select_all.add_widget(self.label_checkbox_select_all)
        self.layout_checkbox_select_all.add_widget(self.checkbox_select_all)

        # Empty Labels for spacing
        self.label_empty_1 = Label(size_hint_x=0.1)
        self.label_empty_2 = Label(size_hint_x=0.1)

        # Machine Filter Block
        self.layout_filter_machine = BoxLayout(orientation='vertical', size_hint_x=0.05)

        self.label_filter_machine = Label(text='Máq.', font_size='15', halign='center', valign='center')

        self.filter_machine = Taskpage_Filter('Machine', group=1)

        self.layout_filter_machine.add_widget(self.label_filter_machine)
        self.layout_filter_machine.add_widget(self.filter_machine)

        # Material Filter Block
        self.layout_filter_material = BoxLayout(orientation='vertical', size_hint_x=0.3)

        self.label_filter_material = Label(text='Material', font_size='15', halign='center', valign='center')

        self.filter_material = Taskpage_Filter('Material', group=1)

        self.layout_filter_material.add_widget(self.label_filter_material)
        self.layout_filter_material.add_widget(self.filter_material)

        # Empty Labels for spacing
        self.label_empty_3 = Label(size_hint_x=0.08)
        self.label_empty_4 = Label(size_hint_x=0.07)

        # Start Date Filter Block
        self.layout_filter_start_date = BoxLayout(orientation='vertical', size_hint_x=0.1)

        self.label_filter_start_date = Label(text='Data Início', font_size='15', halign='center', valign='center')

        self.filter_start_date = Taskpage_Filter('Start Date', group=1)

        self.layout_filter_start_date.add_widget(self.label_filter_start_date)
        self.layout_filter_start_date.add_widget(self.filter_start_date)

        # End Date Filter Block
        self.layout_filter_end_date = BoxLayout(orientation='vertical', size_hint_x=0.1)

        self.label_filter_end_date = Label(text='Data Fim', font_size='15', halign='center', valign='center')

        self.filter_end_date = Taskpage_Filter('End Date', group=1)

        self.layout_filter_end_date.add_widget(self.label_filter_end_date)
        self.layout_filter_end_date.add_widget(self.filter_end_date)

        # Priority Filter Block
        self.layout_filter_priority = BoxLayout(orientation='vertical', size_hint_x=0.05)

        self.label_filter_priority = Label(text='Ordem', font_size='15', halign='center', valign='center') #Anteriormente era a prioridade

        self.filter_priority = Taskpage_Filter('Priority', group=1)

        self.layout_filter_priority.add_widget(self.label_filter_priority)
        self.layout_filter_priority.add_widget(self.filter_priority)

        # Empty Label for spacing
        self.label_empty_5 = Label(size_hint_x=0.05)

        # Add all the combinations of Label + Filter to the Filter bar
        self.add_widget(self.button_expand_collapse_all)
        self.add_widget(self.layout_checkbox_select_all)
        self.add_widget(self.label_empty_1)
        self.add_widget(self.label_empty_2)
        self.add_widget(self.layout_filter_machine)
        self.add_widget(self.layout_filter_material)
        self.add_widget(self.label_empty_3)
        self.add_widget(self.label_empty_4)
        self.add_widget(self.layout_filter_start_date)
        self.add_widget(self.layout_filter_end_date)
        self.add_widget(self.layout_filter_priority)
        self.add_widget(self.label_empty_5)

    def expand_collapse_all(self, *args):
        if self.expand_all:
            self.button_expand_collapse_all.text = '-\n-'
            for added_task in production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.added_tasks:
                if not added_task.is_expanded:
                    added_task.expand()
        else:
            self.button_expand_collapse_all.text = '+\n+'
            for added_task in production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.added_tasks:
                if added_task.is_expanded:
                    added_task.expand()

        self.expand_all = not self.expand_all

    def select_all_tasks(self, *args):
        if self.checkbox_select_all.state == 'down':
            for added_task in production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.added_tasks:
                added_task.checkbox_selected.state = 'down'
                added_task.is_selected = True
        else:
            for added_task in production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.added_tasks:
                added_task.checkbox_selected.state = 'normal'
                added_task.is_selected = False


class Taskpage_Sideframe(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.orientation = 'vertical'

        self.added_tasks = []
        self.aggregated_tasks_dict = return_dictionary_of_two_letter_combinations()

        self.layout_filters_bar = Taskpage_Filters_Bar()
        self.layout_task_list = Homepage_Task_List()

        self.add_widget(self.layout_filters_bar)
        self.add_widget(self.layout_task_list)

        self.update_display()
        self.update_aggregated_tasks_dict()

    def update_display(self, sort_by='priority', *args):
        if len(self.added_tasks) != 0:
            for added_task in self.added_tasks:
                self.layout_task_list.layout_scroll.remove_widget(added_task)

            self.added_tasks = []
            sorted_tasks = []

        for task in load_tasks_from_database():
            # print((task[0], task[1], task[2], task[3], task[4], task[5], task[6], task[7], task[8], task[9], task[10],task[11], task[12], task[-1]))
            t = Task(task[0], task[1], task[2], task[3], task[4], task[5], task[6], task[7], task[8], task[9], task[10], task[11], task[-1], aggregated_index=task[12], checkbox_mode=True)
            self.added_tasks.append(t)

        if sort_by == 'priority':
            self.added_tasks = sorted(self.added_tasks, key=operator.attrgetter('priority'), reverse=True)
        elif sort_by == 'machine':
            self.added_tasks = sorted(self.added_tasks, key=operator.attrgetter('machine'))
        elif sort_by == 'material':
            self.added_tasks = sorted(self.added_tasks, key=operator.attrgetter('material'))
        elif sort_by == 'start_date':
            self.added_tasks = sorted(self.added_tasks, key=operator.attrgetter('start_date'))
        elif sort_by == 'end_date':
            self.added_tasks = sorted(self.added_tasks, key=operator.attrgetter('end_date'))

        for task in self.added_tasks:
            self.layout_task_list.layout_scroll.add_widget(task)

    def update_aggregated_tasks_dict(self):
        rowid_patterns_added_to_dict = []

        for added_task in self.added_tasks:
            # print(f'ADDED TASK AGGREGATED TASKS: {added_task.aggregated_tasks}')
            if len(added_task.aggregated_tasks) > 1:
                rowid_string = ','.join(str(rowid_aggregated_task) for rowid_aggregated_task in sorted(added_task.aggregated_tasks))

                for key in self.aggregated_tasks_dict:
                    if self.aggregated_tasks_dict[key] == '' and rowid_string not in self.aggregated_tasks_dict.values():
                        self.aggregated_tasks_dict[key] = rowid_string
                        rowid_patterns_added_to_dict.append(rowid_string)

        # print(self.aggregated_tasks_dict)

    # def attribute_aggregated_tasks_keys(self, *args):
    #     for added_task in self.added_tasks:
    #         added_task.set_index_aggregated_group()
    #         added_task.set_text_buttons_paths()


class NewTaskPart(BoxLayout):
    def __init__(self, name, quantity, produced_quantity, client, order_num, order_num_client, additional_operations, rowid, due_date='N/A', **kwargs):
        super().__init__(**kwargs)

        self.orientation = 'horizontal'
        self.size_hint_y = None
        self.height = 35

        self.name = name
        self.quantity = quantity
        self.produced_quantity = produced_quantity
        self.pending_quantity = str(self.quantity - self.produced_quantity)
        if self.produced_quantity > self.quantity:
            self.pending_quantity = str(0)
        self.client = client
        self.order_num = order_num
        self.order_num_client = order_num_client
        if additional_operations:
            self.additional_operations = 'Sim'
        else:
            self.additional_operations = 'Não'
        self.is_selected = False
        self.rowid = rowid
        self.due_date = due_date

        self.checkbox_selected = CheckBox(size_hint_x=0.05)
        self.button_name = Button(text=self.name[:66], font_size='16', background_color=color_dark_blue, size_hint_x=0.4)
        self.button_pending_quantity = Button(text=self.pending_quantity, font_size='16', background_color=color_dark_blue, size_hint_x=0.07)
        self.button_client = Button(text=self.client[:26], font_size='16', background_color=color_dark_blue, size_hint_x=0.18)
        self.button_orders = Button(text=f'{self.order_num_client} / {self.order_num}', font_size='16', background_color=color_dark_blue, size_hint_x=0.25)
        self.button_additional_operations = Button(text=self.additional_operations, font_size='16', background_color=color_dark_blue, size_hint_x=0.1)

        self.checkbox_selected.bind(active=self.set_selected)

        self.add_widget(self.checkbox_selected)
        self.add_widget(self.button_name)
        self.add_widget(self.button_pending_quantity)
        self.add_widget(self.button_client)
        self.add_widget(self.button_orders)
        self.add_widget(self.button_additional_operations)

    def set_selected(self, *args):
        self.is_selected = not self.is_selected


class NewTaskpage(BoxLayout):
        def __init__(self, **kwargs):
            super().__init__(**kwargs)
            #TODO Add "notes" TextInput
            self.orientation = 'vertical'

            self.order_parts = load_order_parts_from_database()
            self.tasks = load_tasks_from_database()
            self.added_new_order_parts = []
            self.select_all = False

            self.layout_line_1 = BoxLayout(orientation='horizontal', size_hint_y=0.05, padding=(5, 5))
            self.layout_line_2 = BoxLayout(orientation='horizontal', size_hint_y=0.05, padding=(5, 5))
            self.layout_line_3 = BoxLayout(orientation='horizontal', size_hint_y=0.05)
            self.layout_line_4 = BoxLayout(orientation='horizontal', size_hint_y=0.75, padding=(5, 5))
            self.layout_new_tasks_list = ScrollView()
            self.layout_scroll_new_task = BoxLayout(orientation='vertical', size_hint_y=None, padding=5)
            self.layout_scroll_new_task.bind(minimum_height=self.layout_scroll_new_task.setter('height'))
            self.layout_line_5 = BoxLayout(orientation='horizontal', size_hint_y=0.05)
            self.label_path = Label(text='Caminho: ', size_hint_x=0.1)
            self.ti_path = TextInput(size_hint_x=0.9)
            self.drp_material = DropDown()
            self.button_material = Button(text='Material', font_size='16', halign='center', size_hint_x=0.3, background_color=color_light_blue, padding=(5, 5))
            self.button_material.bind(on_release=self.drp_material.open)
            self.drp_material.bind(on_select=lambda instance, new_val: self.select_material(new_val))
            for material in load_materials_from_database():
                material_type = material[1]
                material_spec = str(material[2])
                material_thickness = str(material[3])
                material_client = material[11]

                material_name = material_type + ' ' + material_spec + ', ' + material_thickness + 'mm ' + material_client

                button_material_name = Button(text=material_name, font_size='16', size_hint_y=None, height=30, background_color=color_dark_blue)
                button_material_name.bind(on_release=lambda material_object: self.drp_material.select(material_object.text))

                self.drp_material.add_widget(button_material_name)

            self.label_estimated_sheets_required = Label(text='Num. estimado de chapas', size_hint_x=0.2)
            self.ti_estimated_sheets_required = TextInput(size_hint_x=0.1)
            self.label_estimated_time = Label(text='Duração Estimada (min): ', size_hint_x=0.2)
            self.ti_estimated_time = TextInput(size_hint_x=0.1)
            self.drp_machine = DropDown()
            self.button_machine = Button(text='Máquina', size_hint_x=0.1, background_color=color_light_blue, padding=(5, 5))
            self.button_machine.bind(on_release=self.drp_machine.open)
            self.drp_machine.bind(on_select=lambda instance, new_val: self.select_machine(new_val))
            for machine_name in ['LF-3015', 'LC5']:
                button_machine_name = Button(text=machine_name, font_size='16', size_hint_y=None, height=30, background_color=color_dark_blue)
                button_machine_name.bind(on_release=lambda machine_object: self.drp_machine.select(machine_object.text))

                self.drp_machine.add_widget(button_machine_name)

            self.check_select_all = CheckBox(size_hint_x=0.065)
            self.label_name = Label(text='Nome', halign='center', size_hint_x=0.4)
            self.label_pending_quantity = Label(text='Qtd.\nPendente', halign='center', size_hint_x=0.07)
            self.label_client = Label(text='Cliente', halign='center', size_hint_x=0.18)
            self.label_orders = Label(text='Enc. Cliente / Enc. Interna', halign='center', size_hint_x=0.25)
            self.label_additional_operations = Label(text='Op. Adicionais', halign='center', size_hint_x=0.1)
            self.button_submit = Button(text='Submeter', halign='center', background_color=color_light_blue)

            self.check_select_all.bind(on_press=self.select_all_task_order_parts)
            self.button_submit.bind(on_press=self.submit_new_task)

            self.layout_line_1.add_widget(self.label_path)
            self.layout_line_1.add_widget(self.ti_path)
            self.layout_line_2.add_widget(self.button_material)
            self.layout_line_2.add_widget(self.label_estimated_sheets_required)
            self.layout_line_2.add_widget(self.ti_estimated_sheets_required)
            self.layout_line_2.add_widget(self.label_estimated_time)
            self.layout_line_2.add_widget(self.ti_estimated_time)
            self.layout_line_2.add_widget(self.button_machine)
            self.layout_line_3.add_widget(self.check_select_all)
            self.layout_line_3.add_widget(self.label_name)
            self.layout_line_3.add_widget(self.label_pending_quantity)
            self.layout_line_3.add_widget(self.label_client)
            self.layout_line_3.add_widget(self.label_orders)
            self.layout_line_3.add_widget(self.label_additional_operations)
            self.layout_new_tasks_list.add_widget(self.layout_scroll_new_task)
            self.layout_line_4.add_widget(self.layout_new_tasks_list)
            self.layout_line_5.add_widget(self.button_submit)
            self.add_widget(self.layout_line_1)
            self.add_widget(self.layout_line_2)
            self.add_widget(self.layout_line_3)
            self.add_widget(self.layout_line_4)
            self.add_widget(self.layout_line_5)

        def select_material(self, new_val, *args):
            setattr(self.button_material, 'text', new_val)
            material_type = self.button_material.text.split(' ')[0]
            material_spec = self.button_material.text.split(' ')[1].split(',')[0]
            material_thickness = float(self.button_material.text.split(' ')[2].split('mm')[0])
            material_client = self.button_material.text.split(' ')[3]
            material_ref = get_material_ref(material_type, material_spec, material_thickness, material_client)

            if material_thickness > 3.0:
                self.button_machine.text = 'LC5'
            else:
                self.button_machine.text = 'LF-3015'

            self.update_display(material_ref)
            self.select_all_task_order_parts()

        def update_display(self, material_ref, *args):
            if len(self.added_new_order_parts) != 0:
                for added_new_order_part in self.added_new_order_parts:
                    self.layout_scroll_new_task.remove_widget(added_new_order_part)

            self.added_new_order_parts = []

            for new_order_part in self.order_parts:
                if new_order_part[2] == material_ref:
                    order_part_already_added = self.check_order_part_in_any_task(new_order_part[-1])
                    if not order_part_already_added:
                        ntp = NewTaskPart(new_order_part[1], new_order_part[3], new_order_part[4], new_order_part[7], new_order_part[5], new_order_part[6], new_order_part[10], new_order_part[-1], due_date=new_order_part[9])
                        if int(ntp.pending_quantity) != 0 and int(ntp.pending_quantity) > 0:
                            self.layout_scroll_new_task.add_widget(ntp)
                            self.added_new_order_parts.append(ntp)

        def check_order_part_in_any_task(self, rowid):
            '''Checks if the rowid passed in is already added to one of the database tasks'''
            for task in self.tasks:
                for task_order_part_rowid in task[10].split(','):
                    if int(task_order_part_rowid) == rowid:
                        return True
            return False

        def select_machine(self, new_val, *args):
            previous_val = self.button_machine.text
            setattr(self.button_machine, 'text', new_val)
            # if new_val == 'LC5' and previous_val == 'LF-3015':
            #     self.ti_estimated_time.text = str(round(int(self.ti_estimated_time.text) / 2))
            # elif new_val == 'LF-3015' and previous_val == 'LC5':
            #     self.ti_estimated_time.text = str(round(int(self.ti_estimated_time.text) * 2))

        def select_all_task_order_parts(self, *args):

            self.select_all = not self.select_all

            if self.select_all:
                state = 'down'
            else:
                state = 'normal'

            # print((f'STATE IS {state}, SELECT ALL IS {self.select_all}'))

            self.check_select_all.state = state
            for part in self.added_new_order_parts:
                part.checkbox_selected.state = state #This already toggles the is_selected state

        def generate_task_order_parts_rowid_string(self):
            order_parts_string = ''
            for new_order_part in self.added_new_order_parts:
                if new_order_part.is_selected:
                    if order_parts_string == '':
                        order_parts_string = str(new_order_part.rowid)
                    else:
                        order_parts_string += f',{str(new_order_part.rowid)}'

            return order_parts_string

        def return_selected_order_parts(self):
            selected_order_parts = []
            for order_part in self.added_new_order_parts:
                if order_part.is_selected:
                    selected_order_parts.append(order_part)

            return selected_order_parts

        def submit_new_task(self, *args):

            selected_order_parts = self.return_selected_order_parts()

            try:
                if len(selected_order_parts) != 0:
                    original_path = self.ti_path.text
                    current_path = original_path
                    machine = self.button_machine.text
                    material_type = self.button_material.text.split(' ')[0]
                    material_spec = self.button_material.text.split(' ')[1].split(',')[0]
                    material_thickness = float(self.button_material.text.split(' ')[2].split('mm')[0])
                    material_client = self.button_material.text.split(' ')[3]
                    material_ref = get_material_ref(material_type, material_spec, material_thickness, material_client)
                    notes = ''
                    estimated_sheets_required = float(self.ti_estimated_sheets_required.text)
                    estimated_time = int(self.ti_estimated_time.text)
                    start_date = 'N/A'
                    end_date = self.return_earliest_date_order_parts(selected_order_parts)
                    priority = 0
                    order_parts = self.generate_task_order_parts_rowid_string()
                    aggregated_tasks = ''

                    # print((current_path, original_path, machine, material_ref, notes, estimated_sheets_required, estimated_time, start_date, end_date, priority, order_parts))
                    insert_new_task_database(current_path, original_path, machine, material_ref, notes, estimated_sheets_required, estimated_time, start_date,
                                             end_date, priority, order_parts, aggregated_tasks)
                    production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.update_display()
                    production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sidebar.pop_new_task.dismiss()
                else:
                    popup_warning = PopupWarningMessage('Nenhuma Tarefa Selecionada')
                    popup_warning.open()
            except IndexError:
                popup_warning = PopupWarningMessage('Há alguma encomenda sem data de entrega')
                popup_warning.open()

        def return_earliest_date_order_parts(self, selected_order_parts):
            earliest_date = datetime.datetime(year=int(selected_order_parts[0].due_date.split('/')[2]), month=int(selected_order_parts[0].due_date.split('/')[1]),
                                              day=int(selected_order_parts[0].due_date.split('/')[0]), hour=17, minute=45)
            for order_part in selected_order_parts:
                new_date = datetime.datetime(year=int(order_part.due_date.split('/')[2]), month=int(order_part.due_date.split('/')[1]), day=int(order_part.due_date.split('/')[0]), hour=17, minute=45)

                if new_date < earliest_date:
                    earliest_date = new_date

            return return_formatted_datetime(earliest_date)


class Taskspage_Sidebar(BoxLayout):
    # TODO Fix issues when joining a group of a tasks with a new task
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.orientation = 'vertical'
        self.size_hint_x = 0.08

        self.total_time_LF = 0
        self.total_time_LC5 = 0
        self.planned_order_parts = {}

        self.button_return = Button(text='<<<', font_size='16', halign='center', background_color=color_light_blue)
        self.button_create_new_task = Button(text='Nova\nTarefa', halign='center', font_size='14', background_color=color_light_blue)
        self.button_join_tasks = Button(text='Agrupar\nTarefas\nSelecionadas', halign='center', font_size='14', background_color=color_light_blue)
        self.button_split_tasks = Button(text='Desagrupar\nTarefas\nSelecionadas', halign='center', font_size='14', background_color=color_light_blue)
        self.button_delete_tasks = Button(text='Apagar\nTarefas\nSelecionadas', halign='center', font_size='14', background_color=color_light_blue)
        self.button_produce_tasks = Button(text='Produzir\nTarefas\nSelecionadas', halign='center', font_size='14', background_color=color_light_blue)
        self.button_plan_tasks = Button(text='Planear\nTarefas\nSelecionadas', halign='center', font_size='14', background_color=color_light_blue)
        self.button_register_tasks = Button(text='Registar\nTarefas\nSelecionadas', halign='center', font_size='14', background_color=color_light_blue)

        self.label_machine_time_LF = Label(text=f'Tempo Total LF:\n {self.total_time_LF} H')
        self.label_machine_time_LC5 = Label(text=f'Tempo Total LC5:\n {self.total_time_LC5} H')

        self.button_return.bind(on_press=self.close_popup_window)
        self.button_create_new_task.bind(on_press=self.display_popup_new_task)
        self.button_join_tasks.bind(on_press=self.join_selected_tasks)
        self.button_split_tasks.bind(on_press=self.split_selected_tasks)
        self.button_delete_tasks.bind(on_press=lambda x: self.display_popup_warning(message='Isto irá apagar as tarefas selecionadas', type='choice', continue_func=self.delete_selected_tasks))
        self.button_produce_tasks.bind(on_press=lambda x: self.display_popup_warning(message='Isto irá produzir as tarefas selecionadas', type='choice', continue_func=self.produce_selected_tasks))
        self.button_plan_tasks.bind(on_press=self.display_popup_plan_tasks)
        self.button_register_tasks.bind(on_press=self.display_popup_register_tasks)

        self.add_widget(self.button_return)
        self.add_widget(self.button_create_new_task)
        self.add_widget(self.button_join_tasks)
        self.add_widget(self.button_split_tasks)
        self.add_widget(self.button_delete_tasks)
        self.add_widget(self.button_produce_tasks)
        self.add_widget(self.button_plan_tasks)
        self.add_widget(self.button_register_tasks)
        self.add_widget(self.label_machine_time_LF)
        self.add_widget(self.label_machine_time_LC5)

        self.set_total_time_machines()

    def close_popup_window(self, *args):
        production_planning.homepage.layout_sidebar.pop_tasks.dismiss()

    def display_popup_new_task(self, *args):
        self.layout_new_task = NewTaskpage()
        self.pop_new_task = Popup(title='Nova tarefa: ', content=self.layout_new_task, size_hint=(0.8, 0.85))
        self.pop_new_task.open()

    def join_selected_tasks(self, *args):
        selected_tasks = []
        current_paths = []
        machines = []
        materials = []
        total_num_sheets = 0
        total_time = 0
        start_dates = []
        end_dates = []
        priorities = []
        time_ratios = []
        num_sheets_ratios = []

        for added_task in production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.added_tasks:
            if added_task.is_selected:
                selected_tasks.append(added_task.rowid)
                current_paths.append(added_task.current_path)
                machines.append(added_task.machine)
                materials.append(added_task.material)
                total_num_sheets += float(added_task.estimated_num_sheets)
                total_time += int(added_task.time)
                start_dates.append(added_task.start_date)
                end_dates.append(added_task.end_date)
                priorities.append(added_task.priority)

        for added_task in production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.added_tasks:
            for selected_task in selected_tasks:
                if selected_task == added_task.rowid:
                    time_ratio = int(added_task.time) / total_time
                    num_sheets_ratio = float(added_task.estimated_num_sheets) / total_num_sheets

                    time_ratios.append((added_task, time_ratio))
                    num_sheets_ratios.append((added_task, num_sheets_ratio))

        # print(f'TIME RATIOS = {time_ratios}')
        # print(f'NUM SHEETS RATIOS = {num_sheets_ratios}')

        current_paths = list(set(current_paths))
        machines = list(set(machines))
        materials = list(set(materials))
        start_dates = list(set(start_dates))
        end_dates = list(set(end_dates))
        priorities = list(set(priorities))
        rowid_string = ','.join(str(selected_task) for selected_task in sorted(selected_tasks))

        layout_join_selected_tasks = BoxLayout(orientation='vertical')
        line_1 = BoxLayout(orientation='horizontal', size_hint_y=0.1, padding=(5, 5))
        line_2 = BoxLayout(orientation='horizontal', size_hint=(0.5, 0.1), padding=(5, 5))
        line_3 = BoxLayout(orientation='horizontal', size_hint_y=0.1, padding=(5, 5))
        line_4 = BoxLayout(orientation='horizontal', size_hint_y=0.1, padding=(5, 5))
        line_5 = BoxLayout(orientation='horizontal', size_hint_y=0.15, padding=(5, 5))

        label_empty_line_2_start = Label(size_hint_x=0.2)
        label_empty_line_2_end = Label(size_hint_x=0.1)

        label_current_path = Label(text='Caminho Atual:', size_hint_x=0.15, halign='right')

        drp_current_path = DropDown()
        button_current_path = Button(text=current_paths[0], font_size='16', halign='center', background_color=color_light_blue, padding=(5, 5))
        button_current_path.bind(on_release=drp_current_path.open)
        drp_current_path.bind(on_select=lambda instance, new_val: setattr(button_current_path, 'text', new_val))
        for current_path in current_paths:
            button_current_path_option = Button(text=current_path, font_size='16', halign='center', size_hint_y=None, height=30, background_color=color_light_blue, padding=(5, 5))
            button_current_path_option.bind(on_release=lambda current_path_object: drp_current_path.select(current_path_object.text))

            drp_current_path.add_widget(button_current_path_option)

        drp_machine = DropDown()
        button_machine = Button(text=machines[0], font_size='16', halign='center', size_hint_x=0.2, background_color=color_light_blue, padding=(5, 5))
        button_machine.bind(on_release=drp_machine.open)
        drp_machine.bind(on_select=lambda instance, new_val: setattr(button_machine, 'text', new_val))
        for machine in machines:
            button_machine_option = Button(text=machine, font_size='16', halign='center', size_hint_y=None, height=30, background_color=color_light_blue, padding=(5, 5))
            button_machine_option.bind(on_release=lambda machine_object: drp_machine.select(machine_object.text))

            drp_machine.add_widget(button_machine_option)

        ti_time = TextInput(text=str(total_time), size_hint_x=0.2)
        label_time = Label(text=' min', font_size='16', size_hint_x=0.05)

        label_empty_line_3_start = Label(size_hint_x=0.125)
        label_empty_line_3_end = Label(size_hint_x=0.275)

        drp_material = DropDown()
        button_material = Button(text=materials[0], font_size='16', halign='center', size_hint_x=0.4, background_color=color_light_blue, padding=(5, 5))
        button_material.bind(on_release=drp_material.open)
        drp_material.bind(on_select=lambda instance, new_val: setattr(button_material, 'text', new_val))
        for material in materials:
            button_material_option = Button(text=material, font_size='16', halign='center', size_hint_y=None, height=30, background_color=color_light_blue, padding=(5, 5))
            button_material_option.bind(on_release=lambda material_object: drp_material.select(material_object.text))

            drp_material.add_widget(button_material_option)

        ti_num_sheets = TextInput(text=str(total_num_sheets), size_hint_x=0.1)
        label_num_sheets = Label(text=' chapas', font_size='16', size_hint_x=0.05)

        label_empty_line_4_start = Label(size_hint_x=0.12)
        label_empty_line_4_end = Label(size_hint_x=0.18)

        label_start_date = Label(text='Data Início: ', font_size='16', size_hint_x=0.1)

        drp_start_date = DropDown()
        button_start_date = Button(text=start_dates[0], font_size='16', halign='center', size_hint_x=0.15, background_color=color_light_blue, padding=(5, 5))
        button_start_date.bind(on_release=drp_start_date.open)
        drp_start_date.bind(on_select=lambda instance, new_val: setattr(button_start_date, 'text', new_val))
        for start_date in start_dates:
            button_start_date_option = Button(text=start_date, font_size='16', halign='center', size_hint_y=None, height=30, background_color=color_light_blue, padding=(5, 5))
            button_start_date_option.bind(on_release=lambda start_date_object: drp_start_date.select(start_date_object.text))

            drp_start_date.add_widget(button_start_date_option)

        label_end_date = Label(text='Data Fim: ', font_size='16', size_hint_x=0.1)

        drp_end_date = DropDown()
        button_end_date = Button(text=end_dates[0], font_size='16', halign='center', size_hint_x=0.15, background_color=color_light_blue, padding=(5, 5))
        button_end_date.bind(on_release=drp_end_date.open)
        drp_end_date.bind(on_select=lambda instance, new_val: setattr(button_end_date, 'text', new_val))
        for end_date in end_dates:
            button_end_date_option = Button(text=end_date, font_size='16', halign='center', size_hint_y=None, height=30, background_color=color_light_blue,
                                              padding=(5, 5))
            button_end_date_option.bind(on_release=lambda end_date_object: drp_end_date.select(end_date_object.text))

            drp_end_date.add_widget(button_end_date_option)

        label_priority = Label(text='Prioridade: ', font_size='16', size_hint_x=0.1)
        ti_priority = TextInput(text=str(max(priorities)), size_hint_x=0.05)

        button_confirm = Button(text='Confirmar', font_size='16', halign='center', size_hint_x=0.2, background_color=color_light_green, padding=(5, 5))
        button_confirm.bind(on_press=lambda inst: self.submit_changes_join_tasks(button_current_path.text, button_machine.text, ti_time.text, button_material.text, ti_num_sheets.text, button_start_date.text, button_end_date.text, ti_priority.text, rowid_string, time_ratios, num_sheets_ratios))

        layout_join_selected_tasks.add_widget(line_1)
        layout_join_selected_tasks.add_widget(line_2)
        layout_join_selected_tasks.add_widget(line_3)
        layout_join_selected_tasks.add_widget(line_4)
        layout_join_selected_tasks.add_widget(line_5)

        line_1.add_widget(label_current_path)
        line_1.add_widget(button_current_path)
        line_2.add_widget(label_empty_line_2_start)
        line_2.add_widget(button_machine)
        line_2.add_widget(ti_time)
        line_2.add_widget(label_time)
        line_2.add_widget(label_empty_line_2_end)
        line_3.add_widget(label_empty_line_3_start)
        line_3.add_widget(button_material)
        line_3.add_widget(ti_num_sheets)
        line_3.add_widget(label_num_sheets)
        line_3.add_widget(label_empty_line_3_end)
        line_4.add_widget(label_empty_line_4_start)
        line_4.add_widget(label_start_date)
        line_4.add_widget(button_start_date)
        line_4.add_widget(label_end_date)
        line_4.add_widget(button_end_date)
        line_4.add_widget(label_priority)
        line_4.add_widget(ti_priority)
        line_4.add_widget(label_empty_line_4_end)
        line_5.add_widget(button_confirm)

        # self.split_selected_tasks() This was put here in an attempt to fix issues when trying to join a single task to an already added group of tasks

        self.popup_join_selected_tasks = Popup(title='Preencher dados comuns para as tarefas selecionadas', content=layout_join_selected_tasks, size_hint=(0.6, 0.3))
        self.popup_join_selected_tasks.open()

    def submit_changes_join_tasks(self, current_path, machine, time, material, num_sheets, start_date, end_date, priority, rowid_string, time_ratios, num_sheets_ratios):
        initial_time = int(time)
        initial_num_sheets = float(num_sheets)
        for rowid in rowid_string.split(','):
            for added_task in production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.added_tasks:
                time = initial_time
                num_sheets = initial_num_sheets
                if added_task.rowid == int(rowid):

                    for time_ratio in time_ratios:
                        if time_ratio[0].rowid == added_task.rowid:
                            time = int(int(time) * time_ratio[1])

                    for num_sheets_ratio in num_sheets_ratios:
                        if num_sheets_ratio[0].rowid == added_task.rowid:
                            num_sheets = float(round(float(num_sheets) * num_sheets_ratio[1]))

                    print('ROW STRING UPON JOINING = ' + rowid_string)
                    added_task.change_value(current_path, 'current_path', added_task.rowid)
                    added_task.change_value(machine, 'machine', added_task.rowid)
                    added_task.change_value(int(time), 'estimated_time', added_task.rowid)
                    added_task.change_value(material, 'material_name', added_task.rowid)
                    added_task.change_value(float(num_sheets), 'estimated_sheets_required', added_task.rowid)
                    added_task.change_value(start_date, 'start_date', added_task.rowid)
                    added_task.change_value(end_date, 'end_date', added_task.rowid)
                    added_task.change_value(int(priority), 'priority', added_task.rowid)
                    added_task.change_value(rowid_string, 'aggregated_tasks', added_task.rowid)
                    added_task.aggregated_tasks_string = rowid_string

                    production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.update_aggregated_tasks_dict()

                    added_task.set_aggregated_index()
                    added_task.set_text_buttons_paths()

        production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.update_display()
        # print((current_path, machine, time, material, num_sheets, start_date, end_date, rowid_string))

        self.popup_join_selected_tasks.dismiss()

    def split_selected_tasks(self, *args):
        for added_task in production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.added_tasks:
            if added_task.is_selected:

                added_task.aggregated_tasks_string = str(added_task.rowid)
                added_task.aggregated_tasks = [added_task.rowid]
                added_task.aggregated_index = ''

                change_value_task_database(added_task.aggregated_tasks_string, 'aggregated_tasks', added_task.rowid)
                change_value_task_database(added_task.aggregated_index, 'aggregated_index', added_task.rowid)

                added_task.set_total_time()
                added_task.set_text_button_time()
                added_task.set_total_num_sheets()
                added_task.set_text_button_num_sheets()
                added_task.set_text_buttons_paths()


        production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.update_display()

    def delete_selected_tasks(self, *args):
        any_selected = False
        for added_task in production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.added_tasks:
            if added_task.is_selected:
                production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.layout_task_list.layout_scroll.remove_widget(added_task)
                remove_task_database(added_task.rowid)
                any_selected = True

        production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.layout_filters_bar.checkbox_select_all.state = 'normal'

        if not any_selected:
            self.display_popup_warning('Nenhuma tarefa selecionada')
        else:
            self.display_popup_warning('Ta feito')

    def produce_selected_tasks(self, *args):
        any_selected = False
        for task in production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.added_tasks:
            if task.is_selected:
                for task_order_part_rowid in task.order_parts_rowids:
                    produced_quantity = load_order_parts_from_database(mode='produced_quantity', rowid=task_order_part_rowid)[0]
                    for task_order_part in task.order_parts_objects:
                        if task_order_part_rowid == task_order_part.rowid:
                            total_produced_quantity = produced_quantity + int(task_order_part.pending_quantity)
                            change_value_order_part_database(total_produced_quantity, 'produced_quantity', task_order_part.rowid)
                production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.layout_task_list.layout_scroll.remove_widget(task)
                remove_task_database(task.rowid)
                any_selected = True
                # TODO After developing productions interface, insert a new production here

        if not any_selected:
            self.display_popup_warning('Nenhuma tarefa selecionada')
        else:
            self.display_popup_warning('Ta feito')

    def display_popup_plan_tasks(self, *args):
        any_selected = False
        for task in production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.added_tasks:
            if task.is_selected:
                any_selected = True

        if not any_selected:
            self.display_popup_warning('Nenhuma tarefa selecionada')
        else:
            self.layout_plan_tasks = Planning()
            self.pop_plan_tasks = Popup(title='Planear Tarefas', content=self.layout_plan_tasks, size_hint=(0.8, 0.85))
            self.pop_plan_tasks.open()

    def plan_selected_tasks(self, starting_time_LF, starting_time_LC5, starting_date_LF, starting_date_LC5, *args):
        selected_tasks_LF = []
        selected_tasks_LC5 = []

        for task in production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.added_tasks:
            if task.is_selected:
                if task.machine == 'LF-3015':
                    selected_tasks_LF.append(task)
                else:
                    selected_tasks_LC5.append(task)

        selected_tasks_LF.sort(key=lambda x: int(x.priority), reverse=False)
        selected_tasks_LC5.sort(key=lambda x: int(x.priority), reverse=False)

        starting_date_obj_LF = datetime.datetime(year=int(starting_date_LF.split('/')[2]), month=int(starting_date_LF.split('/')[1]),
                                                 day=int(starting_date_LF.split('/')[0]), hour=int(starting_time_LF.split(':')[0]),
                                                 minute=int(starting_time_LF.split(':')[1]))
        starting_date_obj_LC5 = datetime.datetime(year=int(starting_date_LC5.split('/')[2]), month=int(starting_date_LC5.split('/')[1]),
                                                 day=int(starting_date_LC5.split('/')[0]), hour=int(starting_time_LC5.split(':')[0]),
                                                 minute=int(starting_time_LC5.split(':')[1]))

        start_times_list_LF = self.get_available_shift_times('LF3015', starting_date_obj_LF)[0]
        start_times_list_LC5 = self.get_available_shift_times('LC5', starting_date_obj_LC5)[0]

        stop_times_list_LF = self.get_available_shift_times('LF3015', starting_date_obj_LF)[1]
        stop_times_list_LC5 = self.get_available_shift_times('LC5', starting_date_obj_LC5)[1]

        start_datetime_LF = round_time(starting_date_obj_LF)
        start_datetime_LC5 = round_time(starting_date_obj_LC5)

        task_num_cells_dict_LF = {}
        total_num_cells_LF = 0
        for task in selected_tasks_LF:
            task_num_cells = math.ceil(int(task.time) / 15)
            total_num_cells_LF += task_num_cells
            task_num_cells_dict_LF.update({task: task_num_cells})

        task_num_cells_dict_LC5 = {}
        total_num_cells_LC5 = 0
        for task in selected_tasks_LC5:
            task_num_cells = math.ceil(int(task.time) / 15)
            total_num_cells_LC5 += task_num_cells
            task_num_cells_dict_LC5.update({task: task_num_cells})

        # print(f'TOTAL NUM CELLS LF: {str(total_num_cells_LF)}')
        # print(f'TOTAL NUM CELLS LC5: {str(total_num_cells_LC5)}')
        #
        # print(task_num_cells_dict_LF)
        # print(task_num_cells_dict_LC5)

        # print(start_times_list_LF)
        # print(stop_times_list_LF)
        # print(start_times_list_LC5)
        # print(stop_times_list_LC5)
        datetimes_LF = self.return_datetimes_list(start_datetime_LF, start_times_list_LF, stop_times_list_LF, total_num_cells_LF)
        datetimes_LC5 = self.return_datetimes_list(start_datetime_LC5, start_times_list_LC5, stop_times_list_LC5, total_num_cells_LC5)

        # print('DATETIMES LF')
        # for d in datetimes_LF:
        #     print(d)
        # print('____________________________________________________________________________')
        # print('DATETIMES LC5')
        # for d in datetimes_LC5:
        #     print(d)
        # print('____________________________________________________________________________')

        self.build_excel_plan(task_num_cells_dict_LF, datetimes_LF, task_num_cells_dict_LC5, datetimes_LC5)

    def return_datetimes_list(self, start_datetime, start_times_list, stop_times_list, total_num_cells):
        datetimes = []

        start_time = start_datetime
        while total_num_cells != 0:
            # print(f'NUM TOTAL CELLS: {total_num_cells}')
            stop_time = self.return_next_available_time(start_time, stop_times_list)

            times_list = return_times_between_times(start_time, stop_time, datetime.timedelta(minutes=15))

            if total_num_cells > len(times_list):
                # print(f'{str(total_num_cells)} > {str(len(times_list))} = {total_num_cells > len(times_list)}')
                for time in times_list:
                    datetimes.append(time)
                    # print(f'APPENDING {str(time)}')
                total_num_cells -= len(times_list)
            else:
                # print(f'{str(total_num_cells)} == {str(len(times_list))} = {total_num_cells == len(times_list)}')
                for i in range(total_num_cells):
                    datetimes.append(times_list[i])
                    # print(f'APPENDING {str(times_list[i])}')
                total_num_cells = 0

            start_time = self.return_next_available_time(stop_time, start_times_list)

        return datetimes

    def return_next_available_time(self, test_datetime, shift_times):
        # print('CHECKING NEXT AVAILABLE TIME:')
        for shift_time in shift_times:
            # print(f'{shift_time} > {test_datetime}')
            # print(shift_time > test_datetime)
            if shift_time > test_datetime:
                return shift_time

        # print('INCREMENTING ONE DAY')
        for i in range(len(shift_times)):
            shift_times[i] = return_next_weekday(shift_times[i] + datetime.timedelta(days=1))
        # print(f'returning next week day: {shift_times[0]}')
        return shift_times[0]

    def get_available_shift_times(self, machine, starting_date_obj):
        start_times = []
        shift_list = self.return_shift_list(machine)

        if len(shift_list) == 3:
            num_shifts = 1
            for shift in shift_list:
                break_start_obj = datetime.datetime(year=starting_date_obj.year, month=starting_date_obj.month, day=starting_date_obj.day,
                                                    hour=int(shift[3].split(':')[0]), minute=int(shift[3].split(':')[1]))
                break_duration = datetime.timedelta(minutes=int(shift[4]))

                start_times.append(break_start_obj + break_duration)
                num_shifts += 1
        else:
            num_shifts = 1
            for shift in shift_list:
                shift_start_obj = datetime.datetime(year=starting_date_obj.year, month=starting_date_obj.month, day=starting_date_obj.day,
                                                    hour=int(shift[1].split(':')[0]), minute=int(shift[1].split(':')[1]))
                break_start_obj = datetime.datetime(year=starting_date_obj.year, month=starting_date_obj.month, day=starting_date_obj.day,
                                                    hour=int(shift[3].split(':')[0]), minute=int(shift[3].split(':')[1]))
                break_duration = datetime.timedelta(minutes=int(shift[4]))

                start_times.append(shift_start_obj)
                start_times.append(break_start_obj + break_duration)
                num_shifts += 1

        if len(shift_list) == 1:
            stop_times = [datetime.datetime(year=starting_date_obj.year, month=starting_date_obj.month, day=starting_date_obj.day,
                                                hour=int(shift_list[0][3].split(':')[0]), minute=int(shift_list[0][3].split(':')[1])),
                          datetime.datetime(year=starting_date_obj.year, month=starting_date_obj.month, day=starting_date_obj.day,
                                            hour=int( shift_list[0][2].split(':')[0]), minute=int( shift_list[0][2].split(':')[1]))]
        elif len(shift_list) == 2:
            stop_times = [datetime.datetime(year=starting_date_obj.year, month=starting_date_obj.month, day=starting_date_obj.day,
                                            hour=int(shift_list[0][3].split(':')[0]), minute=int(shift_list[0][3].split(':')[1])),
                          datetime.datetime(year=starting_date_obj.year, month=starting_date_obj.month, day=starting_date_obj.day,
                                            hour=int(shift_list[1][3].split(':')[0]), minute=int(shift_list[1][3].split(':')[1])),
                          datetime.datetime(year=starting_date_obj.year, month=starting_date_obj.month, day=starting_date_obj.day,
                                            hour=int(shift_list[1][2].split(':')[0]), minute=int(shift_list[1][2].split(':')[1]))]
        else:
            stop_times = [datetime.datetime(year=starting_date_obj.year, month=starting_date_obj.month, day=starting_date_obj.day,
                                            hour=int(shift_list[0][3].split(':')[0]), minute=int(shift_list[0][3].split(':')[1])),
                          datetime.datetime(year=starting_date_obj.year, month=starting_date_obj.month, day=starting_date_obj.day,
                                            hour=int(shift_list[1][3].split(':')[0]), minute=int(shift_list[1][3].split(':')[1])),
                          datetime.datetime(year=starting_date_obj.year, month=starting_date_obj.month, day=starting_date_obj.day,
                                            hour=int(shift_list[2][3].split(':')[0]), minute=int(shift_list[2][3].split(':')[1]))]

        previous_time = start_times[0]
        for time in start_times:
            if time < previous_time:
                start_times[start_times.index(time)] = return_next_weekday(start_times[start_times.index(time)] + datetime.timedelta(days=1))

            previous_time = time

        previous_time = stop_times[0]
        for time in stop_times:
            if time < previous_time:
                stop_times[stop_times.index(time)] = return_next_weekday(stop_times[stop_times.index(time)] + datetime.timedelta(days=1))

            previous_time = time
        #
        # print(f'START TIMES STARTING AT {starting_date_obj} for machine {machine} working with {num_shifts - 1} shifts')
        # print(start_times)
        # print(f'STOP TIMES STARTING AT {starting_date_obj} for machine {machine} working with {num_shifts - 1} shifts')
        # print(stop_times)
        return (start_times, stop_times)

    def return_shift_list(self, machine):
        shift_list = []
        if machine == 'LF3015':
            for shift in self.layout_plan_tasks.added_shifts_LF:
                shift_list.append([machine, shift.ti_time_start.text, shift.ti_time_finish.text, shift.ti_time_break.text, shift.ti_break_duration.text])
        else:
            for shift in self.layout_plan_tasks.added_shifts_LC5:
                shift_list.append([machine, shift.ti_time_start.text, shift.ti_time_finish.text, shift.ti_time_break.text, shift.ti_break_duration.text])

        return shift_list

    def build_excel_plan(self, task_num_cells_dict_LF, datetimes_LF, task_num_cells_dict_LC5, datetimes_LC5):

        outWorkbook = xls.Workbook('Planeamento_Laser.xlsx')

        gen_format_1 = outWorkbook.add_format({
            'bold': 0,
            'border': 2,
            'align': 'center',
            'valign': 'vcenter',
            'fg_color': 'FFFC99',
            'font_size': 12
        })

        gen_format_2 = outWorkbook.add_format({
            'bold': 0,
            'border': 2,
            'align': 'center',
            'valign': 'vcenter',
            'fg_color': 'FFD899',
            'font_size': 12
        })

        gen_format_3 = outWorkbook.add_format({
            'bold': 0,
            'border': 2,
            'align': 'center',
            'valign': 'vcenter',
            'fg_color': 'DBFF99',
            'font_size': 12
        })

        header_format = outWorkbook.add_format({
            'bold': 1,
            'border': 2,
            'align': 'center',
            'valign': 'vcenter',
            'fg_color': 'silver',
            'font_size': 13
        })

        color_iterator = cycle([gen_format_1, gen_format_2, gen_format_3])

        outSheetPlanning = outWorkbook.add_worksheet('Planeamento')

        # LF PLANNING
        outSheetPlanning.write(1, 1, 'LF-3015', header_format)
        outSheetPlanning.write(1, 2, len(production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sidebar.layout_plan_tasks.added_shifts_LF),
                               header_format)
        outSheetPlanning.write(1, 3, 'Turnos', header_format)

        datetimes_LF_iter = cycle(datetimes_LF)
        previous_date = datetimes_LF[0]

        date_row = 2
        time_row = 3
        job_row = 4
        date_col = 1
        time_col = 1
        job_col = 1
        task_counter = 0
        for task in task_num_cells_dict_LF:
            task_counter += 1
            task_name = f'T{str(task_counter)}-{task.material}'
            task_input_message = self.generate_task_input_message(task.order_parts_objects)

            # print(f'{task_name} = {str(task_num_cells_dict_LF[task])} CELLS')
            finish_time = 0
            for _ in range(task_num_cells_dict_LF[task]):
                time = next(datetimes_LF_iter)
                formatted_time = f'{time.hour}:{time.minute}h'
                outSheetPlanning.write(time_row, time_col, formatted_time, header_format)
                # print(f'Writing: {formatted_time}')

                # CHECKING FOR DAY CHANGE
                if time.day != previous_date.day or time.month != previous_date.month or time.year != previous_date.year:
                    # print('CHANGING DAY')
                    formatted_date = f'{previous_date.day}/{previous_date.month}/{previous_date.year}'
                    # print(f'CHANGING PREVIOUS DATE FROM {str(previous_date)} TO {str(time)}')
                    previous_date = time

                    # print(f'{date_col} == {time_col} = {date_col == time_col}')
                    if date_col == time_col:
                        outSheetPlanning.write(date_row, date_col, formatted_date, header_format)
                        # print(f'WRITING {formatted_date} ON {date_row},{date_col}')
                    else:
                        outSheetPlanning.merge_range(date_row, date_col, date_row, time_col - 1, formatted_date, header_format)
                        # print(f'MERGING {formatted_date} ON {date_row},{date_col},{date_row},{time_col - 1}')

                    date_col = time_col

                time_col += 1

                finish_time = time

            # JOB/TASK FILL
            if job_col == time_col - 1:
                outSheetPlanning.write(job_row, job_col, task_name, next(color_iterator))
                # print(f'WRITING {task_name} ON {job_row},{job_col}')
            else:
                outSheetPlanning.merge_range(job_row, job_col, job_row, time_col - 1, task_name, next(color_iterator))
                # print(f'MERGING {task_name} ON {job_row},{job_col},{job_row},{time_col - 1}')

            outSheetPlanning.data_validation(job_row, job_col, job_row, time_col - 1, {'validate': 'integer',
                                                                                   'criteria': '<',
                                                                                   'value': 10,
                                                                                   'input_title': 'Cliente | Encomenda',
                                                                                   'input_message': task_input_message})
            job_col = time_col

            # UPDATING PLANNED ORDER PARTS DICT
            for order_part in task.order_parts_objects:
                self.planned_order_parts.update({order_part: finish_time})

        formatted_date = f'{previous_date.day}/{previous_date.month}/{previous_date.year}'
        if date_col == time_col:
            outSheetPlanning.write(date_row, date_col, formatted_date, header_format)
        else:
            outSheetPlanning.merge_range(date_row, date_col, date_row, time_col - 1, formatted_date, header_format)

        # LC5 PLANNING
        outSheetPlanning.write(6, 1, 'LC5', header_format)
        outSheetPlanning.write(6, 2, len(production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sidebar.layout_plan_tasks.added_shifts_LC5),
                               header_format)
        outSheetPlanning.write(6, 3, 'Turnos', header_format)

        datetimes_LC5_iter = cycle(datetimes_LC5)
        previous_date = datetimes_LC5[0]

        date_row = 7
        time_row = 8
        job_row = 9
        date_col = 1
        time_col = 1
        job_col = 1
        task_counter = 0
        for task in task_num_cells_dict_LC5:
            task_counter += 1
            task_name = f'T{str(task_counter)}-{task.material}'
            task_input_message = self.generate_task_input_message(task.order_parts_objects)

            # print(f'{task_name} = {str(task_num_cells_dict_LF[task])} CELLS')
            finish_time = 0
            for _ in range(task_num_cells_dict_LC5[task]):
                time = next(datetimes_LC5_iter)
                formatted_time = f'{time.hour}:{time.minute}h'
                outSheetPlanning.write(time_row, time_col, formatted_time, header_format)
                # print(f'Writing: {formatted_time}')

                # CHECKING FOR DAY CHANGE
                if time.day != previous_date.day or time.month != previous_date.month or time.year != previous_date.year:
                    # print('CHANGING DAY')
                    formatted_date = f'{previous_date.day}/{previous_date.month}/{previous_date.year}'
                    # print(f'CHANGING PREVIOUS DATE FROM {str(previous_date)} TO {str(time)}')
                    previous_date = time

                    # print(f'{date_col} == {time_col} = {date_col == time_col}')
                    if date_col == time_col:
                        outSheetPlanning.write(date_row, date_col, formatted_date, header_format)
                        # print(f'WRITING {formatted_date} ON {date_row},{date_col}')
                    else:
                        outSheetPlanning.merge_range(date_row, date_col, date_row, time_col - 1, formatted_date, header_format)
                        # print(f'MERGING {formatted_date} ON {date_row},{date_col},{date_row},{time_col - 1}')

                    date_col = time_col

                time_col += 1

                finish_time = time

            # JOB/TASK FILL
            if job_col == time_col - 1:
                outSheetPlanning.write(job_row, job_col, task_name, next(color_iterator))
                # print(f'WRITING {task_name} ON {job_row},{job_col}')
            else:
                outSheetPlanning.merge_range(job_row, job_col, job_row, time_col - 1, task_name, next(color_iterator))
                # print(f'MERGING {task_name} ON {job_row},{job_col},{job_row},{time_col - 1}')

            outSheetPlanning.data_validation(job_row, job_col, job_row, time_col - 1, {'validate': 'integer',
                                                                                       'criteria': '<',
                                                                                       'value': 10,
                                                                                       'input_title': 'Cliente | Encomenda',
                                                                                       'input_message': task_input_message})
            job_col = time_col

            # UPDATING PLANNED ORDER PARTS DICT
            for order_part in task.order_parts_objects:
                self.planned_order_parts.update({order_part: finish_time})

        formatted_date = f'{previous_date.day}/{previous_date.month}/{previous_date.year}'
        if date_col == time_col:
            outSheetPlanning.write(date_row, date_col, formatted_date, header_format)
        else:
            outSheetPlanning.merge_range(date_row, date_col, date_row, time_col - 1, formatted_date, header_format)

        # DATE GENERATED
        now_date = datetime.datetime.now()
        outSheetPlanning.write(12, 1, 'PLANEADO A:', header_format)
        outSheetPlanning.write(12, 2, f'{now_date.hour}:{now_date.minute}', header_format)
        outSheetPlanning.write(12, 3, f'{now_date.day}/{now_date.month}/{now_date.year}', header_format)

        # ORDERS PAGE
        outSheetOrders = outWorkbook.add_worksheet('Estados Enc.')

        starting_row = 1
        starting_col = 1

        # Write header
        outSheetOrders.write(starting_row, starting_col, 'Cliente', header_format)
        outSheetOrders.write(starting_row, starting_col + 1, 'Encomenda', header_format)
        outSheetOrders.write(starting_row, starting_col + 2, 'Estado', header_format)
        outSheetOrders.write(starting_row, starting_col + 3, 'Data Objetivo', header_format)
        outSheetOrders.write(starting_row, starting_col + 4, 'Data Conclusão Corte', header_format)
        outSheetOrders.write(starting_row, starting_col + 5, 'Op. Adic.', header_format)
        outSheetOrders.write(starting_row, starting_col + 6, 'Nome Peça', header_format)
        outSheetOrders.write(starting_row, starting_col + 7, 'Qtd.', header_format)
        outSheetOrders.write(starting_row, starting_col + 8, 'Matéria-prima', header_format)

        starting_row += 1

        for order_part_info in load_order_parts_from_database():
            order_part_ref = order_part_info[0]
            order_part_name = order_part_info[1]
            order_part_material_ref = order_part_info[2]
            order_part_material = get_material_name(order_part_material_ref)
            order_part_quantity = order_part_info[3]
            order_part_produced_quantity = order_part_info[4]
            order_part_order_num = order_part_info[5]
            order_part_order_num_client = order_part_info[6]
            order_part_client = order_part_info[7]
            order_part_date_modified = order_part_info[8]
            order_part_due_date = order_part_info[9]
            order_part_additional_operations = order_part_info[10]
            order_part_rowid = order_part_info[11]

            state = self.select_cell_state(order_part_ref, order_part_name, order_part_material_ref, order_part_quantity, order_part_produced_quantity, order_part_order_num, order_part_order_num_client, order_part_client, order_part_date_modified, order_part_due_date, order_part_additional_operations, order_part_rowid, self.planned_order_parts)
            formatting = outWorkbook.add_format({'bold': 0, 'border': 2, 'align': 'center', 'valign': 'vcenter', 'fg_color': self.select_cell_color(state), 'font_size': 12})

            # COMPARING DUE DATE AND FINISH DATE
            finish_date_obj = self.return_finish_date(order_part_ref, order_part_name, order_part_material_ref, order_part_quantity, order_part_produced_quantity,
                                                      order_part_order_num, order_part_order_num_client, order_part_client, order_part_date_modified,
                                                      order_part_due_date, order_part_additional_operations, order_part_rowid, self.planned_order_parts)

            if order_part_due_date != '' and finish_date_obj != 'N/A':
                due_date_obj = datetime.datetime(year=int(order_part_due_date.split('/')[2]), month=int(order_part_due_date.split('/')[1]),
                                                 day=int(order_part_due_date.split('/')[0]), hour=17, minute=45)

                if finish_date_obj > due_date_obj:
                    formatting.set_font_color('red')

            outSheetOrders.set_column(1, 5, 30)
            outSheetOrders.write(starting_row, starting_col, order_part_client, formatting)
            starting_col += 1
            outSheetOrders.write(starting_row, starting_col, order_part_order_num, formatting)
            starting_col += 1
            outSheetOrders.write(starting_row, starting_col, state, formatting)
            starting_col += 1
            outSheetOrders.write(starting_row, starting_col, '17:45 - ' + order_part_due_date if order_part_due_date != '' else 'N/A', formatting)
            starting_col += 1
            outSheetOrders.write(starting_row, starting_col, return_formatted_datetime(finish_date_obj) if finish_date_obj != 'N/A' else finish_date_obj, formatting)
            starting_col += 1
            outSheetOrders.write(starting_row, starting_col, 'Sim' if order_part_additional_operations else 'Não', formatting)
            starting_col += 1
            outSheetOrders.write(starting_row, starting_col, order_part_name, formatting)
            starting_col += 1
            outSheetOrders.set_column(6, 6, 20)
            outSheetOrders.write(starting_row, starting_col, order_part_quantity, formatting)
            starting_col += 1
            outSheetOrders.set_column(7, 7, 30)
            outSheetOrders.write(starting_row, starting_col, order_part_material, formatting)

            starting_col = 1
            starting_row += 1

        outWorkbook.close()

        layout_popup_inform_success = BoxLayout(orientation='vertical')
        layout_popup_inform_success.add_widget(Label(text='Planeamento com sucesso'))
        button_open_doc = Button(text='Abrir', font_size='16', halign='center', background_color=color_light_blue)
        button_open_doc.bind(on_press=lambda x: os.startfile("Planeamento_Laser.xlsx"))
        layout_popup_inform_success.add_widget(button_open_doc)

        popup_inform_generated_successfully = Popup(title='Sucesso', content=layout_popup_inform_success, size_hint=(0.2, 0.15))
        popup_inform_generated_successfully.open()

    def select_cell_state(self, order_part_ref, order_part_name, order_part_material_ref, order_part_quantity, order_part_produced_quantity, order_part_order_num, order_part_order_num_client, order_part_client, order_part_date_modified, order_part_due_date, order_part_additional_operations, order_part_rowid , planned_order_parts):
        if order_part_produced_quantity >= order_part_quantity:
            if order_part_additional_operations:
                return 'A aguardar operações adicionais'
            else:
                return 'Concluído'
        else:
            for planned_order_part in planned_order_parts:
                if order_part_name == planned_order_part.name and str(order_part_quantity) == planned_order_part.quantity and str(order_part_produced_quantity) == planned_order_part.produced_quantity and order_part_client == planned_order_part.client and order_part_order_num == planned_order_part.order_num and order_part_rowid == planned_order_part.rowid:
                    return 'Planeado'
        return 'Não planeado'

    def select_cell_color(self, cell_state):
        if cell_state == 'Concluído':
            return 'DBFF99'
        elif cell_state == 'A aguardar operações adicionais':
            return 'FFFC99'
        elif cell_state == 'Planeado':
            return 'FFD899'
        elif cell_state == 'Não planeado':
            return 'FFFFFF'

    def return_finish_date(self, order_part_ref, order_part_name, order_part_material_ref, order_part_quantity, order_part_produced_quantity, order_part_order_num, order_part_order_num_client, order_part_client, order_part_date_modified, order_part_due_date, order_part_additional_operations, order_part_rowid, planned_order_parts):
        for planned_order_part in planned_order_parts:
            if order_part_name == planned_order_part.name and str(order_part_quantity) == planned_order_part.quantity and str(order_part_produced_quantity) == planned_order_part.produced_quantity and order_part_client == planned_order_part.client and order_part_order_num == planned_order_part.order_num and order_part_rowid == planned_order_part.rowid:
                return planned_order_parts[planned_order_part]

        return 'N/A'

    def update_times(self, num_total_cells, start_date, shift_times):
        new_shift_times = []
        iterator_shift_times = cycle(shift_times)
        day_multiplier = 0

        index_time_start = shift_times.index(start_date)

        for _ in range(index_time_start):
            next(iterator_shift_times)

        last_date = start_date
        for _ in range(num_total_cells):
            current_date = next(iterator_shift_times)
            current_date = current_date.replace(day=last_date.day, month=last_date.month, year=last_date.year)

            current_date = return_next_weekday(current_date + datetime.timedelta(days=day_multiplier))
            last_date = current_date

            new_shift_times.append(current_date)

            if current_date.hour == shift_times[-1].hour and current_date.minute == shift_times[-1].minute:
                last_date += datetime.timedelta(days=1)

        return new_shift_times

    def return_next_valid_start_time(self, starting_time, shift_times):
        starting_hours = int(starting_time.split(':')[0])
        starting_minutes = int(starting_time.split(':')[1])
        while True:
            for shift_time in shift_times:
                if starting_hours == shift_time.hour and starting_minutes == shift_time.minute:
                    return starting_time

            start_time_obj = datetime.datetime(year=shift_times[0].year, month=shift_times[0].month, day=shift_times[0].day, hour=starting_hours,
                                               minute=starting_minutes)
            start_time_obj += datetime.timedelta(minutes=15)
            starting_hours = start_time_obj.hour
            starting_minutes = start_time_obj.minute
            starting_time = str(starting_hours) + ':' + str(starting_minutes)

    def generate_task_input_message(self, order_parts_objects):
        task_input_message_list = []
        for obj in order_parts_objects:
            task_input_message_list.append(f'{obj.client} | {obj.order_num}\n')

        task_input_message = ''.join(set(task_input_message_list))
        return task_input_message

    def display_popup_warning(self, message, type='warning', continue_func=''):
        self.popup_warning_delete = PopupWarningMessage(message=message, type=type, continue_func=continue_func)
        self.popup_warning_delete.open()

    def set_total_time_machines(self, *args):
        current_tasks = [task for task in load_tasks_from_database()]
        time_dict = {'LF-3015': 0, 'LC5': 0}

        for curr_task in current_tasks:
            time_dict[curr_task[2]] += curr_task[6]

        self.total_time_LF = time_dict['LF-3015']
        self.total_time_LC5 = time_dict['LC5']

        self.label_machine_time_LF.text = str(f'Tempo Total LF:\n {return_formatted_time_string(self.total_time_LF)}')
        self.label_machine_time_LC5.text = str(f'Tempo Total LC5:\n {return_formatted_time_string(self.total_time_LC5)}')

    def display_popup_register_tasks(self, *args):
        any_selected = False
        for task in production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.added_tasks:
            if task.is_selected:
                any_selected = True

        if not any_selected:
            self.display_popup_warning('Nenhuma tarefa selecionada')
        else:
            suggested_date = return_next_weekday(datetime.datetime.now())
            suggested_date_string = f'{suggested_date.day}/{suggested_date.month}/{suggested_date.year}'

            self.layout_register_tasks = BoxLayout(orientation='vertical')
            self.layout_register_tasks.add_widget(Label(text='Insere a data do registo (dd/mm/aaaa)'))
            ti_date = TextInput(text=suggested_date_string, size_hint_y=0.5)
            self.layout_register_tasks.add_widget(ti_date)
            self.layout_register_tasks.add_widget(Label(size_hint_y=0.5))
            self.layout_register_tasks.add_widget(Button(text='OK', background_color=color_light_green, on_press=lambda x: self.register_tasks_excel(ti_date.text)))

            self.pop_register_tasks = Popup(title='Registar Tarefas', content=self.layout_register_tasks, size_hint=(0.3, 0.3))
            self.pop_register_tasks.open()

    def register_tasks_excel(self, string_date_register, *args):
        try:
            df = pd.read_excel(CAMINHO_EXCEL)

            selected_tasks = []
            counter_part = 1

            for task in production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.added_tasks:
                if task.is_selected:
                    selected_tasks.append(task)

            counter_task = 1
            for task in selected_tasks:
                task_name = string_date_register + "_T" + str(counter_task)
                task_machine = task.machine
                task_path = task.current_path
                task_material_name = task.material.split(',')[0].split(' ')[0]
                task_material_spec = task.material.split(',')[0].split(' ')[1]
                task_material_thickness = task.material.split(',')[1].split(' ')[1]
                task_material_owner = task.material.split(',')[1].split(' ')[2]

                for part in task.order_parts_objects:
                    production_order = 'OP_' + string_date_register + '_P' + str(counter_part)

                    info_row = {'Ordem Produção': production_order, 'Trabalho': task_name, 'Máquina': task_machine, 'Caminho': task_path, 'Cliente': part.client, 'Encomenda': part.order_num,
                                'Material': task_material_name, 'Liga (acabamento)': task_material_spec, 'Espessura': task_material_thickness,
                                'Proprietário M.P.': task_material_owner, 'Referência': part.name, 'Quantidade': str(part.quantity)}

                    df = df.append(info_row, ignore_index=True)

                    counter_part += 1

                counter_task += 1

                df.to_excel(CAMINHO_EXCEL, index=False)
                production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sidebar.pop_register_tasks.dismiss()

        except PermissionError:
            self.display_popup_warning("Folha de Excel aberta")


class Taskspage(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.orientation = 'horizontal'

        self.layout_sidebar = Taskspage_Sidebar()
        self.layout_sideframe = Taskpage_Sideframe()

        self.add_widget(self.layout_sidebar)
        self.add_widget(self.layout_sideframe)


class TaskPart(BoxLayout):
    # TODO Complete Task Part class
    def __init__(self, name, quantity, produced_quantity, client, order_num, order_num_client, additional_operations, rowid, is_selected=False, checkbox_mode=True, **kwargs):
        super().__init__(**kwargs)

        self.orientation = 'horizontal'

        self.name = name
        self.quantity = str(quantity)
        self.produced_quantity = str(produced_quantity)
        self.pending_quantity = str(int(self.quantity) - int(self.produced_quantity))
        self.client = client
        self.order_num = order_num
        self.order_num_client = order_num_client
        self.is_selected = is_selected
        self.checkbox_mode = checkbox_mode
        self.additional_operations = additional_operations
        self.rowid = rowid

        self.button_remove_task_part = Button(text='X', font_size='14', background_color=color_dark_red, size_hint_x=0.041)
        self.checkbox_selected = CheckBox(size_hint_x=0.028)
        self.button_name = Button(text=self.name[:66], font_size='14', background_color=color_dark_blue, size_hint_x=0.4)
        self.button_quantity = Button(text=self.pending_quantity, font_size='14', background_color=color_dark_blue, size_hint_x=0.05)
        self.button_client = Button(text=self.client[:30], font_size='14', background_color=color_dark_blue, size_hint_x=0.15)
        self.button_order_num = Button(text=self.order_num, font_size='14', background_color=color_dark_blue, size_hint_x=0.1)
        self.button_order_num_client = Button(text=self.order_num_client, font_size='14', background_color=color_dark_blue, size_hint_x=0.1)
        self.button_additional_operations = Button(text='', font_size='14', background_color=color_dark_blue, size_hint_x=0.05)

        self.add_widget(self.button_remove_task_part)
        self.add_widget(self.checkbox_selected)
        self.add_widget(self.button_name)
        self.add_widget(self.button_quantity)
        self.add_widget(self.button_client)
        self.add_widget(self.button_order_num)
        self.add_widget(self.button_order_num_client)
        self.add_widget(self.button_additional_operations)

        self.set_text_button_additional_operations()

    def set_text_button_additional_operations(self, *args):
        if self.additional_operations:
            self.button_additional_operations.text = 'Sim'
        else:
            self.button_additional_operations.text = 'Não'


class Task(BoxLayout):
    def __init__(self, current_path, original_path, machine, material_ref, notes, estimated_num_sheets, time, start_date, end_date, priority, order_parts, aggregated_tasks, rowid, is_selected=False, checkbox_mode=True, aggregated_index='', **kwargs):
        super().__init__(**kwargs)

        self.orientation = 'vertical'
        self.size_hint_y = None
        self.height = 45
        self.padding = (2, 2)

        self.current_path = current_path
        self.original_path = original_path
        self.machine = machine
        self.material_ref = material_ref
        self.material = get_material_name(self.material_ref)
        self.notes = notes
        self.estimated_num_sheets = str(round_number_if_integer(estimated_num_sheets))
        self.total_num_sheets = 0
        self.time = time #CASO TENHA PROBLEMAS VOLTAR A CONVERTER ISTO PARA STRING
        self.time_display_button = return_formatted_time_string(int(self.time))
        self.total_time = 0
        self.start_date = start_date
        self.end_date = end_date
        self.priority = str(priority)
        self.rowid = rowid
        self.is_selected = is_selected
        self.is_expanded = False
        self.checkbox_mode = checkbox_mode
        self.idle_color = color_light_blue

        self.order_parts_rowids = [int(order_part_rowid) for order_part_rowid in order_parts.split(',')]
        self.order_parts_objects = []

        self.aggregated_tasks_string = aggregated_tasks
        self.aggregated_tasks = [int(task_rowid) for task_rowid in aggregated_tasks.split(',')]
        self.aggregated_index = aggregated_index

        self.layout_task_info = BoxLayout(orientation='horizontal')
        self.layout_task_parts_header = BoxLayout(orientation='horizontal', height=40)
        self.layout_order_parts_info = BoxLayout(orientation='vertical', size_hint_y=0.01)

        # Layout Task Info Widgets
        self.checkbox_selected = CheckBox(size_hint_x=0.03)
        self.button_expand = RightClickableButton(text='+', background_color=color_light_blue, font_size='25', size_hint_x=0.05)
        self.button_current_path = RightClickableButton(text='Caminho\nAtual', font_size='15', halign='center', background_color=self.idle_color, size_hint_x=0.1)
        self.button_original_path = RightClickableButton(text='Caminho\nOriginal', font_size='15', halign='center', background_color=self.idle_color, size_hint_x=0.1)
        self.button_machine = RightClickableButton(text=self.machine, font_size='15', halign='center', background_color=self.idle_color, size_hint_x=0.05)
        self.button_material = RightClickableButton(text=self.material, font_size='15', halign='center', background_color=self.idle_color, size_hint_x=0.3)
        self.button_estimated_num_sheets = RightClickableButton(text=f'{str(self.estimated_num_sheets)} chapas', halign='center', background_color=self.idle_color, size_hint_x=0.08)
        self.button_time = RightClickableButton(text=self.time_display_button, font_size='15', halign='center', background_color=self.idle_color, size_hint_x=0.07)
        self.button_start_date = RightClickableButton(text=self.start_date, font_size='15', halign='center', background_color=self.idle_color, size_hint_x=0.1)
        self.button_end_date = RightClickableButton(text=self.end_date, font_size='15', halign='center', background_color=self.idle_color, size_hint_x=0.1)
        self.button_priority = RightClickableButton(text=str(self.priority), font_size='15', halign='center', background_color=self.idle_color, size_hint_x=0.05)
        self.button_notes = RightClickableButton(text='Notas', font_size='15', halign='center', background_color=self.idle_color, size_hint_x=0.05)

        self.checkbox_selected.bind(on_press=self.set_selected)
        self.button_current_path.bind(on_press=lambda button: self.choose_on_press_function('current_path', self.rowid, button))
        self.button_original_path.bind(on_press=lambda button: self.choose_on_press_function('original_path', self.rowid, button))
        self.button_machine.bind(on_press=lambda button: self.choose_on_press_function('machine', self.rowid, button))
        self.button_material.bind(on_press=lambda button: self.choose_on_press_function('material_name', self.rowid, button))
        self.button_estimated_num_sheets.bind(on_press=lambda button: self.choose_on_press_function('estimated_sheets_required', self.rowid, button))
        self.button_time.bind(on_press=lambda button: self.choose_on_press_function('estimated_time', self.rowid, button))
        self.button_start_date.bind(on_press=lambda button: self.choose_on_press_function('start_date', self.rowid, button))
        self.button_end_date.bind(on_press=lambda button: self.choose_on_press_function('end_date', self.rowid, button))
        self.button_priority.bind(on_press=lambda button: self.choose_on_press_function('priority', self.rowid, button))
        self.button_notes.bind(on_press=lambda button: self.choose_on_press_function('notes', self.rowid, button))


        # Layout Task Parts Header Widgets
        self.label_empty = Label(size_hint_x=0.041)
        self.checkbox_select_all = CheckBox(size_hint_x=0.028)
        self.label_name = Label(text='Nome', halign='center', size_hint_x=0.4)
        self.label_quantity = Label(text='Qtdd.', halign='center', size_hint_x=0.05)
        self.label_client = Label(text='Cliente', halign='center', size_hint_x=0.15)
        self.label_order_num = Label(text='Enc.\nInterna', halign='center', size_hint_x=0.1)
        self.label_order_num_client = Label(text='Enc.\nCliente', halign='center', size_hint_x=0.1)
        self.label_additional_operations = Label(text='Op.\nAdicionais', halign='center', size_hint_x=0.05)

        self.button_expand.bind(on_press=self.expand)

        self.layout_task_info.add_widget(self.button_expand)
        self.add_checkbox_widget()
        self.layout_task_info.add_widget(self.button_current_path)
        self.layout_task_info.add_widget(self.button_original_path)
        self.layout_task_info.add_widget(self.button_machine)
        self.layout_task_info.add_widget(self.button_material)
        self.layout_task_info.add_widget(self.button_estimated_num_sheets)
        self.layout_task_info.add_widget(self.button_time)
        self.layout_task_info.add_widget(self.button_start_date)
        self.layout_task_info.add_widget(self.button_end_date)
        self.layout_task_info.add_widget(self.button_priority)
        self.layout_task_info.add_widget(self.button_notes)
        self.add_widget(self.layout_task_info)
        self.add_widget(self.layout_order_parts_info)

        self.add_task_order_parts()
        self.set_idle_color_buttons()
        self.set_color_button('current_path')
        self.set_color_button('original_path')
        self.set_color_button('notes')
        self.set_total_time()
        self.set_text_button_time()
        self.set_total_num_sheets()
        self.set_text_button_num_sheets()
        self.set_text_buttons_paths()

    def expand(self, *args):
        self.is_expanded = not self.is_expanded

        if self.is_expanded:
            self.button_expand.text = '-'

            self.add_widget(self.layout_task_parts_header)

            self.height += 40
            self.layout_task_parts_header.add_widget(self.label_empty)
            self.layout_task_parts_header.add_widget(self.checkbox_select_all)
            self.layout_task_parts_header.add_widget(self.label_name)
            self.layout_task_parts_header.add_widget(self.label_quantity)
            self.layout_task_parts_header.add_widget(self.label_client)
            self.layout_task_parts_header.add_widget(self.label_order_num)
            self.layout_task_parts_header.add_widget(self.label_order_num_client)
            self.layout_task_parts_header.add_widget(self.label_additional_operations)

            for task_order_part in self.order_parts_objects:
                self.layout_order_parts_info.height += 35
                self.height += 40
                self.add_widget(task_order_part)
        else:
            self.button_expand.text = '+'

            self.remove_widget(self.layout_task_parts_header)

            self.height -= 40
            self.layout_task_parts_header.remove_widget(self.label_empty)
            self.layout_task_parts_header.remove_widget(self.checkbox_select_all)
            self.layout_task_parts_header.remove_widget(self.label_name)
            self.layout_task_parts_header.remove_widget(self.label_quantity)
            self.layout_task_parts_header.remove_widget(self.label_client)
            self.layout_task_parts_header.remove_widget(self.label_order_num)
            self.layout_task_parts_header.remove_widget(self.label_order_num_client)
            self.layout_task_parts_header.remove_widget(self.label_additional_operations)

            for task_order_part in self.order_parts_objects:
                self.layout_order_parts_info.height -= 35
                self.height -= 40
                self.remove_widget(task_order_part)

    def add_checkbox_widget(self):
        if self.checkbox_mode:
            self.layout_task_info.add_widget(self.checkbox_selected)

    def add_task_order_parts(self, *args):
        for task_order_part_rowid in self.order_parts_rowids:
            try:
                task_info = load_order_parts_from_database(mode='rowid', rowid=task_order_part_rowid)[0]
                top = TaskPart(task_info[1], task_info[3], task_info[4], task_info[7], task_info[5], task_info[6], task_info[10], task_order_part_rowid)
                self.order_parts_objects.append(top)
            except IndexError:
                print('Exception Ocurred')

    def set_selected(self, *args):
        self.is_selected = not self.is_selected

    def choose_on_press_function(self, field, rowid, button, *args):
        if button.mouse_button == 'left':
            pass
        elif button.mouse_button == 'right':
            self.display_popup_change_value(field, rowid)

    def display_popup_change_value(self, field, rowid):
        self.layout_new_value = BoxLayout(orientation='vertical')
        self.layout_horizontal_line = BoxLayout(orientation='horizontal', padding=(10, 10))
        self.label_new_value = Label(text='Novo valor: ', font_size='16', halign='right', size_hint_x=0.3)
        if field == 'current_path':
            self.widget_new_value = TextInput(text=self.current_path)
        elif field == 'original_path':
            self.widget_new_value = TextInput(text=self.original_path)
        elif field == 'machine':
            self.drp_machine = DropDown()
            self.widget_new_value = Button(text='Máquina', background_color=color_light_blue, padding=(5, 5))
            self.widget_new_value.bind(on_release=self.drp_machine.open)
            self.drp_machine.bind(on_select=lambda instance, new_val: setattr(self.widget_new_value, 'text', new_val))
            for machine_name in ['LF-3015', 'LC5']:
                button_machine_name = Button(text=machine_name, font_size='16', size_hint_y=None, height=30, background_color=color_dark_blue)
                button_machine_name.bind(on_release=lambda machine_object: self.drp_machine.select(machine_object.text))
                self.drp_machine.add_widget(button_machine_name)
        elif field == 'material_name':
            self.drp_material = DropDown()
            self.widget_new_value = Button(text='', font_size='16', halign='right', size_hint_y=0.75)
            self.widget_new_value.bind(on_release=self.drp_material.open)
            self.drp_material.bind(on_select=lambda instance, new_val: setattr(self.widget_new_value, 'text', new_val))

            for material in load_materials_from_database():
                material_type = material[1]
                material_spec = str(material[2])
                material_thickness = str(material[3])
                material_client = material[11]

                material_name = material_type + ' ' + material_spec + ', ' + material_thickness + 'mm ' + material_client

                button_material_name = Button(text=material_name, font_size='16', size_hint_y=None, height=30, background_color=color_dark_blue)
                button_material_name.bind(on_release=lambda material_object: self.drp_material.select(material_object.text))

                self.drp_material.add_widget(button_material_name)

        elif field == 'estimated_sheets_required':
            self.widget_new_value = TextInput(text=self.estimated_num_sheets)
        elif field == 'estimated_time':
            self.widget_new_value = TextInput(text=str(self.time))
        elif field == 'start_date':
            self.widget_new_value = TextInput(text=self.start_date)
        elif field == 'end_date':
            self.widget_new_value = TextInput(text=self.end_date)
        elif field == 'priority':
            self.widget_new_value = TextInput(text=self.priority)
        elif field == 'notes':
            self.widget_new_value = TextInput(text=self.notes)

        self.button_submit = Button(text='Submeter', font_size='18', background_color=color_light_green)

        self.layout_horizontal_line.add_widget(self.label_new_value)
        self.layout_horizontal_line.add_widget(self.widget_new_value)
        self.layout_new_value.add_widget(self.layout_horizontal_line)
        self.layout_new_value.add_widget(self.button_submit)

        self.button_submit.bind(on_press=lambda x: self.change_value(self.widget_new_value.text, field, rowid))

        self.pop_change_value = Popup(title='Novo valor: ', content=self.layout_new_value, size_hint=(0.50, 0.20))
        self.pop_change_value.bind(on_open=self.set_focus_text_input)
        self.pop_change_value.open()

    def change_value(self, value, field, rowid, *args):
        if field == 'current_path':
            for added_task in production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.added_tasks:
                for aggregated_task_rowid in self.aggregated_tasks:
                    if aggregated_task_rowid == added_task.rowid:
                        change_value_task_database(value, field, added_task.rowid)
                        added_task.current_path = str(value)
                        added_task.set_color_button(field)
        elif field == 'original_path':
            change_value_task_database(value, field, rowid)
            self.original_path = str(value)
            self.set_color_button(field)

            print(f'CURRENT PATH BUTTON TEXT IS {self.current_path}')

            if self.current_path == '':
                self.current_path = self.original_path
                self.set_color_button('current_path')
        elif field == 'machine':
            for added_task in production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.added_tasks:
                for aggregated_task_rowid in self.aggregated_tasks:
                    if aggregated_task_rowid == added_task.rowid:
                        change_value_task_database(value, field, added_task.rowid)

                        added_task.button_machine.text = value
                        added_task.machine = value
        elif field == 'material_name':
            for added_task in production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.added_tasks:
                for aggregated_task_rowid in self.aggregated_tasks:
                    if aggregated_task_rowid == added_task.rowid:
                        material_type = value.split(' ')[0]
                        material_spec = value.split(' ')[1].split(',')[0]
                        material_thickness = float(value.split(' ')[2].split('mm')[0])
                        material_client = value.split(' ')[3]
                        material_ref = get_material_ref(material_type, material_spec, material_thickness, material_client)
                        for order_part_rowid in added_task.order_parts_rowids:
                            change_value_order_part_database(material_ref, 'material_ref', order_part_rowid)

                        change_value_task_database(material_ref, 'material_ref', added_task.rowid)
                        material_name = get_material_name(material_ref)
                        added_task.button_material.text = material_name
                        added_task.material = material_name
        elif field == 'estimated_sheets_required':
            change_value_task_database(float(value), field, rowid)
            self.button_estimated_num_sheets.text = f'{str(round_number_if_integer(value))} chapas'
            self.estimated_num_sheets = str(round_number_if_integer(value))

            for added_task in production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.added_tasks:
                for aggregated_task_rowid in self.aggregated_tasks:
                    if aggregated_task_rowid == added_task.rowid:
                        added_task.set_total_num_sheets()
                        added_task.set_text_button_num_sheets()
        elif field == 'estimated_time':
            change_value_task_database(int(value), field, rowid)
            self.button_time.text = return_formatted_time_string(int(value))
            self.time_display_button = self.button_time.text
            self.time = int(value)

            print(f'TIME == {self.time}')

            for added_task in production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.added_tasks:
                for aggregated_task_rowid in self.aggregated_tasks:
                    if aggregated_task_rowid == added_task.rowid:
                        added_task.set_total_time()
                        added_task.set_text_button_time()

            production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sidebar.set_total_time_machines()
        elif field == 'start_date':
            for added_task in production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.added_tasks:
                for aggregated_task_rowid in self.aggregated_tasks:
                    if aggregated_task_rowid == added_task.rowid:
                        change_value_task_database(value, field, added_task.rowid)

                        added_task.button_start_date.text = str(value)
                        added_task.start_date = str(value)
        elif field == 'end_date':
            for added_task in production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.added_tasks:
                for aggregated_task_rowid in self.aggregated_tasks:
                    if aggregated_task_rowid == added_task.rowid:
                        change_value_task_database(value, field, added_task.rowid)
                        added_task.button_end_date.text = str(value)
                        added_task.end_date = str(value)
        elif field == 'priority':
            for added_task in production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.added_tasks:
                for aggregated_task_rowid in self.aggregated_tasks:
                    if aggregated_task_rowid == added_task.rowid:
                        change_value_task_database(int(value), field, added_task.rowid)
                        added_task.button_priority.text = str(value)
                        added_task.priority = str(value)
        elif field == 'notes':
            change_value_task_database(value, field, rowid)
            self.notes = value
            self.set_color_button(field)
        elif field == 'aggregated_tasks':
            change_value_task_database(value, field, rowid)
            self.aggregated_tasks = [int(task_rowid) for task_rowid in value.split(',')]

        try:
            self.pop_change_value.dismiss()
        except AttributeError:
            pass

    def set_focus_text_input(self, *args): # Return to this later, autoselect and focus all textinput widgets
        try:
            self.widget_new_value.focus = True
            self.widget_new_value.select_all()
        except AttributeError:
            pass

    def set_color_button(self, field, *args):
        if field == 'current_path':
            if self.current_path == '':
                self.button_current_path.background_color = self.idle_color
            else:
                self.button_current_path.background_color = color_light_green
        elif field == 'original_path':
            if self.original_path == '':
                self.button_original_path.background_color = self.idle_color
            else:
                self.button_original_path.background_color = color_light_green
        elif field == 'notes':
            if self.notes == '':
                self.button_notes.background_color = self.idle_color
            elif self.notes.lower().replace(' ', '') == 'folha':
                self.button_notes.background_color = color_light_green
            else:
                self.button_notes.background_color = color_orange

    def set_idle_color_buttons(self, *args):
        if len(self.aggregated_tasks) > 1:
            self.idle_color = color_light_black
        else:
            self.idle_color = color_light_blue

        self.button_expand.background_color = self.idle_color
        self.button_current_path.background_color = self.idle_color
        self.button_original_path.background_color = self.idle_color
        self.button_machine.background_color = self.idle_color
        self.button_material.background_color = self.idle_color
        self.button_estimated_num_sheets.background_color = self.idle_color
        self.button_time.background_color = self.idle_color
        self.button_start_date.background_color = self.idle_color
        self.button_end_date.background_color = self.idle_color
        self.button_priority.background_color = self.idle_color
        self.button_notes.background_color = self.idle_color

    def update_info_buttons(self):
        self.button_machine.text = self.machine
        self.button_material.text = self.material
        self.button_estimated_num_sheets.text = str(self.estimated_num_sheets)
        self.button_time.text = str(self.time)
        self.button_start_date.text = self.start_date
        self.button_end_date.text = self.end_date
        self.buuton_priority.text = str(self.priority)

    def set_total_time(self):
        try:
            total_time = 0
            for aggregated_task_rowid in self.aggregated_tasks:
                aggregated_task = load_tasks_from_database(aggregated_task_rowid)[0]

                total_time += aggregated_task[6]

            self.total_time = total_time
        except IndexError:
            pass

    def set_text_button_time(self, *args):
        if len(self.aggregated_tasks) > 1:
            total_time_string = return_formatted_time_string(self.total_time)

            self.button_time.text = f'{self.time_display_button}\n{total_time_string} total'
        else:
            self.time_display_button = return_formatted_time_string(self.total_time)
            self.button_time.text = self.time_display_button

    def set_total_num_sheets(self, *args):
        try:
            total_num_sheets = 0
            for aggregated_task_rowid in self.aggregated_tasks:
                aggregated_task = load_tasks_from_database(aggregated_task_rowid)[0]

                total_num_sheets += aggregated_task[5]

            self.total_num_sheets = total_num_sheets
        except IndexError:
            pass

    def set_text_button_num_sheets(self, *args):
        if len(self.aggregated_tasks) > 1:
            self.button_estimated_num_sheets.text = f'{self.estimated_num_sheets} chapas\n{str(round_number_if_integer(self.total_num_sheets))} chapas total'
        else:
            self.button_estimated_num_sheets.text = f'{self.estimated_num_sheets} chapas'

    def set_aggregated_index(self, *args):
        for key in production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.aggregated_tasks_dict:
            # print(f'ROWID STRING IN DICT = {production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.aggregated_tasks_dict[key]}')
            # print(f'TASK ROWID STRING {self.aggregated_tasks_string}')
            if production_planning.homepage.layout_sidebar.layout_popup_tasks.layout_sideframe.aggregated_tasks_dict[key] == self.aggregated_tasks_string:
                self.aggregated_index = key
                self.button_current_path.text = f'Caminho\nAtual [{key}]'
                self.button_original_path.text = f'Caminho\nOriginal [{key}]'

                # print(f'THE SELECT KEY WAS {key}')

                change_value_task_database(key, 'aggregated_index', self.rowid)

    def set_text_buttons_paths(self, *args):
        if len(self.aggregated_tasks) > 1:
            self.button_current_path.text = f'Caminho\nAtual [{self.aggregated_index}]'
            self.button_original_path.text = f'Caminho\nOriginal [{self.aggregated_index}]'
        else:
            self.button_current_path.text = 'Caminho\nAtual'
            self.button_original_path.text = 'Caminho\nOriginal'


class Homepage_Task_List(ScrollView):
    def __init__(self, sort_by='order', **kwargs):
        super().__init__(**kwargs)

        self.layout_scroll = BoxLayout(orientation='vertical', padding=10, spacing=(5,5), size=(10,10), size_hint_y=None)
        self.layout_scroll.bind(minimum_height=self.layout_scroll.setter('height'))

        self.added_tasks = []
        self.sort_by = sort_by

        self.add_widget(self.layout_scroll)


class Homepage_Sideframe(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.orientation = 'vertical'

        # self.layout_filters_bar = Taskpage_Filters_Bar()
        # self.layout_task_list = Homepage_Task_List()
        #
        # self.add_widget(self.layout_filters_bar)
        # self.add_widget(self.layout_task_list)


class Homepage_Sidebar(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.orientation = 'vertical'
        self.size_hint_x = 0.08

        # Creating the buttons
        self.button_tasks = Button(text='Tarefas', font_size='16', background_color=color_light_blue)
        self.button_orders = Button(text='Encomendas', font_size='16', background_color=color_light_blue)
        self.button_parts = Button(text='Artigos', font_size='16', background_color=color_light_blue)
        self.button_materials = Button(text='Materiais', font_size='16', background_color=color_light_blue)
        self.button_shifts = Button(text='Definir\nTurnos', font_size='16', background_color=color_light_blue)
        self.label1 = Label()
        self.label2 = Label()
        self.label3 = Label()
        self.label4 = Label()

        # Binding the functions to the buttons
        self.button_tasks.bind(on_press=self.display_popup_tasks)
        self.button_orders.bind(on_press=self.display_popup_order_parts)
        self.button_parts.bind(on_press=self.display_popup_parts)
        self.button_materials.bind(on_press=self.display_popup_materials)
        self.button_shifts.bind(on_press=self.display_popup_shifts)

        self.add_widget(self.button_tasks)
        self.add_widget(self.button_orders)
        self.add_widget(self.button_parts)
        self.add_widget(self.button_materials)
        self.add_widget(self.button_shifts)
        self.add_widget(self.label1)
        self.add_widget(self.label2)
        self.add_widget(self.label3)
        self.add_widget(self.label4)
        # self.add_widget(self.button_consumables)
        # self.add_widget(self.button_machines)
        # self.add_widget(self.button_statistics)
        # self.add_widget(self.button_settings)

    def display_popup_tasks(self, *args):
        self.layout_popup_tasks = Taskspage()
        self.pop_tasks = Popup(title='Tarefas', content=self.layout_popup_tasks, size_hint=(0.95, 0.95), auto_dismiss=False)

        self.pop_tasks.open()

    def display_popup_order_parts(self, *args):
        self.layout_popup_order_parts = OrderPartspage()
        self.pop_order_parts = Popup(title='Encomendas', content=self.layout_popup_order_parts, size_hint=(0.95, 0.95), auto_dismiss=False)

        self.pop_order_parts.open()

    def display_popup_parts(self, *args):
        self.layout_popup_parts = Partspage()
        self.pop_parts = Popup(title='Artigos', content=self.layout_popup_parts, size_hint=(0.95, 0.95), auto_dismiss=False)

        self.pop_parts.open()

    def display_popup_clients(self, *args):
        self.layout_popup_clients = Clientspage()
        self.pop_clients = Popup(title='Clientes', content=self.layout_popup_clients, size_hint=(0.95, 0.95))

        self.pop_clients.open()

    def display_popup_materials(self, *args):
        self.layout_popup_materials = Materialspage()
        self.pop_materials = Popup(title='Materiais', content=self.layout_popup_materials, size_hint=(0.95, 0.95), auto_dismiss=False)

        self.pop_materials.open()

    def display_popup_shifts(self, *args):
        self.layout_popup_shifts = Shiftspage()
        self.pop_shifts = Popup(title='Definir Turnos', content=self.layout_popup_shifts, size_hint=(0.95, 0.85), auto_dismiss=False)

        self.pop_shifts.open()


class Homepage(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.orientation = 'horizontal'
        self.padding = 10

        self.layout_sidebar = Homepage_Sidebar()
        self.layout_sideframe = Homepage_Sideframe()

        self.add_widget(self.layout_sidebar)
        self.add_widget(self.layout_sideframe)


class ProductionPlanning(App):
    def build(self):
        self.title = 'Planeamento de Produção'

        self.screen_manager = ScreenManager(transition=NoTransition())

        self.create_homepage()

        return self.screen_manager

    def create_homepage(self):
        self.homepage = Homepage()
        screen = Screen(name='Homepage')
        screen.add_widget(self.homepage)
        self.screen_manager.add_widget(screen)


class RightClickableButton(Button):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

        self.mouse_button = None
        self.bind(on_touch_down = self.callback_touch_down)

    def callback_touch_down(self, instance, touch):
        self.mouse_button = touch.button


class AutoSelectTextInput(TextInput):
    def on_open(self, widget, parent):
        self.focus = True


class DailyTasksSheet:
    def __init__(self, machine, date, parallel_tasks, general_notes, tasks):
        self.machine = machine
        self.date = date
        self.parallel_tasks = parallel_tasks
        self.general_notes = general_notes

        self.tasks = tasks
        self.tasks_to_be_added = sorted(self.return_list_tasks_to_be_added(), key=operator.attrgetter('priority'), reverse=True)

    def build_sheet(self):
        document = FPDF(orientation='L')
        document.add_page()

        effective_page_width = document.w - 2 * document.l_margin
        col_width = effective_page_width / 7

        document.set_font('Times', 'B', size=20)
        document.cell(effective_page_width, document.font_size, 'Ordem de Trabalhos', align='C')

        document.ln(15)

        document.set_font('Times', 'B', size=17)
        self.machine = 'LC5'
        document.cell(col_width * 3.5, document.font_size, f'Máquina: {self.machine}', align='C')
        document.cell(col_width * 3.5, document.font_size, f'Data: {datetime.datetime.now().day}/{datetime.datetime.now().month}/{datetime.datetime.now().year}', align='C')

        document.ln(15)

        task_index = 1
        for task in self.tasks_to_be_added:
            document.set_font('Times', 'B', size=13)
            order_parts = task.order_parts_objects

            task_id = f'T{task_index}'
            task_path = task.current_path
            task_material = task.material
            task_total_qqty = self.return_string_total_quantity(order_parts)
            task_notes = task.notes

            document.cell(col_width * 0.25, document.font_size, task_id, align='C', border=1)
            document.cell(col_width * 2.75, document.font_size, task_path, align='C', border=1)
            document.cell(col_width * 2.75, document.font_size, task_material, align='C', border=1)
            document.cell(col_width * 1.25, document.font_size, task_total_qqty, align='C', border=1)


            document.ln(document.font_size)

            document.set_font('Times', 'B', size=12)

            order_part_index = 1
            for order_part in order_parts:
                order_part_id = f'P{order_part_index}'
                order_part_name = order_part.name
                order_part_qtty = f'{order_part.quantity} Unid.'
                order_part_client = order_part.client
                order_part_additional_operations = self.return_value_additional_operations(order_part.additional_operations)

                document.cell(col_width * 0.5, document.font_size, order_part_id, align='C', border=1)
                document.cell(col_width * 4, document.font_size, order_part_name, align='C', border=1)
                document.cell(col_width * 1, document.font_size, order_part_qtty, align='C', border=1)
                document.cell(col_width * 1, document.font_size, order_part_client, align='C', border=1)
                document.cell(col_width * 0.5, document.font_size, order_part_additional_operations, align='C', border=1)

                document.ln(document.font_size)

                order_part_index += 1

            task_index += 1

        document.output(f'{datetime.datetime.now().year}_{datetime.datetime.now().month}_{datetime.datetime.now().day}_Ord_Trab.pdf', 'F')

    def return_list_tasks_to_be_added(self):
        rowids_tasks_to_be_added = []

        for task in self.tasks:
            if task.is_selected:
                if task.rowid not in rowids_tasks_to_be_added:
                    rowids_tasks_to_be_added.append(task.rowid)
                    for aggregated_task_rowid in task.aggregated_tasks:
                        if aggregated_task_rowid not in rowids_tasks_to_be_added:
                            rowids_tasks_to_be_added.append(aggregated_task_rowid)

        return self.group_tasks_same_index([task for task in self.tasks if task.rowid in rowids_tasks_to_be_added])

    def group_tasks_same_index(self, tasks):
        indexes = []
        for task in tasks:
            if task.aggregated_index != '' and task.aggregated_index not in indexes:
                indexes.append(task.aggregated_index)

        for index in indexes:
            grouped_order_parts = []
            current_path = ''
            original_path = current_path
            machine = ''
            material_ref = ''
            notes = ''
            estimated_num_sheets = 0
            time = 0
            start_date = ''
            end_date = ''
            priority = 0
            rowid = max([task.rowid for task in tasks]) + 1
            aggregated_tasks_string = str(rowid)

            tasks_to_be_deleted = []
            for task in tasks:
                if task.aggregated_index == index:
                    grouped_order_parts += task.order_parts_objects
                    current_path = task.current_path
                    original_path = current_path
                    machine = task.machine
                    material_ref = task.material_ref
                    estimated_num_sheets = task.estimated_num_sheets
                    time = task.total_time
                    notes += task.notes + '\n'
                    start_date = task.start_date
                    end_date = task.end_date
                    priority = task.priority

                    tasks_to_be_deleted.append(task)

            grouped_order_parts_string = ','.join([str(order_part.rowid) for order_part in grouped_order_parts])

            temp_task = Task(current_path, original_path, machine, material_ref, notes, estimated_num_sheets, time, start_date, end_date, priority,
                             grouped_order_parts_string, aggregated_tasks_string, rowid)

            for task_to_be_deleted in tasks_to_be_deleted:
                if task_to_be_deleted in tasks:
                    del tasks[tasks.index(task_to_be_deleted)]

            tasks.append(temp_task)

            # print([(task.rowid, task.order_parts_rowids) for task in tasks])

            return tasks

    def calculate_total_num_rows(self):
        total_num_tasks_rows = len(self.tasks_to_be_added)
        total_num_order_parts_rows = 0
        total_num_header_rows = 0

        for task in self.tasks_to_be_added:
            total_num_order_parts_rows += len(task.order_parts_rowids)
            total_num_header_rows += 2

        print(f'TOTAL TASKS = {total_num_tasks_rows}')
        print(f'TOTAL NUM ORDER PARTS = {total_num_order_parts_rows}')

        return total_num_tasks_rows + total_num_order_parts_rows + total_num_header_rows + 3

    def return_string_total_quantity(self, order_parts):
        string_total_quantity = ''

        order_part_clients = []

        for order_part in order_parts:
            if order_part.client not in order_part_clients:
                order_part_clients.append(order_part.client)

        client_quantity_dict = {}

        for order_part_client in order_part_clients:
            client_quantity_dict.update({order_part_client: 0})

        for order_part_client in order_part_clients:
            for order_part in order_parts:
                if order_part_client == order_part.client:
                    client_quantity_dict.update({order_part_client: client_quantity_dict[order_part_client] + int(order_part.quantity)})

        for order_part_client in client_quantity_dict:
            string_total_quantity += f'{client_quantity_dict[order_part_client]} Unid. Total {order_part_client}\n'

        return string_total_quantity

    def return_value_additional_operations(self, additional_operations):
        if additional_operations:
            return 'Sim'
        else:
            return 'Não'


class PopupWarningMessage(Popup):
    def __init__(self, message, type='warning', continue_func='', **kwargs):
        super().__init__(**kwargs)

        self.title = 'Aviso'
        self.message = message
        self.type = type

        self.continue_func = continue_func

        self.size_hint = (0.3, 0.2)

        self.build()

    def build(self):
        if self.type == 'warning':
            self.layout_content = BoxLayout(orientation='vertical')

            self.label_message = Label(text=self.message)
            self.button_ok = Button(text='OK', background_color=color_light_green)

            self.button_ok.bind(on_press=self.dismiss)

            self.layout_content.add_widget(self.label_message)
            self.layout_content.add_widget(self.button_ok)

            self.content = self.layout_content
        elif self.type == 'choice':
            self.layout_content = BoxLayout(orientation='vertical')
            self.layout_buttons = BoxLayout(orientation='horizontal')

            self.label_message = Label(text=self.message)
            self.button_continue = Button(text='Continuar', background_color=color_light_green)
            self.button_cancel = Button(text='Cancelar', background_color=color_dark_red)

            if self.continue_func == '':
                self.button_continue.bind(on_press=self.dismiss)
            else:
                self.button_continue.bind(on_press=self.wrapper_continue_function)
            self.button_cancel.bind(on_press=self.dismiss)

            self.layout_content.add_widget(self.label_message)
            self.layout_content.add_widget(self.layout_buttons)
            self.layout_buttons.add_widget(self.button_cancel)
            self.layout_buttons.add_widget(self.button_continue)

            self.content = self.layout_content

    def wrapper_continue_function(self, *args):
        self.continue_func()
        self.dismiss()


if __name__ == '__main__':

    CAMINHO_EXCEL = "D:/PROJETOS/PROGRAMMING/PROJECTS/PROJETOS PYTHON/Projetos Trabalho/SOFTWARE PLANEAMENTO/REGISTOS/registo.xlsx"

    color_dark_gray = convert_rgb_to_kivy_float((128, 128, 128))
    color_light_gray = convert_rgb_to_kivy_float((224, 224, 224))
    color_dark_blue = convert_rgb_to_kivy_float((0, 102, 204))
    color_light_blue = convert_rgb_to_kivy_float((153, 204, 255))
    color_light_yellow = convert_rgb_to_kivy_float((255, 255, 204))
    color_orange = convert_rgb_to_kivy_float((255, 178, 102))
    color_dark_green = convert_rgb_to_kivy_float((0, 102, 0))
    color_light_green = convert_rgb_to_kivy_float((178, 255, 102))
    color_dark_red = convert_rgb_to_kivy_float((255, 51, 51))
    color_light_black = convert_rgb_to_kivy_float((30, 30, 30))

    Window.clearcolor = color_light_black

    conn = sqlite3.connect('Database.db', isolation_level=None)
    cursor = conn.cursor()

    production_planning = ProductionPlanning()
    production_planning.run()

    conn.close()
